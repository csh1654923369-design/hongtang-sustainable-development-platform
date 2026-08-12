from pathlib import Path
import os
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
SCREENSHOT = QA_ROOT / "home-1440.png"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v117-refreshed.docx"


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def find_caption(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Caption not found: {prefix}")


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


if not SCREENSHOT.exists():
    raise FileNotFoundError(SCREENSHOT)

document = Document(MANUAL)
if "V1.17 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.17 manual, got: {document.paragraphs[5].text}")

caption = find_caption(document, "图 1")
image_paragraph = image_paragraph_before(document, caption)
clear_paragraph(image_paragraph)
image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
shape = image_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.3))
shape._inline.docPr.set("name", "图 1 3D高斯实景首页与五类示例点")
shape._inline.docPr.set(
    "descr",
    "红塘村3D高斯实景首页，显示五类示例点、左键平移和滚轮缩放灵敏度设置",
)

document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 19]
assert "V1.17 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "左键拖动沿村庄水平面平移" in all_text
assert "右键拖动围绕当前中心点旋转" in all_text
assert "中键拖动已停用" in all_text

os.replace(OUTPUT, MANUAL)
print(f"Refreshed V1.17 homepage image: {MANUAL}")
