from copy import deepcopy
from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v110.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.9.backup.docx"
MAP_SCREENSHOT = QA_ROOT / "map-aerial-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption_text):
    caption = find_paragraph(document, caption_text)
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError(f"Image paragraph not found before caption: {caption_text}")


def set_image(paragraph, image_path, width, name, description):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("name", name)
    shape._inline.docPr.set("descr", description)


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        copy_cell_format(template_row.cells[index], cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if template_row.cells[index].paragraphs and cell.paragraphs:
            cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
            cell.paragraphs[0].paragraph_format.alignment = template_row.cells[index].paragraphs[0].paragraph_format.alignment
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        east_asia = "SimHei" if is_heading or index == 0 else "SimSun"
        for run in paragraph.runs:
            set_run_font(run, east_asia)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not MAP_SCREENSHOT.exists():
    raise FileNotFoundError(MAP_SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
version_text = document.paragraphs[5].text

if "V1.10 Demo" in version_text:
    set_image(
        image_paragraph_before(document, "图 3  村里一张图与点位详情"),
        MAP_SCREENSHOT,
        6.3,
        "图 3 村里一张图与点位分类",
        "无人机影像上的村里具体事项、行动办理、互助资源和调研资料点位",
    )
    normalize_fonts(document)
    document.save(OUTPUT)
    os.replace(OUTPUT, MANUAL)
    print(f"Refreshed V1.10 map screenshot: {MANUAL}")
    raise SystemExit(0)

if "V1.9 Demo" not in version_text:
    raise ValueError(f"Expected V1.9 manual, got: {version_text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = "版本：V1.10 Demo   |   更新日期：2026年7月22日   |   适用地址：http://localhost:3000"
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页和“村里一张图”使用红塘村无人机正射影像；"
    "地图已增加小花园、茶场、茶厂、村里用水设施、光伏设施、安全隐患和村庄记忆等事项类型。"
    "业务点位的名称与位置仍含演示数据，不能作为真实设施、风险或产权判断依据。建筑调研底图已去除个人字段，仅用于规划讨论。"
)

find_paragraph(
    document,
    "村里一张图：从首页影像、右上角地图按钮或手机底栏进入，查看完整影像、地点和业务记录。",
).text = "村里一张图：从首页影像、右上角地图按钮或手机底栏进入，按具体事项、行动办理、互助资源和调研资料查找点位。"

find_paragraph(
    document,
    "“村里一张图”分为三部分：上方地图可在简化示意图和红塘村无人机正射影像之间切换，并叠加问题、项目、社区微行动和资源供需点位；中部互助板展示可提供资源与需求；下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。当前业务点位仍是演示位置，尚未与真实影像完成校准。",
).text = (
    "“村里一张图”分为三部分：上方地图可在简化示意图和红塘村无人机正射影像之间切换。"
    "左侧点位分为“村里的具体事项、行动与办理、互助资源、公共空间与调研资料”四组；"
    "中部互助板展示可提供资源与需求；下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。"
    "当前业务点位仍是演示位置，尚未与真实影像完成校准。"
)

set_image(
    image_paragraph_before(document, "图 3  村里一张图与点位详情"),
    MAP_SCREENSHOT,
    6.3,
    "图 3 村里一张图与点位分类",
    "无人机影像上的村里具体事项、行动办理、互助资源和调研资料点位",
)

find_paragraph(
    document,
    "在左侧按“行动与问题、社区互助资源、空间与调研资料”勾选或取消图层。",
).text = (
    "先在左侧查看“村里的具体事项”：小花园、茶场、茶厂、村里用水设施、光伏设施、安全隐患和村庄记忆；"
    "也可在“行动与办理、互助资源、公共空间与调研资料”中勾选或取消其他图层。"
)
find_paragraph(
    document,
    "按待轻量核对、招募伙伴、试验中、可提供、需求中等状态，以及所属目标、公众参与或“与我有关”继续筛选。",
).text = "还可按持续记录、资料待补充、待实地确认、持续巡查、运行观察、办理状态、所属目标、公众参与或“与我有关”继续筛选。"
find_paragraph(
    document,
    "点击地图标记，右侧会显示点位类型、状态、位置、更新时间和关联内容。",
).text = "点击地图标记，右侧会显示点位类型、状态、位置、更新时间和关联内容；不同类型使用不同颜色和图标。"
find_paragraph(
    document,
    "点击详情按钮进入相关问题、项目或微行动区；资源点位可提交回应意向。",
).text = "小花园、茶场、茶厂、用水、光伏、安全和村庄记忆点位可进入各自事项页面；问题、项目或微行动进入对应详情，资源点位可提交回应意向。"

map_row = next(row for row in document.tables[1].rows if row.cells[0].text.strip() == "村里一张图")
map_row.cells[2].text = "四组点位筛选、示意图/无人机影像切换、事项详情跳转、互助资源和建筑调研底图"

version_table = document.tables[4]
new_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    new_row.cells,
    [
        "V1.10 Demo",
        "2026-07-22",
        "村里一张图的点位重分为四组；新增小花园、茶场、茶厂、村里用水设施、光伏设施和安全巡查演示点，村庄记忆改为独立类型，并接入各事项详情入口。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert "V1.10 Demo" in check.paragraphs[5].text
assert len(check.tables[4].rows) == 12
assert any("村里的具体事项" in paragraph.text and "茶场" in paragraph.text for paragraph in check.paragraphs)
assert any(row.cells[0].text.strip() == "村里一张图" and "四组点位筛选" in row.cells[2].text for row in check.tables[1].rows)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.10: {MANUAL}")
