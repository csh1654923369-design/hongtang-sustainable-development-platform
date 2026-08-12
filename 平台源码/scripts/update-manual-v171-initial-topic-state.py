from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.71 ｜ 更新日期：2026年8月10日",
)

topic_intro = next(p for p in document.paragraphs if p.text.startswith("桌面端和手机端都在地图左上角"))
startup_note = (
    "首次打开首页时，五个专题会在第一次显示时直接保持全选；地点和用水数据尚在读取期间，"
    "平台不会短暂显示“当前没有显示要素”等空状态提示。只有数据读取完成且筛选结果确实为空时，才显示相应提示。"
)
if startup_note not in topic_intro.text:
    replace_paragraph_text(topic_intro, f"{topic_intro.text}{startup_note}")

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.71"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.71",
    "2026年8月10日",
    "修复首次打开时短暂出现无专题或无要素提示的问题；五个专题从首帧起默认全选，空状态在数据读取完成后再判断。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.71") for p in reloaded.paragraphs)
assert "五个专题会在第一次显示时直接保持全选" in paragraph_text
assert "只有数据读取完成且筛选结果确实为空时" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.71"
print(MANUAL)
