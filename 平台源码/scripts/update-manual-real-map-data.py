from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.V1.29.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.28.backup.docx"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def copy_cell_format(source_cell, target_cell):
    target_properties = target_cell._tc.get_or_add_tcPr()
    for child in list(target_properties):
        target_properties.remove(child)
    for child in source_cell._tc.tcPr:
        target_properties.append(deepcopy(child))


def set_cell_text_like(source_cell, target_cell, text):
    copy_cell_format(source_cell, target_cell)
    target_cell.text = text
    if source_cell.paragraphs and target_cell.paragraphs:
        target_cell.paragraphs[0].style = source_cell.paragraphs[0].style
        target_cell.paragraphs[0].paragraph_format.alignment = (
            source_cell.paragraphs[0].paragraph_format.alignment
        )


def add_row_like(table, template_row, values):
    row = table.add_row()
    for index, value in enumerate(values):
        set_cell_text_like(template_row.cells[index], row.cells[index], value)
    return row


def keep_table_rows_intact(table):
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        if not properties.xpath("./w:cantSplit"):
            properties.append(OxmlElement("w:cantSplit"))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    if not header_properties.xpath("./w:tblHeader"):
        header_properties.append(OxmlElement("w:tblHeader"))


QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
version_text = document.paragraphs[5].text
if "V1.29 Demo" in version_text:
    print("Manual is already V1.29; no changes applied.")
    raise SystemExit(0)
if "V1.28 Demo" not in version_text:
    raise ValueError(f"Expected V1.28 manual, got: {version_text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.29 Demo   |   更新日期：2026年7月29日   |   "
    "适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n"
    "当前网站是可交互演示版本。正式首页使用SparkJS 2.1.0按当前视角加载红塘村轻量化3D高斯模型，"
    "首页的5个三维图钉仍为演示位置。“村里一张图”已经直接使用平台素材中的真实手绘图、56个POI、"
    "93组村景记录和582张现场照片，点位按素材中的经纬度显示；下方建筑底图使用1490栋真实建筑的"
    "脱敏GeoJSON。问题、项目、微行动、互助资源、指标和3D图钉仍含演示数据。"
)

find_paragraph(
    document,
    "以后不必每次输入命令。项目根目录保留“启动网站.bat”和“关闭网站.bat”；网页源代码集中放在“平台源码”文件夹，3D模型、无人机原始影像和地图原始数据集中放在“平台素材”文件夹。3D首页默认从“平台素材/3D高斯展示/轻量化模型”读取模型。",
).text = (
    "以后不必每次输入命令。项目根目录保留“启动网站.bat”和“关闭网站.bat”；网页源代码集中放在"
    "“平台源码”文件夹，3D模型、无人机原始影像、真实POI表格、照片清单和地图服务原始资料集中放在"
    "“平台素材”文件夹。3D首页默认从“平台素材/3D高斯展示/轻量化模型”读取模型。"
)

find_paragraph(
    document,
    "小花园、茶厂、村里用水：进入后首先使用同一张村庄地图，并分别只显示小花园、茶场与茶厂、用水设施点位。",
).text = (
    "小花园、茶厂、村里用水：进入后首先使用同一张真实手绘图，并分别显示35个真实花园、"
    "9个真实茶厂和2个真实用水设施点位；点击点位可查看素材中的现场照片、简介和坐标。"
)
find_paragraph(
    document,
    "村里一张图：从“村庄总览”的无人机影像、右上角地图按钮或手机底栏进入，按具体事项、行动办理、互助资源和调研资料查找点位。",
).text = (
    "村里一张图：从“村庄总览”的无人机影像、右上角地图按钮或手机底栏进入；默认打开真实手绘图，"
    "可切换无人机影像或示意图，并按具体事项、行动办理、互助资源和调研资料查找点位。"
)
find_paragraph(
    document,
    "系统会进入该事项自己的网址，页面第一块内容就是红塘村无人机地图；切换顶部标签时，底图不变，点位自动切换为新事项。",
).text = (
    "系统会进入该事项自己的网址，页面第一块内容就是红塘村真实手绘图；切换顶部标签时，地图框架不变，"
    "点位自动切换为新事项。"
)
find_paragraph(
    document,
    "地图右侧会说明当前只显示哪些点位；点击标记可查看位置、状态和简介。茶厂页面同时显示茶场与茶厂，其他五页只显示各自事项类型。",
).text = (
    "地图右侧会说明当前只显示哪些点位；点击标记可查看位置、状态、简介和现场照片。小花园、茶厂和用水页"
    "优先显示素材中的真实记录；当前素材没有对应点位的光伏、安全和村庄记忆页仍明确标注为演示点位。"
)
find_paragraph(
    document,
    "使用提示\n六个事项共用同一地图底图和点位交互，便于村民保持空间认识；地图下方的功能仍按事项分别设计。所有点位位置目前都是演示数据，需在实地核对后替换。",
).text = (
    "使用提示\n六个事项共用同一套地图框架，便于村民保持空间认识；地图下方的功能仍按事项分别设计。"
    "小花园、茶厂和用水设施已经使用真实素材点位；光伏、安全、村庄记忆及互动业务点位仍为演示数据。"
)

find_paragraph(
    document,
    "“村里一张图”分为三部分：上方地图可在简化示意图和红塘村无人机正射影像之间切换。左侧点位分为“村里的具体事项、行动与办理、互助资源、公共空间与调研资料”四组；中部互助板展示可提供资源与需求；下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。当前业务点位仍是演示位置，尚未与真实影像完成校准。",
).text = (
    "“村里一张图”分为三部分：上方地图默认显示真实手绘图，并可切换红塘村无人机正射影像或简化示意图。"
    "真实POI和村景记录按经纬度显示，点击可轮播现场照片；问题、项目、微行动和互助资源继续作为演示互动"
    "内容显示。中部互助板展示可提供资源与需求；下方建筑调研底图使用1490栋脱敏真实建筑。"
)
find_paragraph(
    document,
    "影像与点位说明\n无人机原始影像保存在“平台素材/Production_1-tif”文件夹；网站展示使用“平台源码/public/hongtang-orthophoto”中的网页切片。地图点位和3D图钉仍为演示数据，待后续实地核对后再替换。",
).text = (
    "底图与点位说明\n真实手绘图和建筑WFS保存在“平台素材/地图服务素材”；POI与照片记录保存在"
    "“平台素材”的Excel和CSV文件中。网站使用生成到“平台源码/public/data”的本地底图、真实点位和脱敏"
    "建筑文件。照片继续使用素材表格中的sannongdata.cn网址，因此查看现场照片需要联网。3D图钉仍为演示位置。"
)
find_paragraph(
    document, "使用地图右上角的“示意图 / 无人机影像”按钮切换底图。"
).text = "使用地图右上角的“手绘图 / 无人机影像 / 示意图”按钮切换底图；默认打开真实手绘图。"
find_paragraph(
    document,
    "切换到无人机影像后，仍可点击原有点位；页面左下角会持续提示“点位尚未按真实位置校准”。",
).text = (
    "真实POI和村景点位按经纬度定位；切换无人机影像时，超出当前影像范围的真实点位会自动隐藏。"
    "互动业务的演示或隐私安全点位仍可继续使用。"
)
find_paragraph(
    document,
    "点击地图标记，右侧会显示点位类型、状态、位置、更新时间和关联内容；不同类型使用不同颜色和图标。",
).text = (
    "点击真实地图标记，右侧会显示点位类型、状态、位置、更新时间、资料来源和真实坐标；有照片时可用"
    "左右按钮轮播。不同类型继续使用不同颜色和图标。"
)

find_paragraph(
    document,
    "资料与隐私说明\n公开底图只包含建筑编号、轮廓、新旧类型、高度、估算占地和中心坐标。户主、家庭住址、家庭成员、人口等字段不会进入网站。Spark首页通过受限接口只读取指定RAD/RADC模型文件，不载入建筑户主等个人属性；轻量化模型和恢复出的PLY保存在源码目录外，不进入Git仓库。",
).text = (
    "资料与隐私说明\n公开建筑底图只包含建筑编号、轮廓、新旧类型、高度、估算占地和中心坐标。"
    "原始WFS中的姓名、电话、住址、家庭成员和人口等字段不会进入网站。真实POI运行文件只使用素材中的"
    "点位名称、分类、简介、坐标、照片网址和资料时间。Spark首页通过受限接口只读取指定RAD/RADC模型文件。"
)

# Current technical implementation chapter.
find_paragraph(document, "10.3 建筑调研底图").text = "10.3 真实地图资料与建筑调研底图"
find_paragraph(
    document,
    "地图原始资料已经集中复制到“平台素材/地图原始数据”，当前建筑调查数据预处理直接从该目录读取。“可持续发展平台_倍伟”仅保留为旧项目参考，不再作为现有平台运行时依赖。建筑属性资料可能含有真实住户信息，不应上传至公开仓库或直接发布到网页。",
).text = (
    "“平台素材/地图服务素材”保存三农数据GeoServer下载的真实手绘图、建筑WFS、坐标范围和来源地址；"
    "POI_1785232551999.xlsx、poi图片.csv和村景图片.csv保存真实点位与照片记录。原始WFS可能含调查字段，"
    "不应上传至公开仓库或直接发布到网页。"
)
find_paragraph(
    document,
    "scripts/prepare-building-survey.mjs 从原始 GeoJSON 和属性表生成 1490 栋建筑的去隐私轻量数据，输出到 public/data/hongtang-building-survey.json。",
).text = (
    "scripts/prepare-map-service-assets.mjs生成hongtang-handdrawn-map.png、hongtang-map-layers.json和"
    "hongtang-buildings-safe.geojson；建筑文件只保留1490栋建筑的公开字段。"
)
find_paragraph(
    document,
    "BuildingSurveyExplorer 使用浏览器 SVG 渲染轮廓，提供新旧类型筛选、编号定位、缩放、点击查看档案和移动端适配。",
).text = (
    "scripts/prepare-real-map-data.py从Excel和CSV生成hongtang-real-map-features.json，当前包含56个POI、"
    "93组村景记录、392张POI照片和190张村景照片；VillageMap按经纬度定位，BuildingSurveyExplorer使用"
    "浏览器SVG渲染脱敏建筑轮廓。"
)
find_paragraph(
    document,
    "当前只完成底图展示，尚未把演示点位转换为经核实的真实坐标。",
).text = (
    "真实POI与村景点位已经按素材经纬度接入手绘图和无人机影像；问题、项目、微行动、互助资源和3D图钉"
    "仍使用演示或隐私安全位置。"
)
find_paragraph(
    document,
    "无人机影像正式公开、上传 GitHub 或部署到互联网前，应确认成果授权、敏感区域处理范围和公开分辨率。",
).text = (
    "无人机影像、真实手绘图、POI名称和现场照片正式公开、上传GitHub或部署到互联网前，应确认成果授权、"
    "个人信息和敏感区域处理范围，以及照片网址的长期可用性。"
)
find_paragraph(
    document, "无人机影像与问题、项目、行动等业务点位的真实位置校准。"
).text = (
    "继续核实真实POI的公开授权、名称和分类，补充光伏、安全与村庄记忆真实资料，并把问题、项目、行动和"
    "3D图钉逐步校准到核实位置。"
)

# Page overview table.
page_table = document.tables[1]
page_updates = {
    "小花园": "真实手绘图首屏显示35个真实花园点位和现场照片；下方为四季变化与经验交流功能骨架",
    "茶厂": "真实手绘图首屏显示9个真实茶厂点位和现场照片；下方为茶园、收茶、加工功能骨架",
    "村里用水": "真实手绘图首屏显示2个真实用水设施点位；下方为水源、问题和维修反馈功能骨架",
    "村里一张图": "真实手绘图/无人机影像/示意图切换；56个POI、93组村景、582张照片、互动点位与1490栋脱敏建筑",
}
for row in page_table.rows:
    page_name = row.cells[0].text.strip()
    if page_name in page_updates:
        row.cells[2].text = page_updates[page_name]

# FAQ table.
faq_table = document.tables[3]
for row in faq_table.rows:
    if row.cells[0].text.strip() == "建筑调研底图无法加载":
        row.cells[1].text = "真实地图运行文件尚未生成，或开发缓存仍是旧版本"
        row.cells[2].text = (
            "运行 npm run prepare:map-service-assets；关闭网站后重新双击“启动网站.bat”"
        )
add_row_like(
    faq_table,
    faq_table.rows[-1],
    [
        "真实点位照片无法显示",
        "现场照片继续使用素材CSV中的外部网址，当前网络不可用或原地址失效",
        "先检查网络；坐标、文字和本地底图仍可使用，必要时重新整理照片网址",
    ],
)

# Version record.
version_table = document.tables[4]
add_row_like(
    version_table,
    version_table.rows[-1],
    [
        "V1.29 Demo",
        "2026-07-29",
        "“村里一张图”接入真实手绘图、56个POI、93组村景和582张照片；小花园、茶厂、用水专题页优先显示真实资料；建筑底图切换为1490栋脱敏真实GeoJSON。",
    ],
)

for table in document.tables:
    keep_table_rows_intact(table)

document.save(OUTPUT)

# Structural QA without LibreOffice rendering, per project instruction.
with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None
check = Document(OUTPUT)
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
table_rows = [len(table.rows) for table in check.tables]
first_line_count = len(
    check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]')
)
assert "V1.29 Demo" in check.paragraphs[5].text
assert "56个POI" in all_text
assert "93组村景" in all_text
assert "hongtang-buildings-safe.geojson" in all_text
assert "scripts/prepare-real-map-data.py" in all_text
assert table_rows == [5, 22, 9, 14, 31]
assert len(check.inline_shapes) == 13
assert first_line_count >= 18
assert check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
assert check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"
assert check.tables[4].rows[-1].cells[0].text == "V1.29 Demo"
assert check.tables[3].rows[-1].cells[0].text == "真实点位照片无法显示"
assert all(
    row._tr.xpath("./w:trPr/w:cantSplit")
    for table in check.tables
    for row in table.rows
)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual: {MANUAL}")
print(f"Backup: {BACKUP}")
print(
    {
        "version": check.paragraphs[5].text,
        "paragraphs": len(check.paragraphs),
        "inline_shapes": len(check.inline_shapes),
        "table_rows": table_rows,
        "native_two_character_indents": first_line_count,
        "zip": "ok",
        "render": "skipped by project instruction",
    }
)
