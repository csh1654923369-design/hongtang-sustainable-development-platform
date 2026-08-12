from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_two_character_indent(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), "200")


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    set_body_font(paragraph.add_run(text))
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_body_font(paragraph.add_run(text))


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.74 ｜ 更新日期：2026年8月11日",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("页面不再设置固定顶栏")),
    "页面不再设置固定顶栏、TAB、搜索或功能按钮。左上角悬浮框显示平台Logo和“红塘村可持续发展平台”名称；其右侧的“地图编辑 / 专业工具”入口默认在新窗口打开GeoLibre空间数据实验室，原平台页面和当前地图状态继续保留。右上角独立显示“2D地图”和“3D实景”切换卡片。点击平台名称可返回首页；其余操作均在首页地图内部完成。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("手机端同样不显示固定顶栏")),
    "手机端同样不显示固定顶栏、底部导航、菜单或搜索。左上角的平台品牌和地图编辑入口会根据屏幕宽度缩为两个并列图标，其中铅笔尺图标用于在新窗口打开GeoLibre空间数据实验室；右上角继续显示2D/3D切换卡片。地图、地点详情和三维工具按钮会按屏幕宽度调整布局。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("平台新增独立的GeoLibre空间数据实验室")),
    "平台新增独立的GeoLibre空间数据实验室。启动网站后，点击首页左上角平台名称右侧的“地图编辑 / 专业工具”入口即可进入，GeoLibre默认在新窗口打开，原首页及其专题、图层和视图状态继续保留；也可直接在浏览器地址栏输入http://localhost:3000/geolibre-lab。实验室内还提供“返回平台”和“新窗口打开”。实验室不是面向村民的首页，而是供规划、调研和数据维护人员检查图层、属性和空间关系的试验工具。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.74"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.74",
    "2026年8月11日",
    "新增首页到GeoLibre空间数据实验室的可见入口；桌面端显示“地图编辑 / 专业工具”，手机端显示铅笔尺图标。入口默认在新窗口打开，原平台页面和地图状态继续保留。",
)
for cell, value in zip(new_row.cells, values):
    set_cell_text(cell, value)

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 按项目约定不进行LibreOffice渲染，只执行内容、样式、结构和压缩包完整性检查。
reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.74") for p in reloaded.paragraphs)
assert "地图编辑 / 专业工具”入口默认在新窗口打开" in paragraph_text
assert "铅笔尺图标用于在新窗口打开GeoLibre空间数据实验室" in paragraph_text
assert "原首页及其专题、图层和视图状态继续保留" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.74"
assert reloaded.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
assert reloaded.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(MANUAL)
print("version=V1.74")
print("zip=ok")
