from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_run_font(run, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_two_character_indent(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), "200")


def replace_paragraph_text(paragraph, text: str, *, heading: bool = False) -> None:
    paragraph.clear()
    set_run_font(paragraph.add_run(text), "黑体" if heading else "宋体")
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_run_font(paragraph.add_run(text))


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.76 ｜ 更新日期：2026年8月11日",
)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("6.4 GeoLibre")),
    "6.4 红塘空间数据编辑器（GeoLibre）",
    heading=True,
)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("平台新增独立的GeoLibre空间数据实验室")),
    "平台设有独立的红塘空间数据编辑器。启动网站后，点击首页右上角地图功能卡片最下方的“地图编辑 / GeoLibre专业工具”入口即可进入。编辑器默认在新窗口打开，原首页及其专题、图层和视图状态继续保留；也可直接在浏览器地址栏输入http://localhost:3000/geolibre-lab。编辑页只保留“返回平台”，不再重复提供“新窗口打开”按钮。它不是面向村民的首页，而是供规划、调研和数据维护人员编辑矢量要素、检查属性和调整图层样式的专业工具。",
)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("实验室左侧图层面板当前包含8层")),
    "编辑器左侧图层面板当前包含8层：供水分区3个、供水线路8条、水系统节点7个、小花园35处、茶产业9处、村里用水设施2处、公共服务设施9处和村景记录149处。每层都可单独显示或隐藏、调整透明度、改变顺序并缩放至范围；可绘制点、线和面，选择并移动节点、修改几何形状、删除要素、撤销或重做操作；也可检查属性、调整颜色与描边，并导入常用矢量数据或导出GeoJSON。界面顶部只保留“项目、编辑、视图、添加数据”，通用处理、三维试验、插件管理、系统设置和帮助等与当前维护任务无关的入口已经隐藏。",
)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("技术实现包括src/app/geolibre-lab")),
    "编辑器网页和GeoLibre v2.5.0完整源代码均保存在平台源码文件夹中：vendor/geolibre保存可维护的上游源码，public/geolibre保存由本地源码生成的网页文件，src/app/geolibre-lab/page.tsx负责平台入口，supabase/functions/geolibre-bridge/index.ts负责数据桥接。精简界面不会调用外部GeoLibre程序代码；启动时按需加载绘图工具，并关闭当前不需要的离线应用缓存和重型栅格数据库引擎。运行npm run test:geolibre可检查云端桥接、8个图层、本地源码、精简工具栏、矢量编辑控件和实际页面渲染。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.76"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.76",
    "2026年8月11日",
    "将GeoLibre精简为红塘专用矢量编辑器：保留图层管理、点线面绘制与几何编辑、属性与样式、导入导出和撤销重做；隐藏通用分析、三维试验、插件与设置入口，并优化主平台和编辑器的按需加载。",
)
for cell, value in zip(new_row.cells, values):
    set_cell_text(cell, value)

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 按项目约定不进行LibreOffice渲染，只执行内容、样式、结构和压缩包完整性检查。
reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.76") for p in reloaded.paragraphs)
assert "6.4 红塘空间数据编辑器（GeoLibre）" in paragraph_text
assert "可绘制点、线和面" in paragraph_text
assert "通用处理、三维试验、插件管理、系统设置和帮助" in paragraph_text
assert "vendor/geolibre保存可维护的上游源码" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.76"
assert reloaded.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
assert reloaded.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(MANUAL)
print("version=V1.76")
print("zip=ok")
