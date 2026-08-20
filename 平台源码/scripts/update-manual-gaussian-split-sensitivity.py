from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v115.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.14.backup.docx"
HOME_SCREENSHOT = QA_ROOT / "home-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


def set_image(paragraph):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(HOME_SCREENSHOT), width=Inches(6.3))
    shape._inline.docPr.set("name", "图 1 3D高斯实景首页与独立灵敏度设置")
    shape._inline.docPr.set(
        "descr",
        "红塘村3D高斯实景首页打开操作设置，右键平移与滚轮缩放分别使用独立滑块",
    )


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        copy_cell_format(template_row.cells[index], cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if template_row.cells[index].paragraphs and cell.paragraphs:
            cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
            cell.paragraphs[0].paragraph_format.alignment = (
                template_row.cells[index].paragraphs[0].paragraph_format.alignment
            )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_first_line_chars(paragraph, chars=200):
    properties = paragraph._p.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    for attribute in ("firstLine", "hanging", "hangingChars"):
        key = qn(f"w:{attribute}")
        if key in indent.attrib:
            del indent.attrib[key]
    indent.set(qn("w:firstLineChars"), str(chars))


def normalize_document(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
        if (
            index >= 8
            and paragraph.style
            and paragraph.style.name == "Normal"
            and paragraph.text.strip()
            and "\n" not in paragraph.text
        ):
            set_first_line_chars(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not HOME_SCREENSHOT.exists():
    raise FileNotFoundError(HOME_SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.14 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.14 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.15 Demo   |   更新日期：2026年7月23日   |   适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页使用红塘村在线3D高斯实景；"
    "鼠标左键旋转、右键沿村庄水平面平移，滚轮按固定比例缩放。"
    "右键平移和滚轮缩放各有一个独立的五档灵敏度设置，滚轮五档为每步6%、10%、15%、22%和32%。"
    "原无人机影像首页完整保留为“村庄总览”。3D首页需要联网，业务点位仍不能作为真实设施、风险或产权判断依据。"
)

caption = find_paragraph(document, "图 1  3D高斯实景首页与操作设置")
set_image(image_paragraph_before(document, caption))
caption.text = "图 1  3D高斯实景首页与独立灵敏度设置"

find_paragraph(
    document,
    "3D首页操作：左键拖动旋转；右键拖动只沿村庄所在地表的水平面平移；"
    "滚轮每次按固定比例缩放，快速连续滚动不会额外放大单步距离。"
    "点击“操作设置”可在五档灵敏度之间调节，并可使用“回到模型”和“全屏查看”。触控设备继续使用双指缩放。",
).text = (
    "3D首页操作：左键拖动旋转；右键拖动只沿村庄所在地表的水平面平移；"
    "滚轮每次按固定比例缩放，快速连续滚动不会额外放大单步距离。点击“操作设置”后，"
    "“右键平移灵敏度”和“滚轮缩放灵敏度”可以分别调节、分别保存，调整其中一项不会改变另一项。"
    "缩放默认“适中”为每步15%，最高档为每步32%。还可使用“回到模型”和“全屏查看”；触控设备继续使用双指缩放。"
)

for row in document.tables[1].rows:
    if row.cells[0].text.strip() == "首页":
        row.cells[2].text = (
            "全屏红塘村3D高斯实景；左键旋转、右键水平平移、滚轮定比缩放，"
            "并可分别调节平移与缩放灵敏度"
        )
        break

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载红塘村高斯模型；服务端保存长期令牌；自定义右键水平平移、"
            "滚轮定比缩放，以及两项独立五档灵敏度"
        )
        break

find_paragraph(
    document,
    "平台从Cesium默认缩放输入中移除鼠标右键和滚轮：右键位移先投影到模型所在地表切平面；"
    "滚轮忽略事件幅度，每120毫秒最多执行一次定比缩放。五档每步比例为3.5%、5%、7%、10%和14%，默认7%。",
).text = (
    "平台从Cesium默认缩放输入中移除鼠标右键和滚轮：右键位移先投影到模型所在地表切平面，"
    "继续使用0.5、0.72、1、1.28和1.6倍五档倍率；滚轮忽略事件幅度，每120毫秒最多执行一次定比缩放，"
    "五档每步比例为6%、10%、15%、22%和32%，默认15%。两项设置分别写入浏览器本地存储。"
)

for row in document.tables[3].rows:
    if row.cells[0].text.strip() == "3D平移或缩放太快、太慢":
        row.cells[1].text = "当前设备与某一项默认“适中”灵敏度不匹配"
        row.cells[2].text = (
            "点击3D画面右上角“操作设置”，分别调整“右键平移灵敏度”或“滚轮缩放灵敏度”；"
            "两项选择会分别保存在当前浏览器"
        )
        break

version_table = document.tables[4]
version_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    version_row.cells,
    [
        "V1.15 Demo",
        "2026-07-23",
        "拆分右键平移与滚轮缩放灵敏度；保留原平移五档，并将滚轮五档提高为每步6%、10%、15%、22%和32%。",
    ],
):
    cell.text = value

normalize_document(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert "V1.15 Demo" in check.paragraphs[5].text
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 17]
assert any("调整其中一项不会改变另一项" in paragraph.text for paragraph in check.paragraphs)
assert any("五档每步比例为6%、10%、15%、22%和32%" in paragraph.text for paragraph in check.paragraphs)
assert any(row.cells[0].text.strip() == "V1.15 Demo" for row in check.tables[4].rows)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.15: {MANUAL}")
