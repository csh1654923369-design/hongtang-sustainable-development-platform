from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v112.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.11.backup.docx"
HOME_SCREENSHOT = QA_ROOT / "home-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


def set_image(paragraph, image_path, width, name, description):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("name", name)
    shape._inline.docPr.set("descr", description)


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        copy_cell_format(template_row.cells[index], cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if template_row.cells[index].paragraphs and cell.paragraphs:
            cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
            cell.paragraphs[0].paragraph_format.alignment = (
                template_row.cells[index].paragraphs[0].paragraph_format.alignment
            )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_first_line_chars(paragraph, chars=200):
    properties = paragraph._p.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    for attribute in ("firstLine", "hanging", "hangingChars"):
        key = qn(f"w:{attribute}")
        if key in indent.attrib:
            del indent.attrib[key]
    indent.set(qn("w:firstLineChars"), str(chars))


def normalize_document(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")

    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        east_asia = "SimHei" if is_heading or index == 0 else "SimSun"
        for run in paragraph.runs:
            set_run_font(run, east_asia)
        if (
            index >= 8
            and paragraph.style
            and paragraph.style.name == "Normal"
            and paragraph.text.strip()
            and "\n" not in paragraph.text
        ):
            set_first_line_chars(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not HOME_SCREENSHOT.exists():
    raise FileNotFoundError(HOME_SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.11 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.11 manual, got: {document.paragraphs[5].text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.12 Demo   |   更新日期：2026年7月23日   |   适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页使用红塘村在线3D高斯实景，"
    "可以旋转、缩放和平移；原无人机影像首页完整保留为“村庄总览”。"
    "3D首页需要联网，业务点位仍未按影像校准真实位置，不能作为真实设施、风险或产权判断依据。"
)

home_caption = find_paragraph_prefix(document, "图 1")
set_image(
    image_paragraph_before(document, home_caption),
    HOME_SCREENSHOT,
    6.3,
    "图 1 3D高斯实景首页",
    "红塘村3D高斯实景铺满首页画面，顶部保留平台导航",
)
home_caption.text = "图 1  3D高斯实景首页"

replacements = {
    "以后不必每次输入命令。在项目根目录双击“启动网站.bat”，脚本会在后台启动网站并自动打开浏览器。":
        "以后不必每次输入命令。在项目根目录双击“启动网站.bat”，脚本会在后台启动网站并自动打开浏览器。首页3D实景需要保持联网；加载失败时仍可进入“村庄总览”。",
    "桌面端顶栏：平台名称、六类事项 TAB 和常用操作位于同一行；点击“村庄总览”返回首页，点击具体事项则先进入该事项的地图首屏。":
        "桌面端顶栏：平台名称、八个 TAB 和常用操作位于同一行；“首页”打开3D高斯实景，“村庄总览”打开原无人机影像首页，具体事项仍先进入各自的地图首屏。",
    "村里一张图：从首页影像、右上角地图按钮或手机底栏进入，按具体事项、行动办理、互助资源和调研资料查找点位。":
        "村里一张图：从“村庄总览”的无人机影像、右上角地图按钮或手机底栏进入，按具体事项、行动办理、互助资源和调研资料查找点位。",
    "更多工具：活动、项目、认识红塘、可持续目标、研究资料和数字沙盘从首页“其他工具”、移动菜单、搜索或页脚进入。":
        "更多工具：活动、项目、认识红塘、可持续目标、研究资料和数字沙盘从“村庄总览”的“其他工具”、移动菜单、搜索或页脚进入。",
    "在手机或窄屏浏览器中，横排事项 TAB 会收进右上角菜单；底部保留“首页、选事项、一张图、我的”四个快捷入口。网站已针对 390px 宽度进行无横向溢出测试。":
        "在手机或窄屏浏览器中，横排 TAB 会收进右上角菜单；底部“首页”打开3D实景，“选事项”进入村庄总览中的事项入口，另保留“一张图、我的”。网站已针对390px宽度完成无横向溢出测试。",
    "资料与隐私说明\n公开底图只包含建筑编号、轮廓、新旧类型、高度、估算占地和中心坐标。户主、家庭住址、家庭成员、人口等字段不会进入网站；大型 Cesium 和三维高斯模型也未打包到当前 Demo。":
        "资料与隐私说明\n公开底图只包含建筑编号、轮廓、新旧类型、高度、估算占地和中心坐标。户主、家庭住址、家庭成员、人口等字段不会进入网站。3D首页只在线加载高斯实景，不载入建筑户主等个人属性。",
    "从首页“更多资料”、右上角移动菜单、搜索或相关项目进入“看看改造后的样子（数字沙盘）”，然后勾选或取消建筑、项目、绿化图层。":
        "从“村庄总览”的“其他工具”、右上角移动菜单、搜索或相关项目进入“看看改造后的样子（数字沙盘）”，然后勾选或取消建筑、项目、绿化图层。",
    "当前边界\n数字沙盘目前是 CSS/HTML 交互占位组件，不是真实三维模型。未来可替换为 Cesium 或三维高斯场景，同时保留其他页面业务逻辑。":
        "当前边界\n首页已经接入红塘村真实3D高斯实景，但“数字沙盘”仍是用于演示改造方案切换的CSS/HTML概念组件，两者用途不同。未来可把经过确认的方案模型接入数字沙盘。",
    "Cesium 或三维高斯数字孪生场景。":
        "数字沙盘中的改造方案A/B真实三维模型，以及与项目数据的联动。",
}
for old, new in replacements.items():
    find_paragraph(document, old).text = new

role_rows = [
    ["角色", "可以做什么", "主要入口"],
    ["游客", "浏览3D首页、村庄总览、六个事项、村里一张图、项目和公开进展", "首页、村庄总览、顶部事项栏、村里一张图"],
    ["村民", "查看3D实景和具体事项、上报问题、发起或加入微行动、共享资源、报名、建议与评价", "首页、顶部事项栏、村里一张图、参与、我的"],
    ["学生/规划协作者", "查看3D实景、参与事项功能共创、发起或加入微行动、共享资源、提交调研成果与建议", "首页、顶部事项栏、村里一张图、项目、研究协作、我的"],
    ["管理员/村委", "轻量核对微行动、处理问题、审核资料、维护项目和查看日志", "管理后台"],
]
for row, values in zip(document.tables[0].rows, role_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

page_rows = [
    ["页面", "地址", "核心用途"],
    ["首页", "/", "全屏红塘村3D高斯实景，可旋转、缩放、平移、回到模型和全屏查看"],
    ["村庄总览", "/village-overview", "原首页完整保留：无人机影像主视觉、六个事项入口和三个常用操作"],
    ["小花园", "/garden", "地图首屏只显示小花园点位；下方为四季变化与经验交流功能骨架"],
    ["茶厂", "/tea-factory", "地图首屏只显示茶场、茶厂点位；下方为茶园、收茶、加工功能骨架"],
    ["村里用水", "/water", "地图首屏只显示用水设施点位；下方为水源、问题和维修反馈功能骨架"],
    ["光伏设施", "/solar", "地图首屏只显示光伏设施点位；下方为运行巡查和信息公开功能骨架"],
    ["安全隐患", "/safety", "地图首屏只显示安全隐患点位；下方为发现、处理和复查功能骨架"],
    ["村庄记忆", "/village-history", "地图首屏只显示村庄记忆点位；下方为照片、讲述、古道与老屋档案骨架"],
    ["村里的事（旧汇总）", "/village-life", "兼容旧链接，不再作为顶部主入口"],
    ["认识红塘", "/village", "村庄资料框架与故事"],
    ["可持续目标", "/goals", "目标、指标、项目联动"],
    ["村里一张图", "/map", "四组点位筛选、示意图/无人机影像切换、事项详情跳转、互助资源和建筑调研底图"],
    ["问题上报", "/report", "五步记录村庄问题"],
    ["发起微行动", "/actions/new", "五步形成可核对、可招募的小行动"],
    ["项目与行动", "/projects", "社区微行动、正式项目、开放任务与资源需求"],
    ["发展进展", "/progress", "社区行动能力、趋势、来源与完整度"],
    ["公众参与", "/participate", "微行动入口、建议、活动、问卷和共创"],
    ["个人中心", "/profile", "我的行动、个人记录与通知"],
    ["研究协作", "/research", "调研成果提交"],
    ["管理后台", "/admin", "问题处理、微行动核对、审核与维护"],
    ["数字沙盘", "/digital-twin", "改造方案三维场景概念演示，与3D实景首页用途不同"],
]
page_table = document.tables[1]
while len(page_table.rows) < len(page_rows):
    add_row_like(page_table, page_table.rows[-1])
for row, values in zip(page_table.rows, page_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

technology_table = document.tables[2]
technology_row = add_row_like(technology_table, technology_table.rows[-1])
for cell, value in zip(
    technology_row.cells,
    [
        "3D实景",
        "CesiumJS 1.143 + Cesium ion",
        "在线加载红塘村高斯模型；长期访问令牌只保存在服务端环境变量中",
    ],
):
    cell.text = value

troubleshooting_table = document.tables[3]
troubleshooting_row = add_row_like(troubleshooting_table, troubleshooting_table.rows[-1])
for cell, value in zip(
    troubleshooting_row.cells,
    [
        "3D首页一直加载或显示失败",
        "未联网、Cesium访问令牌未配置，或在线模型暂时不可访问",
        "检查网络并重新加载；仍失败时点击“进入村庄总览”继续使用其他功能",
    ],
):
    cell.text = value

orthophoto_heading = find_paragraph(document, "10.4 无人机正射影像")
orthophoto_heading.text = "10.5 无人机正射影像"
orthophoto_heading.insert_paragraph_before(
    "外层“3D高斯展示”文件夹保存原在线查看页和建筑轮廓参考；实际高斯模型是Cesium ion在线资产，不复制进源码仓库。",
    style="List Bullet",
)
orthophoto_heading.insert_paragraph_before(
    "src/app/api/cesium-ion/asset/route.ts在服务端读取CESIUM_ION_TOKEN并获取模型地址；public/gaussian-viewer/index.html负责全屏交互显示。",
    style="List Bullet",
)
orthophoto_heading.insert_paragraph_before(
    "首页仅加载高斯实景，不载入building_footprints.geojson，也不展示户主、住址或家庭成员等个人属性。",
    style="List Bullet",
)
orthophoto_heading.insert_paragraph_before("10.4 3D高斯实景首页", style="Heading 2")
find_paragraph(document, "10.5 Supabase、Git 与 GitHub").text = "10.6 Supabase、Git 与 GitHub"

privacy_anchor = find_paragraph(
    document,
    "管理员角色仅为 Demo 切换，不等同于真实登录或安全认证。",
)
privacy_anchor.insert_paragraph_before(
    "3D高斯实景公开或部署到互联网前，应确认航拍成果和在线模型的公开授权，并限制Cesium访问令牌的资产范围和网站来源。",
    style="List Bullet",
)

version_table = document.tables[4]
version_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    version_row.cells,
    [
        "V1.12 Demo",
        "2026-07-23",
        "新增全屏红塘村3D高斯实景首页；原首页迁移为“村庄总览”；顶部增加首页入口，并补充联网、令牌与隐私边界说明。",
    ],
):
    cell.text = value

normalize_document(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert "V1.12 Demo" in check.paragraphs[5].text
assert len(check.tables[1].rows) == 22
assert len(check.tables[2].rows) == 9
assert len(check.tables[3].rows) == 12
assert len(check.tables[4].rows) == 14
assert any(paragraph.text.strip() == "10.4 3D高斯实景首页" for paragraph in check.paragraphs)
assert any(
    row.cells[0].text.strip() == "村庄总览" and row.cells[1].text.strip() == "/village-overview"
    for row in check.tables[1].rows
)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.12: {MANUAL}")
