from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v128.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.27.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def add_version_row(table):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        properties = cell._tc.get_or_add_tcPr()
        for child in list(properties):
            properties.remove(child)
        for child in template_row.cells[index]._tc.tcPr:
            properties.append(deepcopy(child))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        cell.paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.27 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.27 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.28 Demo   |   更新日期：2026年7月29日   |   适用地址：http://localhost:3000"
)

find_paragraph_prefix(document, "以后不必每次输入命令").text = (
    "以后不必每次输入命令。项目根目录保留“启动网站.bat”和“关闭网站.bat”；"
    "网页源代码集中放在“平台源码”文件夹，3D模型、无人机原始影像和地图原始数据集中放在"
    "“平台素材”文件夹。3D首页默认从“平台素材/3D高斯展示/轻量化模型”读取模型。"
)

find_paragraph_prefix(document, "影像与点位说明").text = (
    "影像与点位说明\n无人机原始影像保存在“平台素材/Production_1-tif”文件夹；"
    "网站展示使用“平台源码/public/hongtang-orthophoto”中的网页切片。"
    "地图点位和3D图钉仍为演示数据，待后续实地核对后再替换。"
)

find_paragraph_prefix(document, "原始资料保存在").text = (
    "地图原始资料已经集中复制到“平台素材/地图原始数据”，当前建筑调查数据预处理直接从该目录读取。"
    "“可持续发展平台_倍伟”仅保留为旧项目参考，不再作为现有平台运行时依赖。"
    "建筑属性资料可能含有真实住户信息，不应上传至公开仓库或直接发布到网页。"
)

find_paragraph_prefix(document, "外层“3D高斯展示”").text = (
    "“平台素材/3D高斯展示”文件夹保存Spark官方源码、转换工具、原始模型恢复资料和"
    "“轻量化模型”。正式首页通过“平台源码”中的受限Range接口读取轻量化模型，"
    "不需要从旧项目或其他外部目录调用文件。"
)

find_paragraph_prefix(document, "外层 Production_1-tif").text = (
    "“平台素材/Production_1-tif”文件夹保存无人机原始正射影像、DSM及其辅助文件。"
    "这些文件体积较大，网站运行时使用已经生成到“平台源码/public”中的网页版本；"
    "只有重新生成网页影像时才会读取原始素材。"
)

version_row = add_version_row(document.tables[4])
for cell, value in zip(
    version_row.cells,
    [
        "V1.28 Demo",
        "2026-07-29",
        "新增“平台素材”目录，集中3D模型、无人机原始影像和地图原始数据；更新运行、预处理和使用说明中的素材路径。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 30]
assert "V1.28 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "平台素材/3D高斯展示/轻量化模型" in all_text
assert "平台素材/Production_1-tif" in all_text
assert "平台素材/地图原始数据" in all_text
assert "不再作为现有平台运行时依赖" in all_text
assert any(row.cells[0].text.strip() == "V1.28 Demo" for row in check.tables[4].rows)
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count
assert check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
assert check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.28: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
