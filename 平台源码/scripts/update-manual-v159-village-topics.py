from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        set_body_font(run)


def insert_before(document, anchor, text: str, style: str):
    paragraph = anchor.insert_paragraph_before(text, style=style)
    return paragraph


document = Document(MANUAL)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.59 ｜ 更新日期：2026年8月10日",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("小花园、茶厂、村里用水、公共服务")),
    "首页按照红塘村的具体事情组织为五个专题：小花园、茶产业、村里用水、塌方与安全、茶马古道与村庄历史。"
    "公共服务、生态资源和村景记录保留为地图基础资料，不再与五个主题并列。用户通过首页左上角专题卡片进入，2D与3D共用同一套专题状态和地图内容。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("桌面端和手机端都在地图左上角显示同一张白色圆角“专题”卡片")),
    "桌面端和手机端都在地图左上角显示同一张白色圆角“专题”卡片。展开后首先显示小花园、茶产业、村里用水、塌方与安全、"
    "茶马古道与村庄历史五个专题，下面另列公共服务、生态资源和村景记录等地图基础资料。专题和基础资料均可多选；“全选”恢复全部内容。"
    "每个专题右侧的“进入”用于切换到专题阅读状态，数字表示已接入的空间资料数量，没有已核实资料的专题显示“待调查”。切换2D与3D不会重置当前状态。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("地点：包括小花园、茶厂、村里用水设施")),
    "专题地点：当前已接入35处小花园、9处茶厂和已有水设施资料；茶园、塌方安全点、茶马古道线路及历史地点仍需调查。"
    "公共服务、生态资源和村景记录作为基础资料继续显示。点击已有地点后，2D和3D都会聚焦并在点位旁显示同一气泡详情。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("小花园、茶厂、公共服务、生态资源、村景记录和村里用水不再")),
    "进入首页后点击左上角“专题”，可以多选五个红塘专题，也可以单独勾选地图基础资料。点击某一专题右侧“进入”，地图只保留该专题相关内容，"
    "并显示该专题需要回答的问题、已有资料数量和后续调查重点；关闭专题后恢复进入前的勾选状态。",
)

water_heading = next(p for p in document.paragraphs if p.text == "4.3 村里用水专题")
if not any(p.text == "4.3 小花园与茶产业专题" for p in document.paragraphs):
    insert_before(document, water_heading, "4.3 小花园与茶产业专题", "Heading 2")
    insert_before(
        document,
        water_heading,
        "进入“小花园”后，地图显示现有35处小花园，并提示继续记录空间分布、种植与物种、季度和年度变化。进入“茶产业”后，地图显示现有9处茶厂；"
        "茶园边界、土壤与有机性监测、采青数量、收购流向和农户—茶厂关系尚未形成完整数据，界面明确提示后续补充，不以示例点代替。",
        "Normal",
    )
    insert_before(document, water_heading, "4.4 塌方与安全、茶马古道与村庄历史专题", "Heading 2")
    insert_before(
        document,
        water_heading,
        "这两个专题已经建立入口和数据类型，但目前没有已核实空间对象。进入后地图显示“资料待调查”，并分别提示记录隐患位置、影响对象、处置复查，"
        "以及古道线路、历史地点、口述与影像。后续把核实资料写入共享数据后，2D与3D会自动按相同图标和交互方式显示。",
        "Normal",
    )

replace_paragraph_text(water_heading, "4.5 村里用水专题")
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text == "4.4 已移除的旧板块"),
    "4.6 已移除的旧板块",
)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("“村里一张图”“我要记录”“村庄总览”")),
    "“村里一张图”“我要记录”“村庄总览”及原来分散的小花园、茶厂、村里用水、安全隐患和村庄记忆页面均已删除，相关内容改在首页专题中查看。"
    "旧光伏设施页面及其演示内容同样删除，光伏不纳入当前五个专题。",
)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("当前交付版本只把首页作为面向用户的正式页面")),
    "当前交付版本只把首页作为面向用户的正式页面。首页已经包含2D地图、3D实景、五个红塘专题、地图基础资料、照片气泡详情，"
    "以及饮水供给和排水去向的点线面关系，不再另设地图页、记录页或事项页。",
)

if not any("npm run test:village-topics" in p.text for p in document.paragraphs):
    test_anchor = next(p for p in document.paragraphs if "npm run test:water-topic" in p.text)
    new_test = document.add_paragraph(
        "npm run test:village-topics：验证五个专题入口、光伏专题不存在、小花园和茶产业已有数量、无资料专题提示，以及2D/3D切换时专题状态保持。",
        style="List Bullet",
    )
    test_anchor._p.addprevious(new_test._p)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_version_row = next((row for row in version_table.rows if row.cells[0].text == "V1.59"), None)
if new_version_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_version_row = version_table.rows[1]
replace_cell_text(new_version_row.cells[0], "V1.59")
replace_cell_text(new_version_row.cells[1], "2026年8月10日")
replace_cell_text(
    new_version_row.cells[2],
    "按老师提出的红塘具体事项重组为小花园、茶产业、村里用水、塌方与安全、茶马古道与村庄历史五个专题；删除光伏专题。"
    "已有资料直接接入，缺少核实资料的专题显示待调查，不建立虚构点位。",
)

document.save(MANUAL)

reloaded = Document(MANUAL)
text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.59") for p in reloaded.paragraphs)
assert "4.3 小花园与茶产业专题" in text
assert "4.4 塌方与安全、茶马古道与村庄历史专题" in text
assert "4.5 村里用水专题" in text
assert "光伏不纳入当前五个专题" in text
assert "npm run test:village-topics" in text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.59"
print(MANUAL)
