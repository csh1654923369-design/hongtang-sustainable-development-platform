from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.63 ｜ 更新日期：2026年8月10日",
)

overview = next(p for p in document.paragraphs if p.text.startswith("首页默认进入红塘村2D地图"))
overview_text = overview.text.replace(
    "2D以高德在线地图作为云端底层，可叠加红塘无人机影像或手绘图；",
    "2D以高德在线地图作为云端底层，可切换高德卫星遥感图，也可叠加红塘无人机影像或手绘图；",
).replace(
    "左上角平台悬浮框右侧可在“2D地图”和“3D实景”之间切换；",
    "“2D地图”和“3D实景”切换按钮位于右上角；选中2D时，其下方显示“高德地图”“卫星遥感”“无人机影像”“手绘图”四种底图按钮；",
)
replace_paragraph_text(overview, overview_text)

two_d = next(p for p in document.paragraphs if p.text.startswith("2D模式与3D实景读取同一套地点和专题数据"))
two_d_text = two_d.text.replace(
    "2D地图使用高德在线道路、地名和水系作为云端底层，右上角可选择“高德底图”“无人机影像”或“手绘图”；后两项会以半透明地理配准图层叠加在高德地图上，",
    "2D地图使用高德在线道路、地名、水系和卫星遥感影像作为云端底层。右上角第一层切换2D/3D；选中2D后，第二层可选择“高德地图”“卫星遥感”“无人机影像”或“手绘图”。“卫星遥感”使用高德官方卫星与路网图层；“无人机影像”和“手绘图”会以半透明地理配准图层叠加在高德地图上，",
)
replace_paragraph_text(two_d, two_d_text)

test_note = next(p for p in document.paragraphs if p.text.startswith("npm run test:2d-markers"))
replace_paragraph_text(
    test_note,
    test_note.text.replace(
        "在高德底图、无人机影像和手绘图三种模式下",
        "在高德地图、卫星遥感、无人机影像和手绘图四种模式下",
    ),
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.63"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.63",
    "2026年8月10日",
    "首页2D/3D切换移至右上角；选中2D时在下方显示四种底图按钮，并新增高德卫星遥感与路网图层。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)
    for run in cell.paragraphs[0].runs:
        set_body_font(run)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.63") for p in reloaded.paragraphs)
assert "切换按钮位于右上角" in paragraph_text
assert "高德地图”“卫星遥感”“无人机影像”“手绘图”四种底图按钮" in paragraph_text
assert "卫星遥感”使用高德官方卫星与路网图层" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.63"
print(MANUAL)
