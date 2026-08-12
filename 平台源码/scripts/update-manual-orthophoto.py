from copy import deepcopy
from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.text.paragraph import Paragraph


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v17.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.6.backup.docx"
MAP_SCREENSHOT = QA_ROOT / "map-aerial-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def insert_paragraph_before(target, text="", style=None):
    element = OxmlElement("w:p")
    target._p.addprevious(element)
    paragraph = Paragraph(element, target._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption_text):
    caption = find_paragraph(document, caption_text)
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError(f"Image paragraph not found before caption: {caption_text}")


def set_image(paragraph, image_path, width, name, description):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("name", name)
    shape._inline.docPr.set("descr", description)


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def set_cell_text_like(source_cell, target_cell, text):
    copy_cell_format(source_cell, target_cell)
    target_cell.text = text
    target_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if source_cell.paragraphs and target_cell.paragraphs:
        source = source_cell.paragraphs[0]
        target = target_cell.paragraphs[0]
        target.style = source.style
        target.paragraph_format.alignment = source.paragraph_format.alignment


def add_row_like(table, template_row, values):
    row = table.add_row()
    for index, value in enumerate(values):
        set_cell_text_like(template_row.cells[index], row.cells[index], value)
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        east_asia = "SimHei" if is_heading or index == 0 else "SimSun"
        for run in paragraph.runs:
            set_run_font(run, east_asia)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


def keep_table_rows_intact(table):
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        if not properties.xpath("./w:cantSplit"):
            properties.append(OxmlElement("w:cantSplit"))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    if not header_properties.xpath("./w:tblHeader"):
        header_properties.append(OxmlElement("w:tblHeader"))


if not MAP_SCREENSHOT.exists():
    raise FileNotFoundError(MAP_SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
version_text = document.paragraphs[5].text
if "V1.7 Demo" in version_text:
    set_image(
        image_paragraph_before(document, "图 3  村里一张图与点位详情"),
        MAP_SCREENSHOT,
        6.3,
        "图 3 村里一张图",
        "村里一张图切换到红塘村无人机正射影像后的页面",
    )
    normalize_fonts(document)
    document.save(OUTPUT)
    check = Document(OUTPUT)
    assert len(check.inline_shapes) == 13
    os.replace(OUTPUT, MANUAL)
    print(f"Refreshed V1.7 map screenshot: {MANUAL}")
    raise SystemExit(0)
if "V1.6 Demo" not in version_text:
    raise ValueError(f"Expected V1.6 manual, got: {version_text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = "版本：V1.7 Demo   |   更新日期：2026年7月21日   |   适用地址：http://localhost:3000"

map_intro = find_paragraph(
    document,
    "“村里一张图”分为三部分：上方地图汇集问题、正式项目、社区微行动、已完成行动和资源供需点位；中部互助板展示可提供资源与需求；下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。除建筑轮廓外，坐标与业务点位均为演示数据。",
)
map_intro.text = (
    "“村里一张图”分为三部分：上方地图可在简化示意图和红塘村无人机正射影像之间切换，并叠加问题、项目、社区微行动和资源供需点位；"
    "中部互助板展示可提供资源与需求；下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。当前业务点位仍是演示位置，尚未与真实影像完成校准。"
)

set_image(
    image_paragraph_before(document, "图 3  村里一张图与点位详情"),
    MAP_SCREENSHOT,
    6.3,
    "图 3 村里一张图",
    "村里一张图切换到红塘村无人机正射影像后的页面",
)

filter_step = find_paragraph(document, "在左侧按“行动与问题、社区互助资源、空间与调研资料”勾选或取消图层。")
insert_paragraph_before(filter_step, "使用地图右上角的“示意图 / 无人机影像”按钮切换底图。", "List Number")
insert_paragraph_before(filter_step, "切换到无人机影像后，仍可点击原有点位；页面左下角会持续提示“点位尚未按真实位置校准”。", "List Number")

map_section = find_paragraph(document, "3.1 查看和筛选点位")
insert_paragraph_before(
    map_section,
    "影像与点位说明\n无人机影像来自 Production_1-tif 中的正射成果，网页使用 0.3 米轻量展示版。现有点位坐标和图上位置仍为演示数据，不能据此判断真实问题、产权边界或设施位置。",
    "Normal",
)

supabase_heading = find_paragraph(document, "10.4 Supabase、Git 与 GitHub")
supabase_heading.text = "10.5 Supabase、Git 与 GitHub"
insert_paragraph_before(supabase_heading, "10.4 无人机正射影像", "Heading 2")
insert_paragraph_before(supabase_heading, "外层 Production_1-tif 文件夹保存 4厘米正射影像、DSM、裁剪和重采样成果，不进入源码仓库。", "List Bullet")
insert_paragraph_before(supabase_heading, "scripts/prepare-orthophoto.mjs 从 0.3 米重采样 TIF 生成约 3.7MB 的透明 WebP，并写入坐标系、像素大小、范围和校准状态元数据。", "List Bullet")
insert_paragraph_before(supabase_heading, "网页读取 public/data/hongtang-orthophoto-0.3m.webp；配套元数据为 public/data/hongtang-orthophoto.json，坐标系是 EPSG:32647。", "List Bullet")
insert_paragraph_before(supabase_heading, "当前只完成底图展示，尚未把演示点位转换为经核实的真实坐标。", "List Bullet")

admin_principle = find_paragraph(document, "管理员角色仅为 Demo 切换，不等同于真实登录或安全认证。")
insert_paragraph_before(
    admin_principle,
    "无人机影像正式公开、上传 GitHub 或部署到互联网前，应确认成果授权、敏感区域处理范围和公开分辨率。",
    "List Bullet",
)

map_limit = find_paragraph(document, "互联网瓦片底图、GPS 定位和 PostGIS 空间查询；当前只接入了既有建筑轮廓调研资料。")
map_limit.text = "无人机影像与问题、项目、行动等业务点位的真实位置校准。"
insert_paragraph_before(map_limit, "互联网瓦片底图、GPS 定位和 PostGIS 空间查询。", "List Bullet")

lint_command = find_paragraph(document, "npm run lint：检查代码规范。")
insert_paragraph_before(lint_command, "npm run prepare:orthophoto：从外层 0.3 米 TIF 重新生成网页正射影像和元数据。", "List Bullet")

page_table = document.tables[1]
map_row = next(row for row in page_table.rows if row.cells[0].text.strip() == "村里一张图")
map_row.cells[2].text = "示意图/无人机影像切换、业务点位、互助资源和建筑调研底图"

version_table = document.tables[4]
add_row_like(
    version_table,
    version_table.rows[-1],
    [
        "V1.7 Demo",
        "2026-07-21",
        "“村里一张图”新增示意图与无人机正射影像切换；接入 0.3 米网页展示版及 EPSG:32647 元数据，并明确标注业务点位尚未完成真实位置校准。",
    ],
)

for table in document.tables:
    keep_table_rows_intact(table)
normalize_fonts(document)

captions = [
    paragraph.text.strip()
    for paragraph in document.paragraphs
    if paragraph.style and paragraph.style.name == "Caption" and paragraph.text.strip().startswith("图 ")
]
if len(captions) != len(document.inline_shapes):
    raise AssertionError(f"Caption/image mismatch: {len(captions)} captions, {len(document.inline_shapes)} images")
for index, (shape, caption) in enumerate(zip(document.inline_shapes, captions), start=1):
    shape._inline.docPr.set("name", f"Figure {index}")
    shape._inline.docPr.set("descr", caption)

document.save(OUTPUT)

check = Document(OUTPUT)
assert "V1.7 Demo" in check.paragraphs[5].text
assert len(check.inline_shapes) == 13
assert any(paragraph.text == "10.4 无人机正射影像" for paragraph in check.paragraphs)
assert any("点位尚未按真实位置校准" in paragraph.text for paragraph in check.paragraphs)
assert "无人机影像" in next(row for row in check.tables[1].rows if row.cells[0].text.strip() == "村里一张图").cells[2].text
assert check.tables[4].rows[-1].cells[0].text.strip() == "V1.7 Demo"
assert all(row._tr.xpath("./w:trPr/w:cantSplit") for table in check.tables for row in table.rows)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual: {MANUAL}")
print(f"Backup: {BACKUP}")
print(f"Paragraphs: {len(check.paragraphs)} | tables: {len(check.tables)} | images: {len(check.inline_shapes)}")
