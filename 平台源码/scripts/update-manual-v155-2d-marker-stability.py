from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)

version_paragraph = next(p for p in document.paragraphs if p.text.startswith("版本：V1."))
replace_paragraph_text(version_paragraph, "版本：V1.55 ｜ 更新日期：2026年8月9日")

operation_paragraph = next(
    p for p in document.paragraphs
    if p.text.startswith("2D模式与3D实景读取同一套地点和专题数据。")
)
replace_paragraph_text(
    operation_paragraph,
    "2D模式与3D实景读取同一套地点和专题数据。2D地图使用高德在线道路、地名和水系作为云端底层，"
    "右上角可选择“高德底图”“无人机影像”或“手绘图”；后两项会以半透明地理配准图层叠加在高德地图上，"
    "因此村域外仍能看到周边道路并继续移动。按住鼠标左键可向任意方向拖动，滚轮缩放，点击右下角“回到红塘”"
    "恢复村庄中心视图。地点图钉使用独立于底图缩放的矢量层，并按整数屏幕像素定位；悬停或选中时不再整体放大，"
    "所以连续缩放和拖动后仍能保持清晰。点击小花园、茶厂等地点后，地图会移动并放大到对应坐标，在点位旁弹出与3D"
    "相同结构的气泡详情；点击保护会防止气泡刚打开就被底图空白点击关闭。点击水源、供排水线路或片区也会聚焦对象并"
    "突出关联关系。原始地点数据保留WGS84坐标，2D显示时统一转换为高德使用的GCJ-02，避免底图与村庄资料错位；3D仍读取"
    "原始WGS84坐标。高德服务不可用时，页面自动回退到本地影像地图，不影响筛选和详情查看。",
)

if not any(p.text.startswith("npm run test:2d-markers") for p in document.paragraphs):
    water_test_paragraph = next(p for p in document.paragraphs if p.text.startswith("npm run test:water-topic"))
    marker_test_paragraph = water_test_paragraph.insert_paragraph_before(
        "npm run test:2d-markers：在高德底图、无人机影像和手绘图三种模式下连续缩放并真实点击图钉，检查图标清晰度和气泡响应稳定性。",
        style=water_test_paragraph.style,
    )
    for run in marker_test_paragraph.runs:
        set_body_font(run)

faq_table = next(table for table in document.tables if table.cell(0, 0).text == "现象")
faq_row = next(
    (row for row in faq_table.rows if row.cells[0].text == "2D图钉有时模糊，或点击后没有出现详情"),
    None,
)
if faq_row is None:
    faq_table._tbl.append(deepcopy(faq_table.rows[-1]._tr))
    faq_row = faq_table.rows[-1]
replace_cell_text(faq_row.cells[0], "2D图钉有时模糊，或点击后没有出现详情")
replace_cell_text(
    faq_row.cells[1],
    "先按Ctrl+F5刷新到最新版本。新版图钉采用固定像素矢量渲染，点击后会保护详情气泡不被底图事件立即关闭；"
    "若仍无响应，先取消部分重叠较密的类型，再点击当前位于最上方的图标。",
)

version_table = document.tables[-1]
if version_table.cell(0, 0).text != "版本":
    # Repair the first execution of this updater, which placed a cloned row
    # before tblGrid instead of after the header row.
    misplaced_data_row = version_table.rows[0]
    misplaced_header_row = version_table.rows[1]
    first_original_data_row = version_table.rows[2]
    first_original_data_row._tr.addprevious(misplaced_header_row._tr)
    first_original_data_row._tr.addprevious(misplaced_data_row._tr)
    replace_cell_text(misplaced_header_row.cells[0], "版本")
    replace_cell_text(misplaced_header_row.cells[1], "日期")
    replace_cell_text(misplaced_header_row.cells[2], "主要变化")
    replace_cell_text(misplaced_data_row.cells[0], "V1.55")
    replace_cell_text(misplaced_data_row.cells[1], "2026年8月9日")
    replace_cell_text(
        misplaced_data_row.cells[2],
        "修复2D图钉偶发模糊和点击无响应：高德覆盖层改用整数像素定位的DOM矢量图标，取消悬停、选中及关联状态下的整体缩放和栅格滤镜；"
        "图钉与水专题对象点击采用短时底图事件保护，避免详情气泡刚打开就被关闭；新增三种2D底图模式下的连续真实点击测试。",
    )

new_version_row = next((row for row in version_table.rows if row.cells[0].text == "V1.55"), None)
if new_version_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_version_row = version_table.rows[1]
replace_cell_text(new_version_row.cells[0], "V1.55")
replace_cell_text(new_version_row.cells[1], "2026年8月9日")
replace_cell_text(
    new_version_row.cells[2],
    "修复2D图钉偶发模糊和点击无响应：高德覆盖层改用整数像素定位的DOM矢量图标，取消悬停、选中及关联状态下的整体缩放和栅格滤镜；"
    "图钉与水专题对象点击采用短时底图事件保护，避免详情气泡刚打开就被关闭；新增三种2D底图模式下的连续真实点击测试。",
)

document.save(MANUAL)

reloaded = Document(MANUAL)
assert any(p.text.startswith("版本：V1.55") for p in reloaded.paragraphs)
assert any("整数屏幕像素定位" in p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("npm run test:2d-markers") for p in reloaded.paragraphs)
assert any(row.cells[0].text == "2D图钉有时模糊，或点击后没有出现详情" for row in faq_table.rows)
assert version_table.rows[1].cells[0].text == "V1.55"
print(MANUAL)
