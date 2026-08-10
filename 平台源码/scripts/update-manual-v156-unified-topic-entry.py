from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)

version_paragraph = next(p for p in document.paragraphs if p.text.startswith("版本：V1."))
replace_paragraph_text(version_paragraph, "版本：V1.56 ｜ 更新日期：2026年8月9日")

overview = next(p for p in document.paragraphs if p.text.startswith("首页默认进入红塘村2D地图"))
replace_paragraph_text(
    overview,
    "首页默认进入红塘村2D地图，便于直接浏览并降低设备性能压力。2D以高德在线地图作为云端底层，可叠加红塘无人机影像或手绘图；"
    "3D以Cesium地形和高斯模型作为底层。左上角平台悬浮框右侧可在“2D地图”和“3D实景”之间切换；进入3D后直接定位红塘村，"
    "不播放从全球飞入村庄的开场动画。两种模式读取同一套地点、点线面专题和调研记录数据，并统一采用“地图作为全屏底层—选择专题—"
    "点击对象—视图移动并缩放—点位旁气泡展示详情”的交互顺序。切换时只替换底层地图，左上角“专题”组件始终保留在首页，已经勾选"
    "或取消的图层会继续生效；同时只运行当前地图，避免2D与3D同时占用显卡。",
)

topic_panel = next(p for p in document.paragraphs if p.text.startswith("桌面端筛选采用地图左上角"))
replace_paragraph_text(
    topic_panel,
    "桌面端和手机端都在地图左上角显示同一张白色圆角“专题”卡片，视觉样式与原水专题入口一致。点击卡片后，在下方展开分类多选列表，"
    "只显示当前共享数据中实际存在的类型，可以同时勾选多个类型；“全部显示”恢复全部类型，数字显示各类型已有要素数量，点、线路和片区都会"
    "计入相应专题。“专题”位于2D与3D之上的首页公共层，切换视图不会重建或重置该组件。",
)

topic_summary = next(p for p in document.paragraphs if p.text.startswith("小花园、茶厂、公共服务"))
replace_paragraph_text(
    topic_summary,
    "小花园、茶厂、公共服务、生态资源、村景记录和村里用水不再分别设置顶部标签或分散入口。进入首页后，点击地图左上角“专题”卡片，"
    "可勾选一个或多个普通分类；“村里用水”这一行右侧提供“进入”按钮，用于连续阅读水源、线路、片区和上下游关系。",
)

water_entry = next(p for p in document.paragraphs if p.text.startswith("地图左下角显示“村里用水专题”入口"))
replace_paragraph_text(
    water_entry,
    "点击地图左上角“专题”卡片，在“村里用水”一行点击“进入”，即可打开水专题。进入后，普通专题卡片在同一位置切换为水专题导航，"
    "2D或3D地图只保留与当前水专题视角有关的对象；点击水专题导航右上角关闭按钮即可退出并恢复普通专题卡片。原地图左下角的孤立入口"
    "已经删除。2D与3D之间切换时，当前水专题视角会继续保留。",
)

topic_table = document.tables[1]
water_topic_row = next(row for row in topic_table.rows if row.cells[0].text == "村里用水")
replace_cell_text(
    water_topic_row.cells[2],
    "从地图左上角“专题”展开列表，在“村里用水”一行点击“进入”，查看2个已有设施、7个待核实节点、8条待核实线路和3个待核实片区",
)

faq_table = next(table for table in document.tables if table.cell(0, 0).text == "现象")
shared_ui_row = next(row for row in faq_table.rows if row.cells[0].text.startswith("为什么2D和3D里的图标"))
replace_cell_text(
    shared_ui_row.cells[1],
    "两种模式读取同一套地点、专题记录和水系统点线面数据，并使用同一个持久的左上角“专题”组件。切换视图只替换底层地图，"
    "因此图层勾选状态不会重置；二维底图和三维场景只是呈现方式不同。",
)
water_faq_row = next(row for row in faq_table.rows if row.cells[0].text == "怎样查看某个片区的水从哪里来")
replace_cell_text(
    water_faq_row.cells[1],
    "点击地图左上角“专题”，在“村里用水”一行点击“进入”，选择“饮水从哪来”，再点击供水片区。气泡会列出供水来源、"
    "经过线路和集中供水点；点击其中任一名称即可继续追踪。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_version_row = next((row for row in version_table.rows if row.cells[0].text == "V1.56"), None)
if new_version_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_version_row = version_table.rows[1]
replace_cell_text(new_version_row.cells[0], "V1.56")
replace_cell_text(new_version_row.cells[1], "2026年8月9日")
replace_cell_text(
    new_version_row.cells[2],
    "删除首页左下角孤立的“村里用水专题”入口；左上角原“筛选要素”改为白色圆角“专题”卡片，点击后展开分类列表，并在“村里用水”"
    "一行提供专题入口。进入水专题时同一位置切换为水专题导航，退出后恢复普通专题卡片；2D、3D和手机端共用同一套状态。",
)

document.save(MANUAL)

reloaded = Document(MANUAL)
assert any(p.text.startswith("版本：V1.56") for p in reloaded.paragraphs)
assert any("原地图左下角的孤立入口已经删除" in p.text for p in reloaded.paragraphs)
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.56"
print(MANUAL)
