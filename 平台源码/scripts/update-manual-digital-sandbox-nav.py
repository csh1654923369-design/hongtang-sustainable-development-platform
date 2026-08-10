from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.text.paragraph import Paragraph


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v15.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.4.backup.docx"
HOME_SCREENSHOT = QA_ROOT / "home-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def insert_paragraph_before(target, text, style):
    element = OxmlElement("w:p")
    target._p.addprevious(element)
    paragraph = Paragraph(element, target._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def set_run_font(run, font_name="SimSun"):
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:cs"), font_name)


def format_paragraph(paragraph, font_name="SimSun"):
    for run in paragraph.runs:
        set_run_font(run, font_name)


def image_paragraph_before(document, caption_text):
    index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == caption_text
    )
    for paragraph in reversed(document.paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError(f"Image not found before caption: {caption_text}")


def replace_image(paragraph, image_path, width_inches):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


QA_ROOT.mkdir(parents=True, exist_ok=True)
if not HOME_SCREENSHOT.exists():
    raise FileNotFoundError(HOME_SCREENSHOT)

document = Document(MANUAL)
version_paragraph = document.paragraphs[5]
if "V1.5 Demo" in version_paragraph.text:
    print("Manual is already V1.5; no changes applied.")
    raise SystemExit(0)
if "V1.4 Demo" not in version_paragraph.text:
    raise ValueError(f"Expected V1.4 manual, got: {version_paragraph.text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

version_paragraph.text = "版本：V1.5 Demo   |   更新日期：2026年7月19日   |   适用地址：http://localhost:3000"
format_paragraph(version_paragraph)

profile_navigation = find_paragraph(
    document,
    "个人中心：查看我的上报、我发起或加入的行动、报名、关注和通知。",
)
new_navigation = insert_paragraph_before(
    profile_navigation,
    "数字沙盘：从顶部导航进入，查看重点区域、图层和改造方案的概念演示。",
    "List Bullet",
)
format_paragraph(new_navigation)

first_sandbox_step = find_paragraph(document, "勾选或取消建筑、项目、绿化图层。")
first_sandbox_step.text = "从顶部导航点击“数字沙盘”进入页面，然后勾选或取消建筑、项目、绿化图层。"
format_paragraph(first_sandbox_step)

replace_image(image_paragraph_before(document, "图 1  平台首页"), HOME_SCREENSHOT, 6.3)

version_table = next(
    table
    for table in document.tables
    if [cell.text.strip() for cell in table.rows[0].cells] == ["版本", "日期", "主要变化"]
)
row = version_table.add_row()
values = [
    "V1.5 Demo",
    "2026-07-19",
    "在桌面端顶部导航补充“数字沙盘”入口，并同步到移动端菜单和平台搜索；可直接进入 /digital-twin 页面。",
]
for index, value in enumerate(values):
    cell = row.cells[index]
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        format_paragraph(paragraph)

document.save(OUTPUT)

check = Document(OUTPUT)
assert "V1.5 Demo" in check.paragraphs[5].text
assert any("数字沙盘：从顶部导航进入" in paragraph.text for paragraph in check.paragraphs)
assert any("从顶部导航点击“数字沙盘”" in paragraph.text for paragraph in check.paragraphs)
assert any(row.cells[0].text == "V1.5 Demo" for table in check.tables for row in table.rows)
assert len(check.inline_shapes) == len(document.inline_shapes)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual: {MANUAL}")
print(f"Backup: {BACKUP}")
print(f"Paragraphs: {len(check.paragraphs)} | tables: {len(check.tables)} | images: {len(check.inline_shapes)}")
