from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.67 ｜ 更新日期：2026年8月10日",
)

for paragraph in document.paragraphs:
    text = paragraph.text
    updated = text.replace(
        "“底图”“卫星遥感”“无人机影像”“手绘图”四种按钮",
        "“航拍”“手绘图”“卫星遥感”“底图”四种按钮",
    ).replace(
        "可选择“底图”“卫星遥感”“无人机影像”或“手绘图”",
        "可选择“航拍”“手绘图”“卫星遥感”或“底图”",
    ).replace(
        "在底图、卫星遥感、无人机影像和手绘图四种模式下",
        "在航拍、手绘图、卫星遥感和底图四种模式下",
    )
    if updated != text:
        replace_paragraph_text(paragraph, updated)

overview = next(p for p in document.paragraphs if p.text.startswith("首页默认进入红塘村2D地图"))
if "首次打开网页时默认选中“航拍”" not in overview.text:
    replace_paragraph_text(
        overview,
        overview.text + " 首次打开网页时默认选中“航拍”；高德“底图”仅做轻微淡化处理，不影响航拍、手绘图和卫星遥感。",
    )

topic_intro = next(p for p in document.paragraphs if p.text.startswith("桌面端和手机端都在地图左上角"))
replace_paragraph_text(
    topic_intro,
    "桌面端和手机端都在地图左上角显示一张白色圆角专题大卡片，标题栏与具体内容不再分成两张卡片。点击右侧带倒三角的“展开”后，卡片向下平滑延展，专题内容同时以透明度渐变显现；打开后按钮变为“收起”，点击时执行反向过渡。内容区直接显示小花园、茶产业、村里用水、塌方与安全、历史与文化五个专题，下面另列公共服务和村景记录等其他资料，不重复显示“专题、全选、完成”标题行。各项可多选，每个专题右侧的“进入”用于切换到专题阅读状态；数字表示已接入的空间资料数量，没有已核实资料的专题显示“待调查”。切换2D与3D不会重置当前状态。平台名称卡片与专题大卡片保持相同宽度，左右边缘对齐。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.67"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.67",
    "2026年8月10日",
    "二维图层改为航拍、手绘图、卫星遥感、底图并默认航拍；专题改为带渐变动画的一体化卡片；隐藏本地开发“N”按钮。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.67") for p in reloaded.paragraphs)
assert "“航拍”“手绘图”“卫星遥感”“底图”四种按钮" in paragraph_text
assert "首次打开网页时默认选中“航拍”" in paragraph_text
assert "卡片向下平滑延展" in paragraph_text
assert "透明度渐变显现" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.67"
print(MANUAL)
