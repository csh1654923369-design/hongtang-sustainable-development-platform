from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)

version_paragraph = next(p for p in document.paragraphs if p.text.startswith("版本：V1."))
replace_paragraph_text(version_paragraph, "版本：V1.57 ｜ 更新日期：2026年8月9日")

data_overview = next(p for p in document.paragraphs if p.text.startswith("当前基础地点数据包含205个有效坐标点"))
replace_paragraph_text(
    data_overview,
    "当前基础地点数据包含205个有效坐标点和586张现场照片；首页另载入7个待核实水系统节点、8条待核实线路和3个待核实供水片区，"
    "共形成223个地图要素。205个基础地点中的名称、分类、坐标和照片来自现有资料。586张照片已转换为平台本地WebP文件，2D与3D详情气泡"
    "读取同一套本地照片，不再依赖原照片服务器；水专题点线面用于展示未来的数据关系和调查方式。",
)

data_table = next(table for table in document.tables if table.cell(0, 0).text == "内容")
location_row = next(row for row in data_table.rows if row.cells[0].text == "地点数据")
replace_cell_text(
    location_row.cells[1],
    "平台源码/public/data/hongtang-real-map-features.json；平台源码/public/local-photos；平台素材/现场照片",
)
replace_cell_text(
    location_row.cells[2],
    "二维和三维地图共用的地点分类、坐标、简介及586张本地WebP照片；来源网址仅保存在照片来源清单中用于核验和重新生成",
)

faq_table = next(table for table in document.tables if table.cell(0, 0).text == "现象")
photo_row = next(row for row in faq_table.rows if row.cells[0].text == "地点资料没有照片")
replace_cell_text(
    photo_row.cells[1],
    "该记录可能原本就没有录入照片，或本地照片文件缺失。正常情况下，已有的586张照片直接从平台本地加载，不需要访问原照片服务器；"
    "平台不会自动补一张虚构图片。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_version_row = next((row for row in version_table.rows if row.cells[0].text == "V1.57"), None)
if new_version_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_version_row = version_table.rows[1]
replace_cell_text(new_version_row.cells[0], "V1.57")
replace_cell_text(new_version_row.cells[1], "2026年8月9日")
replace_cell_text(
    new_version_row.cells[2],
    "将205个地点关联的586张现场照片下载并压缩为本地WebP素材；2D与3D详情气泡统一读取平台本地照片，不再依赖三农数据照片服务器。"
    "原始网址单独保留在来源清单中，便于核验和重新生成。",
)

document.save(MANUAL)

reloaded = Document(MANUAL)
all_text = "\n".join(p.text for p in reloaded.paragraphs)
all_table_text = "\n".join(cell.text for table in reloaded.tables for row in table.rows for cell in row.cells)
assert any(p.text.startswith("版本：V1.57") for p in reloaded.paragraphs)
assert "586张照片已转换为平台本地WebP文件" in all_text
assert "平台源码/public/local-photos" in all_table_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.57"
print(MANUAL)
