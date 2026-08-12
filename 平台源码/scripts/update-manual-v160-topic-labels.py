from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


document = Document(MANUAL)

version_paragraph = next(p for p in document.paragraphs if p.text.startswith("版本：V1."))
replace_paragraph_text(version_paragraph, "版本：V1.60 ｜ 更新日期：2026年8月10日")

replacements = (
    ("茶马古道与村庄历史", "历史与文化"),
    ("地图基础资料", "其他资料"),
    ("五个红塘专题", "五个专题"),
    ("红塘专题", "专题"),
)

for paragraph in document.paragraphs:
    updated = paragraph.text
    for old, new in replacements:
        updated = updated.replace(old, new)
    if updated != paragraph.text:
        replace_paragraph_text(paragraph, updated)

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                updated = paragraph.text
                for old, new in replacements:
                    updated = updated.replace(old, new)
                if updated != paragraph.text:
                    replace_paragraph_text(paragraph, updated)
                    for run in paragraph.runs:
                        set_body_font(run)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.60"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.60",
    "2026年8月10日",
    "精简首页专题文字：取消专题列表中的分组小标题，将最后一个专题改为“历史与文化”，将资料分组改为“其他资料”。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)
    for run in cell.paragraphs[0].runs:
        set_body_font(run)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
table_text = "\n".join(cell.text for table in reloaded.tables for row in table.rows for cell in row.cells)
assert any(p.text.startswith("版本：V1.60") for p in reloaded.paragraphs)
assert "4.4 塌方与安全、历史与文化专题" in paragraph_text
assert "其他资料" in paragraph_text
assert "茶马古道与村庄历史" not in paragraph_text + table_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.60"
print(MANUAL)
