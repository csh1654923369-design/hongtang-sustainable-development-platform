from copy import deepcopy
from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.text.paragraph import Paragraph


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v16.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.5.backup.docx"
HOME_SCREENSHOT = QA_ROOT / "home-1440.png"
VILLAGE_LIFE_SCREENSHOT = QA_ROOT / "village-life-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def insert_paragraph_before(target, text="", style=None):
    element = OxmlElement("w:p")
    target._p.addprevious(element)
    paragraph = Paragraph(element, target._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def remove_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)
    paragraph._p = paragraph._element = None


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)


def set_image(paragraph, image_path, width, name, description):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("name", name)
    shape._inline.docPr.set("descr", description)
    return shape


def image_paragraph_before(document, caption_text):
    caption = find_paragraph(document, caption_text)
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError(f"Image paragraph not found before caption: {caption_text}")


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def set_cell_text_like(source_cell, target_cell, text):
    copy_cell_format(source_cell, target_cell)
    target_cell.text = text
    target_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if source_cell.paragraphs and target_cell.paragraphs:
        source = source_cell.paragraphs[0]
        target = target_cell.paragraphs[0]
        target.style = source.style
        target.paragraph_format.alignment = source.paragraph_format.alignment


def add_row_like(table, template_row, values, before_row=None):
    row = table.add_row()
    for index, value in enumerate(values):
        set_cell_text_like(template_row.cells[index], row.cells[index], value)
    if before_row is not None:
        before_row._tr.addprevious(row._tr)
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        east_asia = "SimHei" if is_heading or index == 0 else "SimSun"
        for run in paragraph.runs:
            set_run_font(run, east_asia)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
                    for run in paragraph.runs:
                        set_run_font(run, "SimHei" if is_heading else "SimSun")


def keep_table_rows_intact(table):
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        if not properties.xpath("./w:cantSplit"):
            properties.append(OxmlElement("w:cantSplit"))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    if not header_properties.xpath("./w:tblHeader"):
        header_properties.append(OxmlElement("w:tblHeader"))


for screenshot in (HOME_SCREENSHOT, VILLAGE_LIFE_SCREENSHOT):
    if not screenshot.exists():
        raise FileNotFoundError(screenshot)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
version_text = document.paragraphs[5].text
if "V1.6 Demo" in version_text:
    set_image(
        image_paragraph_before(document, "图 1  平台首页"),
        HOME_SCREENSHOT,
        6.3,
        "图 1 平台首页",
        "村民优先的首页，主导航缩减为四项，并提供清晰常用入口",
    )
    set_image(
        image_paragraph_before(document, "图 2  “村里的事”六类长期记录模块"),
        VILLAGE_LIFE_SCREENSHOT,
        6.3,
        "图 2 村里的事",
        "村里的事页面包含六类长期记录模块和可切换的详情区",
    )
    normalize_fonts(document)
    document.save(OUTPUT)
    check = Document(OUTPUT)
    assert len(check.inline_shapes) == 13
    os.replace(OUTPUT, MANUAL)
    print(f"Refreshed V1.6 screenshots: {MANUAL}")
    raise SystemExit(0)
if "V1.5 Demo" not in version_text:
    raise ValueError(f"Expected V1.5 manual, got: {version_text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

# Cover, startup path, and the new homepage screenshot.
document.paragraphs[5].text = "版本：V1.6 Demo   |   更新日期：2026年7月21日   |   适用地址：http://localhost:3000"
set_image(
    image_paragraph_before(document, "图 1  平台首页"),
    HOME_SCREENSHOT,
    6.3,
    "图 1 平台首页",
    "村民优先的首页，主导航缩减为四项，并提供清晰常用入口",
)
find_paragraph(document, "打开项目文件夹：E:\\BaiduSyncdisk\\研一下\\红塘可持续发展平台。").text = (
    "打开源码文件夹：E:\\BaiduSyncdisk\\研一下\\红塘可持续发展平台\\平台源码。"
)

# Replace nine expert-oriented navigation bullets with five village-first explanations.
old_navigation = [
    "首页：了解平台定位，并从行动地图、发起微行动或公众参与快速开始。",
    "认识红塘：查看村庄资料框架、故事与发展时间线。",
    "可持续目标：查看五个本地目标、指标和关联项目。",
    "行动地图：按行动、问题、资源供需和空间资料分组筛选点位。",
    "项目与行动：浏览居民微行动，或筛选正式项目并查看开放任务与资源需求。",
    "发展进展：查看趋势、来源、完整度和社区行动能力。",
    "公众参与：发起微行动，提出建议、报名活动、填写问卷并参与共创讨论。",
    "数字沙盘：从顶部导航进入，查看重点区域、图层和改造方案的概念演示。",
    "个人中心：查看我的上报、我发起或加入的行动、报名、关注和通知。",
]
navigation_paragraphs = [find_paragraph(document, text) for text in old_navigation]
new_navigation = [
    "首页：从“记录一件事、看看最近变化、参加村里活动”三个大入口开始。",
    "村里的事：查看小花园、茶园与茶厂、村里用水、光伏设施、塌方与安全、红塘村历史。",
    "村里一张图：在地图上找地点、问题、项目、互助资源和正在开展的事情。",
    "我的：查看与当前演示身份有关的记录、活动、关注和通知。",
    "更多功能：认识红塘、可持续目标、项目与行动、发展进展、研究协作和数字沙盘仍可从首页、移动菜单、搜索或页脚进入。",
]
for paragraph, text in zip(navigation_paragraphs[:5], new_navigation):
    paragraph.text = text
for paragraph in navigation_paragraphs[5:]:
    remove_paragraph(paragraph)

find_paragraph(document, "搜索：点击放大镜，搜索演示项目、目标、问题或功能入口。").text = (
    "搜索：点击放大镜，搜索村里的事、地点、项目或其他功能入口。"
)
find_paragraph(document, "我要上报：村民或管理员可直接进入五步问题上报。").text = (
    "我要记录：村民或管理员可直接进入五步问题记录与上报。"
)
find_paragraph(document, "在手机或窄屏浏览器中，页面底部会显示“首页、地图、参与、项目、我的”五个快捷入口；其他页面可通过右上角菜单访问。网站已针对 390px 宽度进行无横向溢出测试。").text = (
    "在手机或窄屏浏览器中，页面底部只保留“首页、村里的事、一张图、我的”四个快捷入口；其他页面可通过右上角菜单访问。网站已针对 390px 宽度进行无横向溢出测试。"
)

# Add a concise guide and screenshot for the new village-matters page.
overview_heading = find_paragraph(document, "2.4 页面功能总览")
overview_heading.text = "2.5 页面功能总览"
insert_paragraph_before(overview_heading, "2.4 使用“村里的事”", "Heading 2")
insert_paragraph_before(overview_heading, "在顶部导航或手机底栏点击“村里的事”。", "List Number")
insert_paragraph_before(overview_heading, "从六个大按钮中选择一类：我家小花园、茶园与茶厂、村里用水、光伏设施、塌方与安全或红塘村历史。", "List Number")
insert_paragraph_before(overview_heading, "页面下方会显示这类事项最近在记录什么、以后要连续记录什么，以及相关地图或项目入口。", "List Number")
insert_paragraph_before(overview_heading, "点击“我要记录一件事”可进入上报流程；只想浏览时，可继续查看“村里最近有什么新情况”。", "List Number")
image_paragraph = insert_paragraph_before(overview_heading)
set_image(
    image_paragraph,
    VILLAGE_LIFE_SCREENSHOT,
    6.3,
    "图 2 村里的事",
    "村里的事页面包含六类长期记录模块和可切换的详情区",
)
insert_paragraph_before(overview_heading, "图 2  “村里的事”六类长期记录模块", "Caption")
insert_paragraph_before(
    overview_heading,
    "使用提示\n页面目前显示演示记录。真实小花园、茶园、用水、光伏、安全和历史资料，需要在村民、村委和调研团队确认后逐步补充。",
    "Normal",
)

# Use the same everyday language on the map chapter.
find_paragraph(document, "3. 行动地图").text = "3. 村里一张图"
find_paragraph(document, "地图页面分为三部分：上方行动地图汇集问题、正式项目、社区微行动、已完成行动和资源供需点位；中部互助板按空间、工具、材料、技能、地方知识和志愿时间展示可提供资源与需求；下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。除建筑轮廓外，坐标与业务点位均为演示数据。").text = (
    "“村里一张图”分为三部分：上方地图汇集问题、正式项目、社区微行动、已完成行动和资源供需点位；中部互助板展示可提供资源与需求；下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。除建筑轮廓外，坐标与业务点位均为演示数据。"
)
find_paragraph(document, "图 2  行动地图与点位详情").text = "图 3  村里一张图与点位详情"
find_paragraph(document, "打开“行动地图”。").text = "打开“村里一张图”。"
find_paragraph(document, "3.4 从地图上报问题").text = "3.4 从一张图记录问题"
find_paragraph(document, "切换为村民后，点击地图右下角“上报问题”或页面顶部“我要上报”。游客点击时会收到权限提示。").text = (
    "切换为村民后，点击地图右下角“上报问题”或页面顶部“我要记录”。游客点击时会收到权限提示。"
)

# Digital sandbox remains available, but is no longer a primary village navigation item.
find_paragraph(document, "从顶部导航点击“数字沙盘”进入页面，然后勾选或取消建筑、项目、绿化图层。").text = (
    "从首页“更多资料”、右上角移动菜单、搜索或相关项目进入“看看改造后的样子（数字沙盘）”，然后勾选或取消建筑、项目、绿化图层。"
)
find_paragraph(document, "src/data/mockData.ts 提供目标、项目、问题和指标；src/data/communityData.ts 提供微行动与互助资源演示种子。").text = (
    "src/data/mockData.ts 提供目标、项目、问题和指标；src/data/villageMatters.ts 提供六类“村里的事”；src/data/communityData.ts 提供微行动与互助资源演示种子。"
)

# Renumber all captions following the inserted figure.
caption_updates = {
    "图 3  社区资源地图与互助板": "图 4  社区资源地图与互助板",
    "图 4  建筑调研底图与建筑档案": "图 5  建筑调研底图与建筑档案",
    "图 5  五步问题上报向导": "图 6  五步问题上报向导",
    "图 6  社区微行动与伙伴招募": "图 7  社区微行动与伙伴招募",
    "图 7  五步发起社区微行动": "图 8  五步发起社区微行动",
    "图 8  公众参与中心": "图 9  公众参与中心",
    "图 9  社区行动能力指标": "图 10  社区行动能力指标",
    "图 10  发展进展与趋势图": "图 11  发展进展与趋势图",
    "图 11  管理后台数据概览": "图 12  管理后台数据概览",
    "图 12  数字沙盘概念演示": "图 13  数字沙盘概念演示",
}
for old, new in caption_updates.items():
    find_paragraph(document, old).text = new

# Role, page overview, and version tables.
role_table = document.tables[0]
role_table.rows[1].cells[1].text = "浏览首页、村里的事、村里一张图、项目和公开进展"
role_table.rows[1].cells[2].text = "首页、村里的事、村里一张图"
role_table.rows[2].cells[2].text = "首页、村里的事、村里一张图、参与、我的"
role_table.rows[3].cells[2].text = "村里的事、村里一张图、项目、研究协作、我的"

page_table = document.tables[1]
home_row = next(row for row in page_table.rows if row.cells[0].text.strip() == "首页")
home_row.cells[2].text = "三个常用入口、六类村庄事项、最近变化和地图预览"
village_row = next(row for row in page_table.rows if row.cells[0].text.strip() == "认识红塘")
add_row_like(
    page_table,
    home_row,
    ["村里的事", "/village-life", "六类长期记录、最近变化和关联入口"],
    before_row=village_row,
)
map_row = next(row for row in page_table.rows if row.cells[0].text.strip() == "行动地图")
map_row.cells[0].text = "村里一张图"

version_table = document.tables[4]
add_row_like(
    version_table,
    version_table.rows[-1],
    [
        "V1.6 Demo",
        "2026-07-21",
        "主导航简化为首页、村里的事、村里一张图和我的；首页改为三个村民常用入口，新增六类村庄事项，并将数字沙盘转移到更多功能。",
    ],
)

for table in document.tables:
    keep_table_rows_intact(table)

normalize_fonts(document)

# Normalize image names and alt text in document order.
captions = [
    paragraph.text.strip()
    for paragraph in document.paragraphs
    if paragraph.style and paragraph.style.name == "Caption" and paragraph.text.strip().startswith("图 ")
]
if len(captions) != len(document.inline_shapes):
    raise AssertionError(f"Caption/image mismatch before save: {len(captions)} captions, {len(document.inline_shapes)} images")
for index, (shape, caption) in enumerate(zip(document.inline_shapes, captions), start=1):
    shape._inline.docPr.set("name", f"Figure {index}")
    shape._inline.docPr.set("descr", caption)

document.save(OUTPUT)

# Structural gate before replacing the formal manual. Rendering is intentionally
# skipped because this project explicitly does not use LibreOffice rendering.
check = Document(OUTPUT)
assert "V1.6 Demo" in check.paragraphs[5].text
assert len(check.inline_shapes) == 13
assert len([p for p in check.paragraphs if p.style and p.style.name == "Caption" and p.text.strip().startswith("图 ")]) == 13
assert any(paragraph.text == "2.4 使用“村里的事”" for paragraph in check.paragraphs)
assert any(paragraph.text == "3. 村里一张图" for paragraph in check.paragraphs)
assert any(row.cells[0].text.strip() == "村里的事" for row in check.tables[1].rows)
assert any(row.cells[0].text.strip() == "村里一张图" for row in check.tables[1].rows)
assert check.tables[4].rows[-1].cells[0].text == "V1.6 Demo"
assert all(row._tr.xpath("./w:trPr/w:cantSplit") for table in check.tables for row in table.rows)
assert all(
    run.font.color.rgb == RGBColor(0, 0, 0)
    for paragraph in check.paragraphs
    for run in paragraph.runs
    if run.text
)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual: {MANUAL}")
print(f"Backup: {BACKUP}")
print(f"Paragraphs: {len(check.paragraphs)} | tables: {len(check.tables)} | images: {len(check.inline_shapes)}")
