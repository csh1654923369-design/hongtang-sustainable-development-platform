from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.62 ｜ 更新日期：2026年8月10日",
)

for paragraph in document.paragraphs:
    if "回到红塘" in paragraph.text:
        replace_paragraph_text(paragraph, paragraph.text.replace("回到红塘", "回到中心"))

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if "回到红塘" in paragraph.text:
                    replace_paragraph_text(paragraph, paragraph.text.replace("回到红塘", "回到中心"))
                    for run in paragraph.runs:
                        set_body_font(run)

overview = next(p for p in document.paragraphs if p.text.startswith("首页默认进入红塘村2D地图"))
base_text = overview.text.split("右下角工具", 1)[0].rstrip()
replace_paragraph_text(
    overview,
    base_text
    + " 右下角工具在2D和3D之间采用同一视觉样式：44像素深色圆角图标按钮、相同阴影和左侧悬停提示。"
    + "2D保留“回到中心”，3D提供查看全部地点、回到中心、操作设置和全屏查看。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.62"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.62",
    "2026年8月10日",
    "统一2D与3D右下角工具样式：采用同尺寸深色圆角图标、相同阴影和悬停提示；2D复位入口改名为“回到中心”。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)
    for run in cell.paragraphs[0].runs:
        set_body_font(run)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
table_text = "\n".join(cell.text for table in reloaded.tables for row in table.rows for cell in row.cells)
assert any(p.text.startswith("版本：V1.62") for p in reloaded.paragraphs)
assert "右下角工具在2D和3D之间采用同一视觉样式" in paragraph_text
assert "2D保留“回到中心”" in paragraph_text
assert "回到红塘" not in paragraph_text + table_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.62"
print(MANUAL)
