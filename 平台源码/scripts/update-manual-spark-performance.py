from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v127.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.26.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def add_version_row(table):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        properties = cell._tc.get_or_add_tcPr()
        for child in list(properties):
            properties.remove(child)
        for child in template_row.cells[index]._tc.tcPr:
            properties.append(deepcopy(child))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        cell.paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.26 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.26 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.27 Demo   |   更新日期：2026年7月28日   |   适用地址：http://localhost:3000"
)

find_paragraph_prefix(document, "阅读提示").text = (
    "阅读提示\n当前网站是可交互演示版本。正式首页使用SparkJS 2.1.0按当前视角加载红塘村轻量化"
    "3D高斯模型。首页展示小花园、茶场、茶厂、村里用水和光伏设施5个可点击示例点。左键拖动沿模型"
    "水平面平移；右键围绕当前屏幕中心目标点旋转；滚轮按固定比例缩放，中键拖动已停用。点击图钉只"
    "打开事项详情，点击“定位到此处”才让镜头飞近。默认“均衡”档以约35万个可见高斯点和最高30帧"
    "运行；停止操作和模型细化后，页面会完全停止连续绘制。电脑配置较低时可在右下角齿轮中选择“省电”。"
    "所有三维图钉均为演示位置，待实地核实。"
)

find_paragraph_prefix(document, "3D首页操作：").text = (
    "3D首页操作：打开首页后，SparkJS先读取轻量化模型索引，再根据相机位置载入当前视角需要的模型分块。"
    "界面右下角纵向排列“回到中心、操作设置、全屏查看”三个图标按钮。点击齿轮可以分别调整左键平移和"
    "滚轮缩放灵敏度，并在“省电、均衡、清晰”三档画面质量之间切换。桌面端三档分别以约18万、35万和"
    "65万个可见高斯点为目标，最高帧率分别为24、30和40帧；移动端会进一步降低。左键拖动沿模型水平面"
    "平移；右键围绕当前屏幕中心目标点旋转；滚轮每次按固定比例缩放，中键不执行操作。查看器使用"
    "唯一请求式刷新队列，停止操作和分块细化后会完全暂停连续绘制。点击事项图钉只打开详情，点击“定位到此处”"
    "才移动镜头。"
)

find_paragraph_prefix(document, "正式首页不再使用Cesium地球").text = (
    "正式首页不再使用Cesium地球、世界地形或在线高斯资产。恢复工具从现有Cesium 3D Tiles的1,036个末级"
    "瓦片提取22,519,595点标准高斯PLY，再用Spark build-lod生成476个流式分块。性能修复后，“均衡”档"
    "当前视角实测请求25个模型响应、约31.08MiB，稳定后等待两秒新增渲染帧为0。重复动画循环已经消除，"
    "分块缓存由56页下调为24页，默认高斯点预算由约80万降至35万，帧率上限为30帧。模型仍使用"
    "Three.js局部坐标完成居中；原三维场景内的Cesium地形、全球卫星影像和单独正射影像不参与首页渲染，"
    "二维“村庄总览”和“村里一张图”继续保留无人机正射影像。"
)

find_paragraph_prefix(document, "轻量化与正式接入：").text = (
    "轻量化与正式接入：从现有Cesium 3D Tiles恢复出22,519,595点标准高斯PLY，再使用Spark 2.1.0生成"
    "476个流式分块。RAD索引约54KiB，分块合计531.70MiB，比原735.31MiB瓦片包减少约27.7%。正式首页"
    "通过受限Range接口按当前视角加载；性能优化后的“均衡”档首屏实测约31.08MiB，并在静止后停止连续"
    "绘制。轻量化解决网络与存储负担，点数预算、分辨率、帧率上限和单一刷新队列共同控制显卡负担。"
)

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[1].text = "SparkJS 2.1.0 + Three.js 0.180.0"
        row.cells[2].text = (
            "RAD/RADC按视角流式加载；三档分别控制可见高斯点、帧率上限和像素倍率；唯一请求式刷新队列"
            "避免重复动画循环，静止后停止绘制；保留5个图钉、详情、定位、相机操作和全屏功能"
        )
        break
else:
    raise ValueError("3D technology row not found")

version_row = add_version_row(document.tables[4])
for cell, value in zip(
    version_row.cells,
    [
        "V1.27 Demo",
        "2026-07-28",
        "修复SparkJS重复渲染循环；降低三档点数、像素倍率和帧率上限；缩小分块缓存，并增加静止零持续渲染检查。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 29]
assert "V1.27 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "默认“均衡”档以约35万个可见高斯点" in all_text
assert "稳定后等待两秒新增渲染帧为0" in all_text
assert "唯一请求式刷新队列" in all_text
assert any(row.cells[0].text.strip() == "V1.27 Demo" for row in check.tables[4].rows)
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count
assert check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
assert check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.27: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
