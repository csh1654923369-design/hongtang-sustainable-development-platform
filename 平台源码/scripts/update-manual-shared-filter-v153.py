from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"

REPLACEMENTS = {
    "版本：V1.52 ｜ 更新日期：2026年8月9日": "版本：V1.53 ｜ 更新日期：2026年8月9日",
    "首页默认进入红塘村2D地图，便于直接浏览并降低设备性能压力。左上角平台悬浮框右侧可在“2D地图”和“3D实景”之间切换；进入3D后直接定位红塘村，不播放从全球飞入村庄的开场动画。两种模式读取同一套地点、点线面专题和调研记录数据，并统一采用“地图作为全屏底层—筛选要素—点击对象—视图移动并缩放—点位旁气泡展示详情”的交互顺序；切换只改变二维或三维的呈现方式，并且只运行当前地图，避免同时占用显卡。": "首页默认进入红塘村2D地图，便于直接浏览并降低设备性能压力。左上角平台悬浮框右侧可在“2D地图”和“3D实景”之间切换；进入3D后直接定位红塘村，不播放从全球飞入村庄的开场动画。两种模式读取同一套地点、点线面专题和调研记录数据，并统一采用“地图作为全屏底层—筛选要素—点击对象—视图移动并缩放—点位旁气泡展示详情”的交互顺序。切换时只替换底层地图，“筛选要素”组件始终保留在首页，不会重新加载，已经勾选或取消的图层会继续生效；同时只运行当前地图，避免2D与3D同时占用显卡。",
    "桌面端筛选采用地图左上角的深色悬浮多选列表，只显示当前共享数据中实际存在的类型，可以同时勾选多个类型；手机端2D和3D均通过“筛选要素”按钮打开或关闭同一组分类。“全选”恢复全部类型，“清空”取消全部类型，数字统计当前可见地图要素，点、线路和片区都会计入相应专题。": "桌面端筛选采用地图左上角的深色悬浮多选列表，只显示当前共享数据中实际存在的类型，可以同时勾选多个类型；手机端2D和3D通过同一个“筛选要素”按钮打开或关闭分类列表。“筛选要素”位于2D与3D之上的首页公共层，切换视图不会关闭、重建或重置该组件。“全选”恢复全部类型，数字显示各类型已有要素数量，点、线路和片区都会计入相应专题。",
}


def find_paragraph(document: Document, exact: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == exact]
    if len(matches) != 1:
        raise ValueError(f"应找到1处段落，实际找到{len(matches)}处：{exact}")
    return matches[0]


def set_two_character_indent(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), "200")


def format_run(run) -> None:
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(0, 0, 0)


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    format_run(paragraph.add_run(text))
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    format_run(paragraph.add_run(text))


document = Document(MANUAL)
for original, updated in REPLACEMENTS.items():
    replace_paragraph(find_paragraph(document, original), updated)

faq_matches = [row for row in document.tables[4].rows if row.cells[0].text.strip() == "为什么2D和3D里的图标、筛选和详情相同"]
if len(faq_matches) != 1:
    raise ValueError(f"应找到1条2D/3D常见问题，实际找到{len(faq_matches)}条")
set_cell_text(faq_matches[0].cells[1], "两种模式读取同一套地点、专题记录和水系统点线面数据，并使用同一个持久的“筛选要素”组件。切换视图只替换底层地图，因此图层勾选状态不会重置；二维底图和三维场景只是呈现方式不同。")

version_table = document.tables[5]
template_row = deepcopy(version_table.rows[1]._tr)
version_table.rows[0]._tr.addnext(template_row)
new_row = version_table.rows[1]
set_cell_text(new_row.cells[0], "V1.53")
set_cell_text(new_row.cells[1], "2026年8月9日")
set_cell_text(new_row.cells[2], "将“筛选要素”提升为2D与3D共用的首页持久组件；切换视图只替换底层地图，筛选组件不重新加载，各图层勾选状态持续生效；隐藏3D查看器内部的重复筛选面板，并统一桌面端与手机端入口。")

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 用户已明确不需要LibreOffice渲染，因此只执行结构、内容与压缩包完整性校验。
reopened = Document(MANUAL)
paragraph_texts = [paragraph.text.strip() for paragraph in reopened.paragraphs]
assert "版本：V1.53 ｜ 更新日期：2026年8月9日" in paragraph_texts
assert any("筛选要素”组件始终保留在首页" in text for text in paragraph_texts)
assert reopened.tables[5].rows[1].cells[0].text.strip() == "V1.53"
assert "图层勾选状态不会重置" in faq_matches[0].cells[1].text
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(f"updated={MANUAL}")
print(f"version={reopened.tables[5].rows[1].cells[0].text.strip()}")
