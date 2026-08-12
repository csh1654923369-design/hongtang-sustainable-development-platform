from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.64 ｜ 更新日期：2026年8月10日",
)

overview = next(p for p in document.paragraphs if p.text.startswith("首页默认进入红塘村2D地图"))
overview_text = overview.text.replace(
    "选中2D时，其下方显示“高德地图”“卫星遥感”“无人机影像”“手绘图”四种底图按钮；",
    "选中2D时，其下方以独立卡片纵向显示“高德地图”“卫星遥感”“无人机影像”“手绘图”四种底图按钮；",
).replace(
    "2D保留“回到中心”，3D提供查看全部地点、回到中心、操作设置和全屏查看。",
    "2D保留“回到中心、全屏查看”；3D仅保留“操作设置、回到中心、全屏查看”，并将操作设置放在最上方。",
)
replace_paragraph_text(overview, overview_text)

two_d = next(p for p in document.paragraphs if p.text.startswith("2D模式与3D实景读取同一套地点和专题数据"))
two_d_text = two_d.text.replace(
    "右上角第一层切换2D/3D；选中2D后，第二层可选择“高德地图”“卫星遥感”“无人机影像”或“手绘图”。",
    "右上角保留独立的2D/3D切换卡片；选中2D后，其下方另设纵向底图列表，可选择“高德地图”“卫星遥感”“无人机影像”或“手绘图”。",
).replace(
    "点击右下角“回到中心”恢复村庄中心视图。",
    "点击右下角“回到中心”恢复村庄中心视图；点击“全屏查看”可进入或退出浏览器全屏。",
)
replace_paragraph_text(two_d, two_d_text)

fullscreen = next(p for p in document.paragraphs if p.text.startswith("全屏查看："))
replace_paragraph_text(
    fullscreen,
    "全屏查看：2D和3D右下角均保留全屏图标，点击后进入或退出浏览器全屏。鼠标悬停在图标上可查看功能名称。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.64"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.64",
    "2026年8月10日",
    "2D底图按钮改为右上角独立纵向列表；2D右下角保留回到中心和全屏，3D仅保留设置、回到中心和全屏，并将设置置顶。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)
    for run in cell.paragraphs[0].runs:
        set_body_font(run)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.64") for p in reloaded.paragraphs)
assert "其下方以独立卡片纵向显示" in paragraph_text
assert "2D保留“回到中心、全屏查看”" in paragraph_text
assert "3D仅保留“操作设置、回到中心、全屏查看”" in paragraph_text
assert "查看全部地点" not in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.64"
print(MANUAL)
