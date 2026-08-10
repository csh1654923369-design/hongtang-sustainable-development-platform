from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.text.paragraph import Paragraph


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v116.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.15.backup.docx"
HOME_SCREENSHOT = QA_ROOT / "home-1440.png"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


def set_home_image(paragraph):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(HOME_SCREENSHOT), width=Inches(6.3))
    shape._inline.docPr.set("name", "图 1 3D高斯实景首页与五类示例点")
    shape._inline.docPr.set(
        "descr",
        "红塘村3D高斯实景首页展示小花园、茶场、茶厂、村里用水和光伏设施五类可点击演示图钉",
    )


def insert_paragraph_after(paragraph, text, style_name):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.style = style_name
    new_paragraph.add_run(text)
    return new_paragraph


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        copy_cell_format(template_row.cells[index], cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if template_row.cells[index].paragraphs and cell.paragraphs:
            cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
            cell.paragraphs[0].paragraph_format.alignment = (
                template_row.cells[index].paragraphs[0].paragraph_format.alignment
            )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_first_line_chars(paragraph, chars=200):
    properties = paragraph._p.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    for attribute in ("firstLine", "hanging", "hangingChars"):
        key = qn(f"w:{attribute}")
        if key in indent.attrib:
            del indent.attrib[key]
    indent.set(qn("w:firstLineChars"), str(chars))


def normalize_document(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")

    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
        if (
            index >= 8
            and paragraph.style
            and paragraph.style.name == "Normal"
            and paragraph.text.strip()
            and "\n" not in paragraph.text
        ):
            set_first_line_chars(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not HOME_SCREENSHOT.exists():
    raise FileNotFoundError(HOME_SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.15 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.15 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.16 Demo   |   更新日期：2026年7月24日   |   适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页使用红塘村在线3D高斯实景，并展示小花园、茶场、茶厂、"
    "村里用水和光伏设施5个可点击示例点。点击图钉只打开事项详情；点击“定位到此处”才会让镜头飞近，"
    "关闭按钮或按Esc可取消选择。所有三维图钉均为演示位置，待实地核实，不能作为真实设施、风险或产权判断依据。"
    "原无人机影像首页完整保留为“村庄总览”；3D首页需要保持联网。"
)

caption = find_paragraph_prefix(document, "图 1")
set_home_image(image_paragraph_before(document, caption))
caption.text = "图 1  3D高斯实景首页与五类示例点"

operation = find_paragraph_prefix(document, "3D首页操作：")
operation.text = (
    "3D首页操作：左键拖动旋转；右键拖动只沿村庄所在地表的水平面平移；滚轮每次按固定比例缩放。"
    "点击小花园、茶场、茶厂、村里用水或光伏设施图钉，会在桌面右侧或手机底部打开与“村里一张图”一致的事项详情，"
    "但不会自动移动镜头；只有点击“定位到此处”才会平滑靠近。关闭按钮或Esc会取消高亮。点击“操作设置”后，"
    "右键平移和滚轮缩放灵敏度可分别调节、分别保存；还可使用“回到模型”和“全屏查看”。"
)

technical_anchor = find_paragraph_prefix(document, "平台从Cesium默认缩放输入中移除")
insert_paragraph_after(
    technical_anchor,
    "平台在Cesium中固定创建map-13至map-17五个三维锚点；Cesium负责图钉绘制、拾取、高亮和显式镜头定位，"
    "React父页面继续从二维mapFeatures读取标题、状态和说明。普通选择不调用相机，定位动画为1.2秒。"
    "桌面详情在右侧打开，390px手机端详情在底部打开且最大高度不超过可视区的55%。",
    "List Bullet",
)

for row in document.tables[1].rows:
    if row.cells[0].text.strip() == "首页":
        row.cells[2].text = (
            "全屏红塘村3D高斯实景；5个分类示例图钉可打开事项详情和显式定位；"
            "左键旋转、右键水平平移、滚轮定比缩放，并可分别调节两项灵敏度"
        )
        break

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载红塘村高斯模型；显示5个可点击分类锚点；二维事项数据复用、显式定位、"
            "Esc清除和移动端底部详情；服务端保存长期令牌"
        )
        break

for row in document.tables[3].rows:
    if row.cells[0].text.strip() == "3D首页一直加载或显示失败":
        row.cells[0].text = "3D首页一直加载、图钉不显示或显示失败"
        row.cells[1].text = "未联网、Cesium访问配置不可用，或模型和点位尚未完成加载"
        row.cells[2].text = "保持联网并重新加载；仍失败时先进入“村庄总览”，不要把演示图钉当作真实位置"
        break

version_table = document.tables[4]
version_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    version_row.cells,
    [
        "V1.16 Demo",
        "2026-07-24",
        "3D首页新增小花园、茶场、茶厂、村里用水和光伏设施5个可点击示例点；复用二维事项详情，新增显式定位、Esc清除和手机底部详情。",
    ],
):
    cell.text = value

normalize_document(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 18]
assert "V1.16 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "演示位置，待实地核实" in all_text
assert "定位到此处" in all_text
assert any(row.cells[0].text.strip() == "V1.16 Demo" for row in check.tables[4].rows)
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.16: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
