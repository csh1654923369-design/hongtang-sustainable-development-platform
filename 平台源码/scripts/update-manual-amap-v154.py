from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"

REPLACEMENTS = {
    "版本：V1.53 ｜ 更新日期：2026年8月9日": "版本：V1.54 ｜ 更新日期：2026年8月9日",
    "首页默认进入红塘村2D地图，便于直接浏览并降低设备性能压力。左上角平台悬浮框右侧可在“2D地图”和“3D实景”之间切换；进入3D后直接定位红塘村，不播放从全球飞入村庄的开场动画。两种模式读取同一套地点、点线面专题和调研记录数据，并统一采用“地图作为全屏底层—筛选要素—点击对象—视图移动并缩放—点位旁气泡展示详情”的交互顺序。切换时只替换底层地图，“筛选要素”组件始终保留在首页，不会重新加载，已经勾选或取消的图层会继续生效；同时只运行当前地图，避免2D与3D同时占用显卡。": "首页默认进入红塘村2D地图，便于直接浏览并降低设备性能压力。2D以高德在线地图作为云端底层，可叠加红塘无人机影像或手绘图；3D以Cesium地形和高斯模型作为底层。左上角平台悬浮框右侧可在“2D地图”和“3D实景”之间切换；进入3D后直接定位红塘村，不播放从全球飞入村庄的开场动画。两种模式读取同一套地点、点线面专题和调研记录数据，并统一采用“地图作为全屏底层—筛选要素—点击对象—视图移动并缩放—点位旁气泡展示详情”的交互顺序。切换时只替换底层地图，“筛选要素”组件始终保留在首页，已经勾选或取消的图层会继续生效；同时只运行当前地图，避免2D与3D同时占用显卡。",
    "2D模式与3D实景读取同一套数据，并把手绘图或无人机影像作为完整地图底层。2D地图在任何缩放级别都可以按住鼠标左键向任意方向拖动；滚轮以鼠标所在位置为中心缩放，双击地图或点击“复位视图”可返回初始位置。平台会自动区分拖动与点击，移动地图时不会误开地点详情。点击小花园、茶厂等地点后，二维地图会平滑移动并放大到对应坐标，在点位旁弹出与3D相同结构的气泡详情；点击水源、供排水线路或片区也会聚焦到对象，同时突出与其关联的上下游和服务片区。图钉会保持清晰的屏幕尺寸，不随底图缩放而变得过大。没有资料的字段不会由系统虚构补齐。": "2D模式与3D实景读取同一套地点和专题数据。2D地图使用高德在线道路、地名和水系作为云端底层，右上角可选择“高德底图”“无人机影像”或“手绘图”；后两项会以半透明地理配准图层叠加在高德地图上，因此村域外仍能看到周边道路并继续移动。按住鼠标左键可向任意方向拖动，滚轮缩放，点击右下角“回到红塘”恢复村庄中心视图。点击小花园、茶厂等地点后，地图会移动并放大到对应坐标，在点位旁弹出与3D相同结构的气泡详情；点击水源、供排水线路或片区也会聚焦对象并突出关联关系。原始地点数据保留WGS84坐标，2D显示时统一转换为高德使用的GCJ-02，避免底图与村庄资料错位；3D仍读取原始WGS84坐标。高德服务不可用时，页面自动回退到本地影像地图，不影响筛选和详情查看。",
}


def find_paragraph(document: Document, exact: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == exact]
    if len(matches) != 1:
        raise ValueError(f"应找到1处段落，实际找到{len(matches)}处：{exact}")
    return matches[0]


def set_two_character_indent(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), "200")


def format_run(run) -> None:
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(0, 0, 0)


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    format_run(paragraph.add_run(text))
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    format_run(paragraph.add_run(text))


document = Document(MANUAL)
for original, updated in REPLACEMENTS.items():
    replace_paragraph(find_paragraph(document, original), updated)

data_table = document.tables[2]
map_row = next(row for row in data_table.rows if row.cells[0].text.strip() == "地图底图")
set_cell_text(map_row.cells[1], "高德在线地图；平台源码/public/data")
set_cell_text(map_row.cells[2], "高德道路、地名和水系；手绘图、无人机影像及地理配准信息")

module_table = document.tables[3]
map_module_row = next(row for row in module_table.rows if row.cells[0].text.strip().startswith("MapExplorer.tsx"))
set_cell_text(map_module_row.cells[0], "MapExplorer.tsx / AmapVillageMap.tsx / MapFilterPanel.tsx")
set_cell_text(map_module_row.cells[1], "高德云端2D底图、无人机与手绘覆盖层、共享多选筛选、点线面点击聚焦、气泡详情及移动端交互")
fallback_row = next(row for row in module_table.rows if row.cells[0].text.strip().startswith("VillageMap.tsx"))
set_cell_text(fallback_row.cells[0], "VillageMap.tsx / MapMarker.tsx / useMapZoom.ts / amap.ts")
set_cell_text(fallback_row.cells[1], "WGS84转GCJ-02、统一图钉和圆点、云端地图加载、本地影像回退及二维气泡锚点跟踪")

faq_table = document.tables[4]
move_faq = next(row for row in faq_table.rows if row.cells[0].text.strip() == "2D地图怎样移动、缩放和复位")
set_cell_text(move_faq.cells[1], "把鼠标放在地图画面上，按住左键即可向任意方向拖动，滚动滚轮进行缩放；点击右下角“回到红塘”恢复村庄中心视图。")

new_faq_xml = deepcopy(faq_table.rows[1]._tr)
faq_table._tbl.append(new_faq_xml)
new_faq = faq_table.rows[-1]
set_cell_text(new_faq.cells[0], "为什么2D地图能看到村外道路，村内又有无人机或手绘图")
set_cell_text(new_faq.cells[1], "2D以高德在线地图作为连续的云端底层，无人机影像和手绘图按红塘村地理范围叠加在上方。右上角可切换三种显示方式；若高德暂时不可用，页面会自动显示本地影像地图。")

version_table = document.tables[5]
template_row = deepcopy(version_table.rows[1]._tr)
version_table.rows[0]._tr.addnext(template_row)
new_version = version_table.rows[1]
set_cell_text(new_version.cells[0], "V1.54")
set_cell_text(new_version.cells[1], "2026年8月9日")
set_cell_text(new_version.cells[2], "首页2D接入高德在线地图，并将无人机影像和手绘图改为可切换的地理配准覆盖层；点位、水系统线面、筛选和气泡详情继续共用；新增WGS84转GCJ-02、回到红塘和高德不可用时的本地地图回退。")

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 用户已明确不需要LibreOffice渲染，因此只执行结构、内容、样式与压缩包完整性校验。
reopened = Document(MANUAL)
paragraph_texts = [paragraph.text.strip() for paragraph in reopened.paragraphs]
assert "版本：V1.54 ｜ 更新日期：2026年8月9日" in paragraph_texts
assert any("使用高德在线道路、地名和水系作为云端底层" in text for text in paragraph_texts)
assert reopened.tables[2].rows[2].cells[1].text.strip().startswith("高德在线地图")
assert any("为什么2D地图能看到村外道路" in row.cells[0].text for row in reopened.tables[4].rows)
assert reopened.tables[5].rows[1].cells[0].text.strip() == "V1.54"
assert len(reopened.tables[4].rows) == 13
assert len(reopened.tables[5].rows) == 17
assert reopened.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) in {"宋体", "SimSun"}
assert reopened.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) in {"黑体", "SimHei"}
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(f"updated={MANUAL}")
print(f"version={reopened.tables[5].rows[1].cells[0].text.strip()}")
print(f"faq_rows={len(reopened.tables[4].rows)}")
print("zip=ok")
