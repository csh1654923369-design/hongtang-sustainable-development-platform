from copy import deepcopy
from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v19.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.7.backup.docx"
BACKUP_V18 = QA_ROOT / "红塘村可持续发展平台使用手册.V1.8.backup.docx"
HOME_SCREENSHOT = QA_ROOT / "home-1440.png"
TOPIC_SCREENSHOT = QA_ROOT / "topic-tea-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption_text):
    caption = find_paragraph(document, caption_text)
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError(f"Image paragraph not found before caption: {caption_text}")


def set_image(paragraph, image_path, width, name, description):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("name", name)
    shape._inline.docPr.set("descr", description)


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def set_cell_text_like(source_cell, target_cell, text):
    copy_cell_format(source_cell, target_cell)
    target_cell.text = text
    target_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if source_cell.paragraphs and target_cell.paragraphs:
        target_cell.paragraphs[0].style = source_cell.paragraphs[0].style
        target_cell.paragraphs[0].paragraph_format.alignment = source_cell.paragraphs[0].paragraph_format.alignment


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        set_cell_text_like(template_row.cells[index], cell, "")
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        east_asia = "SimHei" if is_heading or index == 0 else "SimSun"
        for run in paragraph.runs:
            set_run_font(run, east_asia)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


def replace_screenshots(document, home_caption, topic_caption):
    set_image(
        image_paragraph_before(document, home_caption),
        HOME_SCREENSHOT,
        6.3,
        "图 1 平台首页",
        "首页首屏展示红塘村无人机影像和顶部事项导航",
    )
    set_image(
        image_paragraph_before(document, topic_caption),
        TOPIC_SCREENSHOT,
        6.3,
        "图 2 茶厂独立页面",
        "从顶部事项导航进入茶厂独立页面后的页面骨架",
    )


for screenshot in (HOME_SCREENSHOT, TOPIC_SCREENSHOT):
    if not screenshot.exists():
        raise FileNotFoundError(screenshot)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
version_text = document.paragraphs[5].text

if "V1.9 Demo" in version_text:
    replace_screenshots(document, "图 1  平台首页", "图 2  顶部事项导航与茶厂独立页面")
    normalize_fonts(document)
    document.save(OUTPUT)
    check = Document(OUTPUT)
    assert len(check.inline_shapes) == 13
    os.replace(OUTPUT, MANUAL)
    print(f"Refreshed V1.9 interface screenshots: {MANUAL}")
    raise SystemExit(0)

if "V1.8 Demo" in version_text:
    if not BACKUP_V18.exists():
        shutil.copy2(MANUAL, BACKUP_V18)
    document.paragraphs[5].text = "版本：V1.9 Demo   |   更新日期：2026年7月22日   |   适用地址：http://localhost:3000"
    replace_screenshots(document, "图 1  平台首页", "图 2  顶部事项导航与茶厂独立页面")
    find_paragraph(
        document,
        "村庄总览：首页首屏直接展示红塘村无人机影像，并提供六个事项入口和三个常用操作。",
    ).text = "桌面端顶栏：平台名称、六类事项 TAB 和常用操作位于同一行；点击“村庄总览”返回首页。"
    find_paragraph(
        document,
        "在手机或窄屏浏览器中，顶部事项栏可以左右滑动；底部保留“首页、选事项、一张图、我的”四个快捷入口。其他页面可通过右上角菜单访问。网站已针对 390px 宽度进行无横向溢出测试。",
    ).text = "在手机或窄屏浏览器中，横排事项 TAB 会收进右上角菜单；底部保留“首页、选事项、一张图、我的”四个快捷入口。网站已针对 390px 宽度进行无横向溢出测试。"
    version_table = document.tables[4]
    new_row = add_row_like(version_table, version_table.rows[-1])
    for cell, value in zip(
        new_row.cells,
        [
            "V1.9 Demo",
            "2026-07-22",
            "桌面端事项 TAB 改回与平台名称和操作按钮同一条顶栏；窄屏时事项入口收进右上角菜单，不再占用第二行。",
        ],
    ):
        cell.text = value
    normalize_fonts(document)
    document.save(OUTPUT)
    check = Document(OUTPUT)
    assert len(check.inline_shapes) == 13
    assert "V1.9 Demo" in check.paragraphs[5].text
    assert len(check.tables[4].rows) == 11
    os.replace(OUTPUT, MANUAL)
    print(f"Updated manual to V1.9: {MANUAL}")
    raise SystemExit(0)

if "V1.7 Demo" not in version_text:
    raise ValueError(f"Expected V1.7 manual, got: {version_text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = "版本：V1.8 Demo   |   更新日期：2026年7月22日   |   适用地址：http://localhost:3000"
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页和“村里一张图”使用红塘村无人机正射影像；"
    "问题、项目、微行动、资源供需、进展指标和业务点位仍含演示数据。建筑调研底图来自既有调研资料并已去除个人字段，"
    "仅用于规划讨论，不作为权属、测绘或审批依据。"
)

replace_screenshots(document, "图 1  平台首页", "图 2  “村里的事”六类长期记录模块")

navigation_replacements = {
    "首页：从“记录一件事、看看最近变化、参加村里活动”三个大入口开始。": "村庄总览：首页首屏直接展示红塘村无人机影像，并提供六个事项入口和三个常用操作。",
    "村里的事：查看小花园、茶园与茶厂、村里用水、光伏设施、塌方与安全、红塘村历史。": "小花园、茶厂、村里用水：分别进入自己的独立页面，不再先打开同一个通用模块。",
    "村里一张图：在地图上找地点、问题、项目、互助资源和正在开展的事情。": "光伏设施、安全隐患、村庄记忆：分别进入设施管理、风险跟进和村庄档案的独立页面骨架。",
    "我的：查看与当前演示身份有关的记录、活动、关注和通知。": "村里一张图：从首页影像、右上角地图按钮或手机底栏进入，查看完整影像、地点和业务记录。",
    "更多功能：认识红塘、可持续目标、项目与行动、发展进展、研究协作和数字沙盘仍可从首页、移动菜单、搜索或页脚进入。": "更多工具：活动、项目、认识红塘、可持续目标、研究资料和数字沙盘从首页“其他工具”、移动菜单、搜索或页脚进入。",
    "在手机或窄屏浏览器中，页面底部只保留“首页、村里的事、一张图、我的”四个快捷入口；其他页面可通过右上角菜单访问。网站已针对 390px 宽度进行无横向溢出测试。": "在手机或窄屏浏览器中，顶部事项栏可以左右滑动；底部保留“首页、选事项、一张图、我的”四个快捷入口。其他页面可通过右上角菜单访问。网站已针对 390px 宽度进行无横向溢出测试。",
    "2.4 使用“村里的事”": "2.4 从顶部选择具体事项",
    "在顶部导航或手机底栏点击“村里的事”。": "在顶部事项栏直接点击“小花园、茶厂、村里用水、光伏设施、安全隐患、村庄记忆”中的任意一项。",
    "从六个大按钮中选择一类：我家小花园、茶园与茶厂、村里用水、光伏设施、塌方与安全或红塘村历史。": "系统会进入该事项自己的网址和页面文件；切换顶部标签即可进入另一项，不需要返回通用汇总页。",
    "页面下方会显示这类事项最近在记录什么、以后要连续记录什么，以及相关地图或项目入口。": "当前独立页面只显示功能分区骨架和“功能待讨论”提示，具体表单、数据和操作将在确认真实需求后逐项开发。",
    "点击“我要记录一件事”可进入上报流程；只想浏览时，可继续查看“村里最近有什么新情况”。": "需要地图、问题记录或近期变化时，可使用首页三个常用操作；原“村里的事”汇总页仅为兼容旧链接而保留。",
    "图 2  “村里的事”六类长期记录模块": "图 2  顶部事项导航与茶厂独立页面",
    "使用提示\n页面目前显示演示记录。真实小花园、茶园、用水、光伏、安全和历史资料，需要在村民、村委和调研团队确认后逐步补充。": "使用提示\n六个事项目前已经分开，但页面中的功能名称只是讨论起点，不代表最终业务流程。每一项都需要与村民、村委、老师和调研团队确认后单独设计。",
}
for old, new in navigation_replacements.items():
    find_paragraph(document, old).text = new

role_rows = [
    ["角色", "可以做什么", "主要入口"],
    ["游客", "浏览首页、六个事项、村里一张图、项目和公开进展", "村庄总览、顶部事项栏、村里一张图"],
    ["村民", "查看具体事项、上报问题、发起或加入微行动、共享资源、报名、建议与评价", "顶部事项栏、村里一张图、参与、我的"],
    ["学生/规划协作者", "参与事项功能共创、发起或加入微行动、共享资源、提交调研成果与建议", "顶部事项栏、村里一张图、项目、研究协作、我的"],
    ["管理员/村委", "轻量核对微行动、处理问题、审核资料、维护项目和查看日志", "管理后台"],
]
for row, values in zip(document.tables[0].rows, role_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

page_rows = [
    ["页面", "地址", "核心用途"],
    ["首页", "/", "无人机影像主视觉、六个独立事项入口和三个常用操作"],
    ["小花园", "/garden", "小花园独立页面骨架，具体功能待讨论"],
    ["茶厂", "/tea-factory", "茶园、收茶、加工三个阶段的独立页面骨架"],
    ["村里用水", "/water", "水源设施、用水问题和维修反馈的独立页面骨架"],
    ["光伏设施", "/solar", "设施位置、运行巡查和信息公开的独立页面骨架"],
    ["安全隐患", "/safety", "隐患发现、跟进处理和现场复查的独立页面骨架"],
    ["村庄记忆", "/village-history", "老照片、村民讲述、古道与老屋的独立档案骨架"],
    ["村里的事（旧汇总）", "/village-life", "兼容旧链接，不再作为顶部主入口"],
    ["认识红塘", "/village", "村庄资料框架与故事"],
    ["可持续目标", "/goals", "目标、指标、项目联动"],
    ["村里一张图", "/map", "示意图/无人机影像切换、业务点位、互助资源和建筑调研底图"],
    ["问题上报", "/report", "五步记录村庄问题"],
    ["发起微行动", "/actions/new", "五步形成可核对、可招募的小行动"],
    ["项目与行动", "/projects", "社区微行动、正式项目、开放任务与资源需求"],
    ["发展进展", "/progress", "社区行动能力、趋势、来源与完整度"],
    ["公众参与", "/participate", "微行动入口、建议、活动、问卷和共创"],
    ["个人中心", "/profile", "我的行动、个人记录与通知"],
    ["研究协作", "/research", "调研成果提交"],
    ["管理后台", "/admin", "问题处理、微行动核对、审核与维护"],
    ["数字沙盘", "/digital-twin", "三维场景概念演示"],
]
page_table = document.tables[1]
template_row = page_table.rows[1]
while len(page_table.rows) < len(page_rows):
    add_row_like(page_table, template_row)
for row, values in zip(page_table.rows, page_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

version_table = document.tables[4]
new_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    new_row.cells,
    [
        "V1.8 Demo",
        "2026-07-22",
        "首页首屏改为红塘村无人机影像；顶部导航改为六类具体事项；小花园、茶厂、村里用水、光伏设施、安全隐患和村庄记忆分别建立独立路由与页面骨架。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)
check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert "V1.8 Demo" in check.paragraphs[5].text
assert any("顶部事项导航与茶厂独立页面" in paragraph.text for paragraph in check.paragraphs)
assert len(check.tables[1].rows) == 21
os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.8: {MANUAL}")
