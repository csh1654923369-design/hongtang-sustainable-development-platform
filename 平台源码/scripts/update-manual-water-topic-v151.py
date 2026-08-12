from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def find_paragraph(document: Document, exact: str) -> Paragraph:
    normalized = " ".join(exact.split())
    for paragraph in document.paragraphs:
        if " ".join(paragraph.text.split()) == normalized:
            return paragraph
    raise ValueError(f"未找到段落：{exact}")


def set_two_character_indent(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    for attr in ("firstLine", "hanging"):
        ind.attrib.pop(qn(f"w:{attr}"), None)
    ind.set(qn("w:firstLineChars"), "200")


def format_run(run, *, heading: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体" if heading else "宋体")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.highlight_color = WD_COLOR_INDEX.AUTO


def replace_paragraph(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    format_run(run, heading=paragraph.style.name.startswith(("Heading", "标题")))
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def insert_before(reference: Paragraph, text: str, style: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    reference._p.addprevious(new_element)
    paragraph = Paragraph(new_element, reference._parent)
    paragraph.style = style
    run = paragraph.add_run(text)
    format_run(run, heading=style.startswith(("Heading", "标题")))
    if style == "Normal":
        set_two_character_indent(paragraph)
    return paragraph


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    format_run(run)


document = Document(MANUAL)

replacements = {
    "版本：V1.50 ｜ 更新日期：2026年8月9日": "版本：V1.51 ｜ 更新日期：2026年8月9日",
    "本手册区分“已有资料”和“示例数据”。已有地点名称、坐标、分类与照片来自当前资料；尚未实测的供水节点、线路、片区和专题字段统一标注为“示例矢量”或“示例数据”，只用于演示信息组织与交互方式，不得作为红塘村现状结论。其他尚未录入或尚未核实的业务页面仍统一提示“暂无已核实内容”。": "本手册区分“已有资料”和“待核实结构示意”。已有地点名称、坐标、分类与照片来自当前资料；尚未实测的水源、供排水线路、片区边界、使用对象和维护关系统一标注为“待实地核实”，只用于验证信息组织与交互方式，不得作为红塘村现状结论。其他尚未录入或尚未核实的内容仍使用明确提示。",
    "当前基础地点数据包含205个有效坐标点；首页另载入7个示例供水节点、5条示例供水线路和3个示例供水片区，共形成220个可筛选地图要素。205个基础地点中的名称、分类、坐标和照片来自现有资料；新增供水点线面只用于展示未来数据结构。": "当前基础地点数据包含205个有效坐标点；首页另载入7个待核实水系统节点、8条待核实线路和3个待核实供水片区，共形成223个地图要素。205个基础地点中的名称、分类、坐标和照片来自现有资料；水专题点线面用于展示未来的数据关系和调查方式。",
    "三维页面使用项目内置的CesiumJS 1.143运行文件，不再依赖cesium.com程序CDN。场景优先从Cesium ion加载红塘村高斯3D Tiles、世界地形和航空影像，并在村庄核心区域叠加本地无人机正射影像；共享的地点、供水线路和供水片区则按同一套地理坐标叠加在三维场景中。若在线地形或模型连接超时，页面会自动改用基础球面、本地高斯3D Tiles和本地无人机影像。在线完整环境仍需要联网和可用的Cesium ion访问令牌。": "三维页面使用项目内置的CesiumJS 1.143运行文件，不再依赖cesium.com程序CDN。场景优先从Cesium ion加载红塘村高斯3D Tiles、世界地形和航空影像，并在村庄核心区域叠加本地无人机正射影像；共享的地点、供排水线路和供水片区则按同一套地理坐标叠加在三维场景中。若在线地形或模型连接超时，页面会自动改用基础球面、本地高斯3D Tiles和本地无人机影像。在线完整环境仍需要联网和可用的Cesium ion访问令牌。",
    "2D模式与3D实景读取同一套数据，并把手绘图或无人机影像作为完整地图底层。点击小花园、茶厂等地点后，二维地图会平滑移动并放大到对应坐标，在点位旁弹出与3D相同结构的气泡详情；点击供水线路或片区也会聚焦到对象并显示关系说明。图钉会保持清晰的屏幕尺寸，不随底图缩放而变得过大。没有资料的字段不会由系统虚构补齐。": "2D模式与3D实景读取同一套数据，并把手绘图或无人机影像作为完整地图底层。点击小花园、茶厂等地点后，二维地图会平滑移动并放大到对应坐标，在点位旁弹出与3D相同结构的气泡详情；点击水源、供排水线路或片区也会聚焦到对象，同时突出与其关联的上下游和服务片区。图钉会保持清晰的屏幕尺寸，不随底图缩放而变得过大。没有资料的字段不会由系统虚构补齐。",
    "关系线路：当前用5条示例线路表达水源、调蓄节点和供水片区之间的连接。点击后可查看起点、终点、服务范围、估算落差和说明。": "关系线路：当前用3条待核实供水线路表达“水源—供水点—供水片区”，用5条待核实排水线路表达“房前屋后—支沟—主沟—水塘—末端去向”。点击后可查看起点、终点、服务范围、流向、估算落差和维护问题。",
    "管理片区：当前用3个示例片区表达北部、中心和南部供水范围。点击片区后，详情会直接显示“该片区供水来源”及对应供水节点。": "管理片区：当前用3个待核实片区表达北部、中心和南部供水范围。点击片区后，详情会串联显示供水来源、经过线路、集中供水点、使用对象和维护事项。",
    "供水节点、线路和片区依据现有地形资料搭建为演示结构，尚未经过管线测绘、村民访谈或管理部门核实。界面统一显示“示例矢量”，后续取得水源点、管径、阀门、流向、服务户数和片区边界等资料后，应在共享数据文件中替换，而不是分别修改2D和3D。": "水系统节点、线路和片区依据现有地形资料搭建为结构示意，尚未经过管线测绘、村民访谈或管理部门核实。界面统一显示“待实地核实”，后续取得水源点、管径、阀门、流向、服务户数、片区边界、共同建设过程和维护责任等资料后，应在共享数据文件中替换，而不是分别修改2D和3D。",
    "小花园、茶厂、村里用水、公共服务、生态资源和村景记录不再分别设置顶部标签或独立页面。进入首页后，在左上角多选列表中勾选一个或多个分类；2D和3D会同步显示同一批地点、线路和片区。": "小花园、茶厂、公共服务、生态资源和村景记录不再分别设置顶部标签或独立页面。进入首页后，可在左上角多选列表中勾选一个或多个普通分类；“村里用水专题”从地图左下角单独进入，以便连续阅读水源、线路、片区和上下游关系。",
    "点击地点后，2D与3D都会先把视图移动并缩放到对象，再在对象旁的气泡中查看名称、照片、坐标、简介和已关联的专题示例字段，不再出现“进入对应板块”按钮。点击供水线路或片区同样会聚焦并以气泡展示连接关系、服务范围和供水来源；点击地图空白处、气泡关闭按钮或其他对象即可取消或切换当前选择。": "点击地点后，2D与3D都会先把视图移动并缩放到对象，再在对象旁的气泡中查看名称、照片、坐标、简介和已关联字段，不再出现“进入对应板块”按钮。点击水系统节点、线路或片区后，地图会同时突出相关上下游和服务片区；气泡中的关联名称可以继续点击并自动移动到新对象。点击地图空白处、气泡关闭按钮或其他对象即可取消或切换当前选择。",
    "4.3 已移除的旧板块": "4.4 已移除的旧板块",
    "当前交付版本只把首页作为面向用户的正式页面。首页内已经包含3D实景、2D地图、地点分类、照片详情、专题示例字段和供水点线面关系，不再另设地图页、记录页或事项页。": "当前交付版本只把首页作为面向用户的正式页面。首页内已经包含3D实景、2D地图、地点分类、照片详情、专题字段，以及饮水供给和排水去向的点线面关系，不再另设地图页、记录页或事项页。",
}

for original, updated in replacements.items():
    replace_paragraph(find_paragraph(document, original), updated)

old_section = find_paragraph(document, "4.4 已移除的旧板块")
insert_before(old_section, "4.3 村里用水专题", "Heading 2")
insert_before(old_section, "地图左下角显示“村里用水专题”入口。点击后，普通地点筛选暂时收起，2D或3D地图只保留与当前水专题视角有关的对象；点击右上角关闭按钮即可退出专题并恢复普通地图。2D与3D之间切换时，当前水专题视角会继续保留。", "Normal")
insert_before(old_section, "水系统全貌：同时查看饮水供给和排水去向，理解水源、住户、房前屋后、沟渠、水塘、处理节点和下游空间如何连接。", "List Bullet")
insert_before(old_section, "饮水从哪来：只显示5个供水节点、3条供水线路和3个供水片区，回答“每个片区的水从哪里来、经过哪条线路、供应给谁”。", "List Bullet")
insert_before(old_section, "排水到哪里：显示2个待核实排水节点、5条待核实排水线路，并保留2个已有污水设施点位，回答“房前屋后的水在哪里汇集、怎样流动、最终到哪里”。", "List Bullet")
insert_before(old_section, "选择任一水源、供水点、沟渠、水塘、线路或片区后，地图先聚焦该对象，再突出同一关系链中的其他对象。", "List Number")
insert_before(old_section, "在气泡的“顺着水继续看”区域点击关联名称，可以沿着“从哪里来—经过哪里—服务谁或流向哪里”的顺序继续查看。", "List Number")
insert_before(old_section, "继续阅读“这个节点有什么作用”“这个片区怎样用水”“共同建设与维护”和“还需要向村民了解”，把空间设施与使用者、维护者和待调查问题联系起来。", "List Number")
insert_before(old_section, "当前水专题是待核实的关系原型，不是红塘村现状管网。正式调查时，应邀请不同片区村民沿图核对水源、季节变化、供水户数、雨污去向、共同建设过程和维护责任。", "Normal")

# 数据数量与说明表
set_cell_text(document.tables[0].rows[3].cells[1], "20个地图要素")
set_cell_text(document.tables[0].rows[3].cells[2], "2个已有设施与7个待核实节点使用水滴图钉；另有8条待核实线路和3个待核实片区")
set_cell_text(document.tables[1].rows[3].cells[1], "20个地图要素")
set_cell_text(document.tables[1].rows[3].cells[2], "从地图左下角进入专题，查看2个已有设施、7个待核实节点、8条待核实线路和3个待核实片区")
set_cell_text(document.tables[2].rows[6].cells[0], "水专题关系数据")
set_cell_text(document.tables[2].rows[6].cells[2], "首页2D、3D共用的7个节点、8条线路和3个片区；用于验证饮水、排水、上下游、服务对象和维护关系，尚未核实")
set_cell_text(document.tables[3].rows[7].cells[0], "spatialData.ts / WaterTopicNavigator.tsx / WaterSpatialDetail.tsx")
set_cell_text(document.tables[3].rows[7].cells[1], "定义2D与3D共用的水系统关系、三种专题视角、上下游突出、关联跳转，以及使用、维护和待访谈问题详情")
set_cell_text(document.tables[4].rows[8].cells[0], "首页中的水系统节点、线路和片区是真实管网吗")
set_cell_text(document.tables[4].rows[8].cells[1], "不是。它们统一标注为“待实地核实”，用于演示怎样沿着水源、线路、片区和排水去向理解村庄。取得测绘、访谈或管理资料后再替换共享JSON。")
new_faq = document.tables[4].add_row()
set_cell_text(new_faq.cells[0], "怎样查看某个片区的水从哪里来")
set_cell_text(new_faq.cells[1], "点击地图左下角“村里用水专题”，选择“饮水从哪来”，再点击供水片区。气泡会列出供水来源、经过线路和集中供水点；点击其中任一名称即可继续追踪。")

# 常用检查命令
check_heading = find_paragraph(document, "9.2 每次功能变更后")
insert_before(check_heading, "npm run test:water-topic：验证水专题三种视角、2D/3D模式保持、关系突出、关联跳转和气泡详情。", "List Bullet")

# 在变更记录表顶部加入V1.51
version_table = document.tables[5]
template_row = deepcopy(version_table.rows[1]._tr)
version_table.rows[0]._tr.addnext(template_row)
new_row = version_table.rows[1]
set_cell_text(new_row.cells[0], "V1.51")
set_cell_text(new_row.cells[1], "2026年8月9日")
set_cell_text(new_row.cells[2], "新增统一的“村里用水专题”入口及“水系统全貌、饮水从哪来、排水到哪里”三种视角；补充3条供水线路，使水专题形成7个节点、8条线路和3个片区；2D与3D同步筛选并突出上下游和服务关系，气泡支持沿关联对象继续跳转，并增加作用、使用对象、共同维护和待访谈问题。所有内容明确标注为待实地核实。")

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 不进行LibreOffice渲染；只做结构、内容和压缩包完整性校验。
reopened = Document(MANUAL)
paragraph_texts = [paragraph.text.strip() for paragraph in reopened.paragraphs]
assert "4.3 村里用水专题" in paragraph_texts
assert "4.4 已移除的旧板块" in paragraph_texts
assert paragraph_texts[1].startswith("版本：V1.51")
assert any("npm run test:water-topic" in text for text in paragraph_texts)
assert reopened.tables[5].rows[1].cells[0].text.strip() == "V1.51"
assert reopened.tables[0].rows[3].cells[1].text.strip() == "20个地图要素"
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(f"updated={MANUAL}")
print(f"paragraphs={len(reopened.paragraphs)} tables={len(reopened.tables)} version={reopened.tables[5].rows[1].cells[0].text.strip()}")
