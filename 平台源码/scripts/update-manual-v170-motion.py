from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.70 ｜ 更新日期：2026年8月10日",
)

overview = next(p for p in document.paragraphs if p.text.startswith("首页默认进入红塘村2D地图"))
overview_text = overview.text.replace(
    "选中2D时，其下方以独立卡片纵向显示“航拍”“手绘图”“卫星遥感”“底图”四种按钮；",
    "选中2D时，同一卡片下排横向显示“航拍”“手绘”“卫星”“底图”四个按钮；切换到3D时，下排按钮通过高度与透明度过渡自然收起，返回2D时自然展开；",
)
replace_paragraph_text(overview, overview_text)

topic_intro = next(p for p in document.paragraphs if p.text.startswith("桌面端和手机端都在地图左上角"))
topic_intro_text = topic_intro.text.replace(
    "点击右侧带倒三角的“展开”后，卡片以简短的缓入缓出动画自然展开，内容同步淡入；打开后按钮变为“收起”，收起时执行反向过渡。动画不使用位移或回弹效果。",
    "点击右侧带倒三角的“展开”后，卡片通过外层高度和透明度自然展开；内部文字与列表始终保持固定位置，不参与位移、缩放或回弹动画。打开后按钮变为“收起”，点击时执行相同节奏的反向过渡。",
)
replace_paragraph_text(topic_intro, topic_intro_text)

two_d = next(p for p in document.paragraphs if p.text.startswith("2D模式与3D实景读取同一套地点和专题数据"))
two_d_text = two_d.text.replace(
    "右上角保留独立的2D/3D切换卡片；选中2D后，其下方另设纵向底图列表，可选择“航拍”“手绘图”“卫星遥感”或“底图”。“卫星遥感”使用高德官方卫星与路网图层；“无人机影像”和“手绘图”会以半透明地理配准图层叠加在高德地图上，",
    "右上角使用一张统一功能卡片：上排切换2D/3D，选中2D后，下排横向显示“航拍”“手绘”“卫星”“底图”；进入3D时下排自然收起，返回2D时自然展开，当前底图选择保持不变。“卫星”使用高德官方卫星与路网图层；“航拍”和“手绘”会以半透明地理配准图层叠加在高德地图上，",
)
replace_paragraph_text(two_d, two_d_text)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.70"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.70",
    "2026年8月10日",
    "专题文字改为固定位置，仅由外层高度与透明度控制展开收起；2D底图按钮在同一卡片内横向排列，并在2D/3D切换时自然展开或收起。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.70") for p in reloaded.paragraphs)
assert "内部文字与列表始终保持固定位置" in paragraph_text
assert "同一卡片下排横向显示“航拍”“手绘”“卫星”“底图”" in paragraph_text
assert "进入3D时下排自然收起，返回2D时自然展开" in paragraph_text
assert "独立卡片纵向显示" not in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.70"
print(MANUAL)
