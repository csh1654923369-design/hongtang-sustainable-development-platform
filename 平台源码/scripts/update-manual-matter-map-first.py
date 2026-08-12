from copy import deepcopy
from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v111.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.10.backup.docx"
TOPIC_SCREENSHOT = QA_ROOT / "topic-tea-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption_text):
    caption = find_paragraph(document, caption_text)
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError(f"Image paragraph not found before caption: {caption_text}")


def set_image(paragraph, image_path, width, name, description):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("name", name)
    shape._inline.docPr.set("descr", description)


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        copy_cell_format(template_row.cells[index], cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if template_row.cells[index].paragraphs and cell.paragraphs:
            cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
            cell.paragraphs[0].paragraph_format.alignment = template_row.cells[index].paragraphs[0].paragraph_format.alignment
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        east_asia = "SimHei" if is_heading or index == 0 else "SimSun"
        for run in paragraph.runs:
            set_run_font(run, east_asia)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not TOPIC_SCREENSHOT.exists():
    raise FileNotFoundError(TOPIC_SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
version_text = document.paragraphs[5].text

if "V1.11 Demo" in version_text:
    set_image(
        image_paragraph_before(document, "图 2  茶厂页面的事项地图首屏"),
        TOPIC_SCREENSHOT,
        6.3,
        "图 2 茶厂页面的事项地图首屏",
        "茶厂页面首先展示村庄无人机地图并只显示茶场与茶厂点位",
    )
    normalize_fonts(document)
    document.save(OUTPUT)
    os.replace(OUTPUT, MANUAL)
    print(f"Refreshed V1.11 topic map screenshot: {MANUAL}")
    raise SystemExit(0)

if "V1.10 Demo" not in version_text:
    raise ValueError(f"Expected V1.10 manual, got: {version_text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = "版本：V1.11 Demo   |   更新日期：2026年7月22日   |   适用地址：http://localhost:3000"
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页和六个具体事项页面都使用红塘村无人机正射影像；"
    "进入小花园、茶厂、村里用水、光伏设施、安全隐患或村庄记忆后，页面首先显示同一张村庄地图，并只保留当前事项相关演示点位。"
    "业务点位尚未按影像校准真实位置，不能作为真实设施、风险或产权判断依据。"
)

replacements = {
    "桌面端顶栏：平台名称、六类事项 TAB 和常用操作位于同一行；点击“村庄总览”返回首页。": "桌面端顶栏：平台名称、六类事项 TAB 和常用操作位于同一行；点击“村庄总览”返回首页，点击具体事项则先进入该事项的地图首屏。",
    "小花园、茶厂、村里用水：分别进入自己的独立页面，不再先打开同一个通用模块。": "小花园、茶厂、村里用水：进入后首先使用同一张村庄地图，并分别只显示小花园、茶场与茶厂、用水设施点位。",
    "光伏设施、安全隐患、村庄记忆：分别进入设施管理、风险跟进和村庄档案的独立页面骨架。": "光伏设施、安全隐患、村庄记忆：进入后首先只显示各自地图点位，再向下进入设施管理、风险跟进或村庄档案功能。",
    "系统会进入该事项自己的网址和页面文件；切换顶部标签即可进入另一项，不需要返回通用汇总页。": "系统会进入该事项自己的网址，页面第一块内容就是红塘村无人机地图；切换顶部标签时，底图不变，点位自动切换为新事项。",
    "当前独立页面只显示功能分区骨架和“功能待讨论”提示，具体表单、数据和操作将在确认真实需求后逐项开发。": "地图右侧会说明当前只显示哪些点位；点击标记可查看位置、状态和简介。茶厂页面同时显示茶场与茶厂，其他五页只显示各自事项类型。",
    "需要地图、问题记录或近期变化时，可使用首页三个常用操作；原“村里的事”汇总页仅为兼容旧链接而保留。": "点击“继续看本板块功能”向下查看该事项自己的功能骨架；点击“打开全部村庄地图”可回到包含四组点位的完整地图。",
    "图 2  顶部事项导航与茶厂独立页面": "图 2  茶厂页面的事项地图首屏",
    "使用提示\n六个事项目前已经分开，但页面中的功能名称只是讨论起点，不代表最终业务流程。每一项都需要与村民、村委、老师和调研团队确认后单独设计。": "使用提示\n六个事项共用同一地图底图和点位交互，便于村民保持空间认识；地图下方的功能仍按事项分别设计。所有点位位置目前都是演示数据，需在实地核对后替换。",
}
for old, new in replacements.items():
    find_paragraph(document, old).text = new

set_image(
    image_paragraph_before(document, "图 2  茶厂页面的事项地图首屏"),
    TOPIC_SCREENSHOT,
    6.3,
    "图 2 茶厂页面的事项地图首屏",
    "茶厂页面首先展示村庄无人机地图并只显示茶场与茶厂点位",
)

page_updates = {
    "小花园": "地图首屏只显示小花园点位；下方为四季变化与经验交流功能骨架",
    "茶厂": "地图首屏只显示茶场、茶厂点位；下方为茶园、收茶、加工功能骨架",
    "村里用水": "地图首屏只显示用水设施点位；下方为水源、问题和维修反馈功能骨架",
    "光伏设施": "地图首屏只显示光伏设施点位；下方为运行巡查和信息公开功能骨架",
    "安全隐患": "地图首屏只显示安全隐患点位；下方为发现、处理和复查功能骨架",
    "村庄记忆": "地图首屏只显示村庄记忆点位；下方为照片、讲述、古道与老屋档案骨架",
}
for row in document.tables[1].rows:
    page_name = row.cells[0].text.strip()
    if page_name in page_updates:
        row.cells[2].text = page_updates[page_name]

version_table = document.tables[4]
new_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    new_row.cells,
    [
        "V1.11 Demo",
        "2026-07-22",
        "六个具体事项页面改为共用村庄地图首屏，并按当前事项自动过滤点位；桌面端地图与说明并列，手机端先显示地图再显示事项功能。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert "V1.11 Demo" in check.paragraphs[5].text
assert len(check.tables[4].rows) == 13
assert any("底图不变，点位自动切换" in paragraph.text for paragraph in check.paragraphs)
assert any(row.cells[0].text.strip() == "茶厂" and "茶场、茶厂点位" in row.cells[2].text for row in check.tables[1].rows)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.11: {MANUAL}")
