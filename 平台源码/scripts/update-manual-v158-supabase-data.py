from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)

version_paragraph = next(p for p in document.paragraphs if p.text.startswith("版本：V1."))
replace_paragraph_text(version_paragraph, "版本：V1.58 ｜ 更新日期：2026年8月9日")

data_overview = next(p for p in document.paragraphs if p.text.startswith("当前基础地点数据包含205个有效坐标点"))
replace_paragraph_text(
    data_overview,
    "当前基础地点数据包含205个有效坐标点和586张现场照片；首页另载入7个待核实水系统节点、8条待核实线路和3个待核实供水片区，"
    "共形成223个地图要素。2D与3D现在优先读取同一个Supabase云端数据库：地点、供水专题和专题记录存放在数据表中，586张WebP照片"
    "存放在Storage素材桶中；本地JSON和本地照片继续保留为断网或云端临时不可用时的备用资料。",
)

github_paragraph = next(p for p in document.paragraphs if p.text.startswith("GitHub可作为远程代码仓库"))
replace_paragraph_text(
    github_paragraph,
    "GitHub保存平台源码并用于在线部署；本地版与GitHub在线版使用同一个Supabase项目，因此地图资料和照片无需分别维护两份。",
)

supabase_paragraph = next(p for p in document.paragraphs if p.text.startswith("Supabase后续可用于账号登录"))
replace_paragraph_text(
    supabase_paragraph,
    "Supabase目前负责首页共享数据和现场照片：platform_datasets表保存地点、供水专题和专题记录，hongtang-photos素材桶保存586张WebP照片。"
    "浏览器只使用可公开读取的项目地址和Publishable Key；Service Role或Secret Key只允许在管理员上传数据时临时使用，不能写入源码、网页或公开仓库。",
)

secret_paragraph = next(p for p in document.paragraphs if p.text.startswith("CESIUM_ION_TOKEN等密钥应放在本地环境变量文件中"))
replace_paragraph_text(
    secret_paragraph,
    "CESIUM_ION_TOKEN等私密令牌应放在本地环境变量文件中，不应上传到公开Git仓库。NEXT_PUBLIC_SUPABASE_URL和"
    "NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY用于网页公开读取，可在本地与在线构建中配置；SUPABASE_SERVICE_KEY不得进入浏览器或Git仓库。",
)

data_table = next(table for table in document.tables if table.cell(0, 0).text == "内容")
location_row = next(row for row in data_table.rows if row.cells[0].text == "地点数据")
replace_cell_text(
    location_row.cells[1],
    "Supabase：platform_datasets表、hongtang-photos素材桶；备用：平台源码/public/data、平台源码/public/local-photos、平台素材/现场照片",
)
replace_cell_text(
    location_row.cells[2],
    "本地版和在线版共用地点分类、坐标、简介、水专题、专题记录及586张云端WebP照片；本地文件仅在连接失败时回退使用",
)

faq_table = next(table for table in document.tables if table.cell(0, 0).text == "现象")
photo_row = next(row for row in faq_table.rows if row.cells[0].text == "地点资料没有照片")
replace_cell_text(
    photo_row.cells[1],
    "该记录可能原本就没有录入照片，或网络暂时无法访问Supabase。已有照片会优先从云端素材桶读取；云端不可用时平台尝试读取本地备用照片，"
    "但不会自动补一张虚构图片。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_version_row = next((row for row in version_table.rows if row.cells[0].text == "V1.58"), None)
if new_version_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_version_row = version_table.rows[1]
replace_cell_text(new_version_row.cells[0], "V1.58")
replace_cell_text(new_version_row.cells[1], "2026年8月9日")
replace_cell_text(
    new_version_row.cells[2],
    "本地版与GitHub在线版统一改为优先读取Supabase；上传586张WebP现场照片，并将地点、供水专题和专题记录写入共享数据表。"
    "保留本地JSON和照片作为连接失败时的备用，不改变2D与3D现有交互。",
)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
table_text = "\n".join(cell.text for table in reloaded.tables for row in table.rows for cell in row.cells)
assert any(p.text.startswith("版本：V1.58") for p in reloaded.paragraphs)
assert "优先读取同一个Supabase云端数据库" in paragraph_text
assert "hongtang-photos素材桶" in paragraph_text
assert "SUPABASE_SERVICE_KEY不得进入浏览器或Git仓库" in paragraph_text
assert "Supabase：platform_datasets表" in table_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.58"
print(MANUAL)
