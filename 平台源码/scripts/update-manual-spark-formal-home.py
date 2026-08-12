from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
SCREENSHOT = QA_ROOT / "gaussian-home-1440.png"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v126.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.25.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(
        index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p
    )
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def add_version_row(table):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        properties = cell._tc.get_or_add_tcPr()
        for child in list(properties):
            properties.remove(child)
        for child in template_row.cells[index]._tc.tcPr:
            properties.append(deepcopy(child))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        cell.paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


QA_ROOT.mkdir(parents=True, exist_ok=True)
if not SCREENSHOT.exists():
    raise FileNotFoundError(SCREENSHOT)

document = Document(MANUAL)
if "V1.25 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.25 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.26 Demo   |   更新日期：2026年7月28日   |   适用地址：http://localhost:3000"
)

find_paragraph_prefix(document, "阅读提示").text = (
    "阅读提示\n当前网站是可交互演示版本。正式首页现使用SparkJS 2.1.0加载红塘村轻量化3D高斯模型："
    "先读取约54KiB的RAD索引，再按当前视角载入所需RADC分块，不再调用Cesium ion在线模型。首页展示"
    "小花园、茶场、茶厂、村里用水和光伏设施5个可点击示例点。左键拖动沿模型水平面平移；右键左右或"
    "上下拖动围绕当前屏幕中心目标点旋转；滚轮按固定比例缩放，中键拖动已停用。点击图钉只打开事项详情，"
    "点击“定位到此处”才让镜头飞近。电脑发热或风扇转速较高时，点击右下角齿轮并选择“省电”。"
    "所有三维图钉均为演示位置，待实地核实。"
)

find_paragraph_prefix(document, "以后不必每次输入命令").text = (
    "以后不必每次输入命令。在项目根目录双击“启动网站.bat”，脚本会在后台启动网站并自动打开浏览器。"
    "3D首页默认从外层“3D高斯展示/轻量化模型”读取本地流式文件，不需要Cesium ion令牌；模型目录缺失时"
    "仍可进入“村庄总览”。"
)

find_paragraph_prefix(document, "3D首页操作：").text = (
    "3D首页操作：打开首页后，SparkJS先读取轻量化模型索引，再根据相机位置载入当前视角需要的模型分块。"
    "界面右下角纵向排列“回到中心、操作设置、全屏查看”三个图标按钮。点击齿轮可以分别调整左键平移和"
    "滚轮缩放灵敏度，并在“省电、均衡、清晰”三档画面质量之间切换；三档会改变Spark细节预算和WebGL"
    "渲染像素倍率。左键拖动沿模型的水平XZ平面移动；右键从画面任意位置按下后，围绕当前屏幕中心对应的"
    "OrbitControls目标点旋转，向上接近水平，向下接近垂直俯视；滚轮每次按固定比例缩放，中键不执行操作。"
    "停止操作和分块细化后，页面会暂停连续绘制。点击事项图钉只打开详情，点击“定位到此处”才移动镜头。"
)

find_paragraph_prefix(document, "资料与隐私说明").text = (
    "资料与隐私说明\n公开底图只包含建筑编号、轮廓、新旧类型、高度、估算占地和中心坐标。户主、家庭住址、"
    "家庭成员、人口等字段不会进入网站。Spark首页通过受限接口只读取指定RAD/RADC模型文件，不载入建筑"
    "户主等个人属性；轻量化模型和恢复出的PLY保存在源码目录外，不进入Git仓库。"
)

find_paragraph_prefix(document, "外层“3D高斯展示”文件夹保存").text = (
    "外层“3D高斯展示”文件夹保存Spark官方源码、模型恢复工具、标准PLY和476个流式分块；这些大体积文件"
    "位于源码目录外，不进入Git仓库或Next.js构建产物。正式首页已经使用其中的RAD/RADC轻量化结果。"
)

find_paragraph_prefix(document, "src/app/api/cesium-ion/asset/route.ts").text = (
    "src/app/api/gaussian-model/[filename]/route.ts只允许读取hongtang-recovered-sh0-lod.rad及对应数字编号"
    "的.radc文件，支持HTTP Range、长期缓存和流式返回；public/gaussian-viewer/index.html使用本地"
    "SparkJS 2.1.0、Three.js 0.180.0和OrbitControls完成全屏显示。"
)

find_paragraph_prefix(document, "首页仅加载高斯实景").text = (
    "首页只加载Spark轻量化高斯实景和5个事项图钉，不载入building_footprints.geojson，也不展示户主、"
    "住址或家庭成员等个人属性。SparkJS、Three.js和所需模块保存在public/gaussian-viewer/vendor，"
    "启动时不依赖外部CDN。"
)

find_paragraph_prefix(document, "平台参照旧村庄规划互动平台的Cesium实现").text = (
    "正式首页不再使用Cesium地球、世界地形或在线高斯资产。恢复工具从现有Cesium 3D Tiles的1,036个末级"
    "瓦片提取22,519,595点标准高斯PLY，再用Spark build-lod的--quality、--rad-chunked和--max-sh=0"
    "生成476个分块。分块合计531.70MiB，比原735.31MiB瓦片包减少约27.7%。正式首页实测当前视角请求"
    "52个模型响应、约64.45MiB。模型以Z轴向上的原坐标恢复，查看器绕X轴旋转-90°转换为Three.js的Y轴"
    "向上坐标，并使用模型包围盒中心完成居中。原三维场景内的Cesium地形、全球卫星影像和单独正射影像"
    "已经退出首页；二维“村庄总览”和“村里一张图”继续保留无人机正射影像。"
)

find_paragraph_prefix(document, "平台在Cesium中固定创建").text = (
    "平台在Spark/Three.js局部坐标中固定创建map-13至map-17五个演示锚点；清晰的DOM/SVG图钉每次绘制"
    "后投影到屏幕，React父页面继续从二维mapFeatures读取标题、状态和说明。普通选择不移动相机，"
    "“定位到此处”才执行1.2秒平滑飞近。桌面详情在右侧打开，390px手机端详情在底部打开且最大高度"
    "不超过可视区的55%。取得实测位置后，应把比例位置替换为经核实的模型局部坐标和高程。"
)

find_paragraph_prefix(document, "轻量化试验：").text = (
    "轻量化与正式接入：从现有Cesium 3D Tiles的1,036个末级瓦片恢复出22,519,595点标准高斯PLY，"
    "再使用Spark 2.1.0生成476个流式分块。RAD索引约54KiB，分块合计531.70MiB，比原735.31MiB"
    "瓦片包减少约27.7%。正式首页已使用这套RAD/RADC文件，并通过受限Range接口按当前视角加载；"
    "本机正式首页首屏实测约64.45MiB。"
)

find_paragraph_prefix(document, "3D高斯实景公开或部署到互联网前").text = (
    "3D高斯实景公开或部署到互联网前，应确认航拍成果和轻量化模型的公开授权，并为模型流式接口增加正式"
    "登录、访问控制、带宽限制或对象存储策略；不要把恢复出的PLY和全部RADC文件直接提交到公开Git仓库。"
)

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[1].text = "SparkJS 2.1.0 + Three.js 0.180.0"
        row.cells[2].text = (
            "通过受限Range接口流式加载RAD索引和476个RADC分块；当前视角按需载入，不一次性下载全部模型；"
            "5个DOM/SVG图钉与二维事项详情联动；保留左键水平平移、右键中心旋转、定比缩放、三档画质和"
            "静止后暂停连续绘制"
        )
        break
else:
    raise ValueError("3D technology row not found")

caption = find_paragraph_prefix(document, "图 1")
caption.text = "图 1  正式平台SparkJS流式3D首页、五类示例点与画面质量设置"
image_paragraph = image_paragraph_before(document, caption)
clear_paragraph(image_paragraph)
image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
shape = image_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.3))
shape._inline.docPr.set("name", "图 1 SparkJS流式3D首页")
shape._inline.docPr.set(
    "descr",
    "正式平台SparkJS流式3D首页，显示红塘村高斯模型、五类示例点和三档画面质量设置",
)

version_row = add_version_row(document.tables[4])
for cell, value in zip(
    version_row.cells,
    [
        "V1.26 Demo",
        "2026-07-28",
        "正式首页切换为SparkJS流式加载轻量化RAD/RADC模型；保留五类图钉、详情、相机操作、三档画质和全屏功能。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 28]
assert "V1.26 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "正式首页现使用SparkJS 2.1.0" in all_text
assert "api/gaussian-model/[filename]/route.ts" in all_text
assert "正式首页实测当前视角请求" in all_text
assert "原三维场景内的Cesium地形" in all_text
assert "轻量化与正式接入" in all_text
assert any(
    row.cells[0].text.strip() == "V1.26 Demo" for row in check.tables[4].rows
)
first_line_count = len(
    check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]')
)
assert first_line_count >= 18, first_line_count
assert (
    check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
)
assert (
    check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia"))
    == "SimHei"
)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.26: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
