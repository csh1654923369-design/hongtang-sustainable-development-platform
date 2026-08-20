from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.68 ｜ 更新日期：2026年8月10日",
)

three_d_intro = next(
    p for p in document.paragraphs if p.text.startswith("三维页面使用项目内置的CesiumJS")
)
visibility_note = (
    "在3D中，专题片区与线路会同时覆盖地形和三维模型表面；片区采用更深的专题色和更高的不透明度，"
    "并增加清晰边界，线路则加粗并带浅色描边和遮挡补偿。选中或关联对象会进一步增强，便于在建筑、"
    "植被和航拍纹理上辨认。"
)
if visibility_note not in three_d_intro.text:
    replace_paragraph_text(three_d_intro, three_d_intro.text + visibility_note)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.68"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.68",
    "2026年8月10日",
    "增强3D专题图层可见性：片区与线路同时覆盖地形和三维模型，提高颜色饱和度、片区不透明度、边界与线路宽度，并增加遮挡补偿。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.68") for p in reloaded.paragraphs)
assert "专题片区与线路会同时覆盖地形和三维模型表面" in paragraph_text
assert "浅色描边和遮挡补偿" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.68"
print(MANUAL)
