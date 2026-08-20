from pathlib import Path
import zipfile

from docx import Document
from docx.oxml.ns import qn


source_root = Path(__file__).resolve().parents[1]
manual = source_root.parent / "红塘村可持续发展平台使用手册.docx"

with zipfile.ZipFile(manual) as archive:
    assert archive.testzip() is None

document = Document(manual)
paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
table_text = "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)
first_line_count = len(document.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))

assert any(paragraph.text.startswith("版本：V1.90") for paragraph in document.paragraphs)
assert "首页默认进入红塘村2D地图" in paragraph_text
assert "从上到下排列的通栏悬浮控制区" in paragraph_text
assert "专题卡片与上述控制区保持相同宽度" in paragraph_text
assert "高德在线地图" in paragraph_text
assert "整数屏幕像素定位" in paragraph_text
assert "npm run test:2d-markers" in paragraph_text
assert "原地图左下角的孤立入口已经删除" in paragraph_text
assert "优先读取同一个Supabase云端数据库" in paragraph_text
assert "hongtang-photos素材桶" in paragraph_text
assert "SUPABASE_SERVICE_KEY不得进入浏览器、源码或公开Git仓库" in paragraph_text
assert "4.3 小花园与茶产业专题" in paragraph_text
assert "4.4 塌方与安全、历史与文化专题" in paragraph_text
assert "其他资料" in paragraph_text
assert "其他资料保持原有颜色，并统一使用低遮挡小圆点" in paragraph_text
assert "点击右侧带倒三角的“展开”" in paragraph_text
assert "打开后按钮变为“收起”" in paragraph_text
assert "不重复显示“专题、全选、完成”标题行" in paragraph_text
assert "生态资源" not in paragraph_text
assert "生态资源" not in table_text
assert "2D和3D使用同一规则" in paragraph_text
assert "右下角工具在2D和3D之间采用同一视觉样式" in paragraph_text
assert "2D保留“回到中心、全屏查看”" in paragraph_text
assert "切换按钮位于右上角" in paragraph_text
assert "左组“航拍／手绘”控制村庄覆盖层" in paragraph_text
assert "右组“卫星／底图”控制云端底层" in paragraph_text
assert "每组内部二选一，但两组选择互不重置" in paragraph_text
assert "首次打开网页时默认选择“航拍＋卫星”" in paragraph_text
assert "约260毫秒的高度缓动和轻微透明度变化" in paragraph_text
assert "文字和列表始终保持固定位置" in paragraph_text
assert "不会出现文字、列表或箭头回弹" in paragraph_text
assert "平台名称卡片与专题大卡片保持相同宽度" in paragraph_text
assert "供水片区采用与三维坐标同步的前景专题覆盖层" in paragraph_text
assert "片区基础不透明度为0.31" in paragraph_text
assert "高斯点云不会再遮住片区中央" in paragraph_text
assert "浅色描边和遮挡补偿" in paragraph_text
assert "“卫星”使用高德官方卫星与路网图层" in paragraph_text
assert "例如可同时显示“手绘＋卫星”" in paragraph_text
assert "五个专题会在第一次显示时直接保持全选" in paragraph_text
assert "只有数据读取完成且筛选结果确实为空时" in paragraph_text
assert "2D保留“回到中心、全屏查看”" in paragraph_text
assert "3D仅保留“操作设置、回到中心、全屏查看”" in paragraph_text
assert "6.4 红塘轻量地图数据编辑器" in paragraph_text
assert "http://localhost:3000/map-editor" in paragraph_text
assert "不向Supabase发送新增、修改或删除请求" in paragraph_text
assert "刷新首页或编辑页后会恢复原始数据" in paragraph_text
assert "npm run test:map-editor" in paragraph_text
assert "专题 → 图层 → 要素" in paragraph_text
assert "塌方与安全只允许创建点要素" in paragraph_text
assert "小花园只保留已有的35处小花园位置点" in paragraph_text
assert "不设置花园空间范围、地方知识传播线路" in paragraph_text
assert "保留18个" in paragraph_text
assert "不应被当作红塘村现状事实" in paragraph_text
assert "npm run test:map-editor-layout" in paragraph_text
assert "轻量编辑器完全由红塘项目自身代码实现" in paragraph_text
assert "原首页及其专题、图层和视图状态继续保留" in paragraph_text
assert "最下方第三行为“地图编辑 / 临时试验工具”" in paragraph_text
assert "专题卡片与上述控制区保持相同宽度并紧接在下方" in paragraph_text
assert "最多保留最近5步" in paragraph_text
assert "顶部固定显示“工具栏”文字" in paragraph_text
assert "连续输入同一字段、一次完整拖动分别只算一步" in paragraph_text
assert "对应箭头会自动置灰" in paragraph_text
assert "先用约260毫秒完成右上角卡片过渡" in paragraph_text
assert "稳定状态下只运行当前地图" in paragraph_text
assert "查看全部地点" not in paragraph_text
assert "光伏不纳入当前五个专题" in paragraph_text
assert "npm run test:village-topics" in paragraph_text
assert "Supabase：platform_datasets表" in table_text
assert "2D图钉有时模糊，或点击后没有出现详情" in table_text

assert "手绘图蓝色水系" in paragraph_text
assert "不用于推断地下供水管道" in paragraph_text
assert "先向南、再转向东南" in paragraph_text
assert "19个地图要素" in table_text
assert "7个节点、7条线路和3个片区" in table_text
assert "当前“航拍／手绘”和“卫星／底图”组合保持不变" in paragraph_text
assert "公开访问地址：https://hongtang-sdg-platform.vercel.app" in paragraph_text
assert "6.3 Git、Vercel与Supabase" in paragraph_text
assert "发布网站到Vercel.bat" in paragraph_text
assert "不必先上传GitHub" in paragraph_text
assert "Cesium官方发布地址加载运行库" in paragraph_text
assert "Allowed URLs应包含https://hongtang-sdg-platform.vercel.app" in paragraph_text
assert "403通常表示域名未获准" in table_text
assert len(document.tables) == 6
assert [cell.text for cell in document.tables[0].rows[0].cells] == ["地点类型", "当前数量", "显示方式"]
assert [len(table.rows) for table in document.tables] == [6, 6, 7, 9, 14, 53]
assert [cell.text for cell in document.tables[4].rows[0].cells] == ["现象", "处理方法"]
assert [cell.text for cell in document.tables[5].rows[0].cells] == ["版本", "日期", "主要变化"]
assert document.tables[5].rows[1].cells[0].text == "V1.90"
assert first_line_count >= 22
assert document.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
assert document.styles["Normal"].font.name == "Times New Roman"
assert document.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
assert document.styles["Heading 2"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"

print(
    {
        "version": next(p.text for p in document.paragraphs if p.text.startswith("版本：")),
        "paragraphs": len(document.paragraphs),
        "table_rows": [len(table.rows) for table in document.tables],
        "native_two_character_indents": first_line_count,
        "zip": "ok",
    }
)
