from copy import deepcopy
from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v13.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.2.backup.docx"
SCREENSHOTS = {
    "home": QA_ROOT / "home-1440.png",
    "resources": QA_ROOT / "community-resources-1440.png",
    "actions": QA_ROOT / "micro-actions-1440.png",
    "action_form": QA_ROOT / "micro-action-form-1440.png",
    "capacity": QA_ROOT / "community-capacity-1440.png",
}


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)


def set_image(paragraph, image_path, width, name, description):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("name", name)
    shape._inline.docPr.set("descr", description)
    return shape


def insert_image_before(anchor, image_path, width, name, description):
    paragraph = anchor.insert_paragraph_before()
    set_image(paragraph, image_path, width, name, description)
    return paragraph


def image_paragraph_before(document, caption_text):
    caption = find_paragraph(document, caption_text)
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError(f"Image paragraph not found before caption: {caption_text}")


def copy_cell_format(source_cell, target_cell):
    target_properties = target_cell._tc.get_or_add_tcPr()
    for child in list(target_properties):
        target_properties.remove(child)
    for child in source_cell._tc.tcPr:
        target_properties.append(deepcopy(child))


def set_cell_text_like(source_cell, target_cell, text):
    copy_cell_format(source_cell, target_cell)
    target_cell.text = text
    if source_cell.paragraphs and target_cell.paragraphs:
        source = source_cell.paragraphs[0]
        target = target_cell.paragraphs[0]
        target.style = source.style
        target.paragraph_format.alignment = source.paragraph_format.alignment


def add_row_like(table, template_row, values, before_row=None):
    row = table.add_row()
    for index, value in enumerate(values):
        set_cell_text_like(template_row.cells[index], row.cells[index], value)
    if before_row is not None:
        before_row._tr.addprevious(row._tr)
    return row


def remove_extra_page_breaks(document):
    """Keep the cover break, then let later chapters flow without blank pages."""
    kept_cover_break = False
    for paragraph in list(document.paragraphs):
        page_breaks = paragraph._p.xpath(".//w:br[@w:type='page']")
        if not page_breaks:
            continue
        if not kept_cover_break:
            kept_cover_break = True
            page_breaks = page_breaks[1:]
        if not page_breaks:
            continue
        for page_break in page_breaks:
            page_break.getparent().remove(page_break)
        if not paragraph.text.strip() and not paragraph._p.xpath(".//w:drawing"):
            paragraph._element.getparent().remove(paragraph._element)


def keep_table_rows_intact(table):
    """Avoid splitting a role or version-history record across pages."""
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        if not properties.xpath("./w:cantSplit"):
            properties.append(OxmlElement("w:cantSplit"))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    if not header_properties.xpath("./w:tblHeader"):
        header_properties.append(OxmlElement("w:tblHeader"))


for screenshot in SCREENSHOTS.values():
    if not screenshot.exists():
        raise FileNotFoundError(screenshot)

document = Document(MANUAL)
version_text = document.paragraphs[5].text
if "V1.3 Demo" in version_text:
    print("Manual is already V1.3; no changes applied.")
    raise SystemExit(0)
if "V1.2 Demo" not in version_text:
    raise ValueError(f"Expected V1.2 manual, got: {version_text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

# Cover and global positioning.
set_image(image_paragraph_before(document, "图 1  平台首页"), SCREENSHOTS["home"], 6.3, "图 1 平台首页", "平台首页新增发起小行动入口")
document.paragraphs[5].text = "版本：V1.3 Demo   |   更新日期：2026年7月18日   |   适用地址：http://localhost:3000"
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。问题、项目、微行动、资源供需、进展指标和坐标均含演示数据。"
    "建筑调研底图来自既有调研资料并已去除个人字段，仅用于规划讨论，不作为权属、测绘或审批依据。"
)
find_paragraph(document, "重要：演示会话\n问题上报、活动报名和项目关注保存在当前浏览器内存中。页面内跳转时会保留；刷新网页、关闭标签页或重启服务后可能恢复为初始演示数据。").text = (
    "重要：演示会话\n问题上报、微行动、资源登记、任务认领、活动报名和项目关注保存在当前浏览器内存中。"
    "页面内跳转时会保留；刷新网页、关闭标签页或重启服务后会恢复为初始演示数据。"
)

# Navigation copy.
find_paragraph(document, "首页：了解平台定位和近期行动入口。").text = "首页：了解平台定位，并从行动地图、发起微行动或公众参与快速开始。"
find_paragraph(document, "行动地图：按图层、状态和目标筛选空间点位。").text = "行动地图：按行动、问题、资源供需和空间资料分组筛选点位。"
find_paragraph(document, "项目与行动：筛选项目并进入项目详情。").text = "项目与行动：浏览居民微行动，或筛选正式项目并查看开放任务与资源需求。"
find_paragraph(document, "发展进展：查看趋势图、指标来源和完整度。").text = "发展进展：查看趋势、来源、完整度和社区行动能力。"
find_paragraph(document, "公众参与：提出建议、报名活动、填写问卷、参与共创讨论。").text = "公众参与：发起微行动，提出建议、报名活动、填写问卷并参与共创讨论。"
find_paragraph(document, "个人中心：查看当前角色相关的上报、报名、关注和通知。").text = "个人中心：查看我的上报、我发起或加入的行动、报名、关注和通知。"

# Map chapter: add community resource board before the building survey section.
map_intro = find_paragraph(document, "地图页面分为两部分：上方行动地图把村庄问题、建设项目、已完成行动、公共服务、生态资源、文化资源、调研照片、建筑、道路和水体放在同一个空间界面中，其底图、坐标与点位均为演示数据；下方建筑调研底图复用既有建筑轮廓资料，并已移除个人调查字段。")
map_intro.text = (
    "地图页面分为三部分：上方行动地图汇集问题、正式项目、社区微行动、已完成行动和资源供需点位；"
    "中部互助板按空间、工具、材料、技能、地方知识和志愿时间展示可提供资源与需求；"
    "下方建筑调研底图复用既有建筑轮廓资料并移除个人字段。除建筑轮廓外，坐标与业务点位均为演示数据。"
)
find_paragraph(document, "在左侧勾选或取消图层，例如村庄问题、建设项目或生态资源。").text = "在左侧按“行动与问题、社区互助资源、空间与调研资料”勾选或取消图层。"
find_paragraph(document, "按状态、所属目标、公众参与或“与我有关”进一步筛选。").text = "按待轻量核对、招募伙伴、试验中、可提供、需求中等状态，以及所属目标、公众参与或“与我有关”继续筛选。"
find_paragraph(document, "点击详情按钮进入相关问题、项目或目标页面。").text = "点击详情按钮进入相关问题、项目或微行动区；资源点位可提交回应意向。"

building_heading = find_paragraph(document, "3.2 查看建筑调研底图")
building_heading.text = "3.3 查看建筑调研底图"
building_heading.insert_paragraph_before("3.2 查看社区资源地图与互助板", style="Heading 2")
building_heading.insert_paragraph_before("向下滚动到“社区资源地图与互助板”，可在“全部、可提供、正在寻找”之间切换。", style="List Number")
building_heading.insert_paragraph_before("按空间、工具、材料、技能、地方知识和志愿时间筛选资源类型。", style="List Number")
building_heading.insert_paragraph_before("村民或协作者可登记可提供资源或发布需求，填写说明、时间和模糊位置。", style="List Number")
building_heading.insert_paragraph_before("点击“我需要这项资源”或“我可以帮忙”提交回应；平台或行动小组代为转达，不公开个人联系方式。", style="List Number")
insert_image_before(building_heading, SCREENSHOTS["resources"], 6.3, "图 3 社区资源地图与互助板", "社区资源互助板展示可提供资源和正在寻找的需求")
building_heading.insert_paragraph_before("图 3  社区资源地图与互助板", style="Caption")
privacy = building_heading.insert_paragraph_before()
privacy.add_run("资源隐私边界\n").bold = True
privacy.add_run("公开页面只显示授权字段和模糊位置。标记为“小组内确认”的资源不显示坐标，匹配后再确认具体位置和联系人。")
find_paragraph(document, "3.3 从地图上报问题").text = "3.4 从地图上报问题"

# Project and action chapter.
project_image = image_paragraph_before(document, "图 5  项目筛选与项目卡片")
set_image(project_image, SCREENSHOTS["actions"], 6.3, "图 6 社区微行动", "项目与行动页的居民微行动招募区")
find_paragraph(document, "图 5  项目筛选与项目卡片").text = "图 6  社区微行动与伙伴招募"
find_paragraph(document, "5.1 查找项目").text = "5.1 浏览和加入社区微行动"
find_paragraph(document, "打开“项目与行动”。").text = "打开“项目与行动”，先查看顶部“社区微行动”区。"
find_paragraph(document, "使用所属目标、项目状态、项目类型和“仅看招募中”筛选条件。").text = "微行动卡片会说明 30—90 天试验周期、希望发生的变化、已有与所需资源、招募角色和下一步。"
find_paragraph(document, "点击项目卡片进入详情；点击“重置”恢复全部项目。").text = "切换为村民或协作者，点击“加入行动”；行动随后出现在“个人中心 → 我的行动”。"

project_detail_heading = find_paragraph(document, "5.2 项目详情与参与")
project_detail_heading.text = "5.4 正式项目详情与参与"
project_detail_heading.insert_paragraph_before("5.2 五步发起一个小行动", style="Heading 2")
project_detail_heading.insert_paragraph_before("进入“发起社区微行动”，填写行动名称、关联目标、简要说明和希望发生的变化。", style="List Number")
project_detail_heading.insert_paragraph_before("在简化地图选择大致位置，并填写地点概称。", style="List Number")
project_detail_heading.insert_paragraph_before("盘点已经具备的伙伴、工具或资料，再填写仍需资源和希望招募的角色。", style="List Number")
project_detail_heading.insert_paragraph_before("选择 30、60 或 90 天周期，说明第一次行动、共同决策方式和维护或退出安排。", style="List Number")
project_detail_heading.insert_paragraph_before("核对后提交。系统生成 HTA-2026-xxxx 编号，先进入“待轻量核对”；通过后开放招募。", style="List Number")
insert_image_before(project_detail_heading, SCREENSHOTS["action_form"], 6.3, "图 7 五步微行动向导", "发起社区微行动的五步表单")
project_detail_heading.insert_paragraph_before("图 7  五步发起社区微行动", style="Caption")
project_detail_heading.insert_paragraph_before("5.3 查找正式项目", style="Heading 2")
project_detail_heading.insert_paragraph_before("向下查看“村庄共建项目”，使用所属目标、项目状态、类型和“仅看招募中”筛选条件。", style="List Number")
project_detail_heading.insert_paragraph_before("项目卡会直接显示开放任务名额和仍需匹配的资源数量；点击卡片进入详情。", style="List Number")
find_paragraph(document, "查看项目概况、负责人、预算标签、参与主体、进度和更新时间线。").text = "查看项目来源、概况、负责人、预算标签、参与主体、进度和更新时间线。"
find_paragraph(document, "拖动前后对比组件，比较现状与方案占位场景。").text = "在“依据、感受与行动约定”中同时查看事实依据、不同使用者声音、共同决策方式和维护责任。"
find_paragraph(document, "村民、协作者或管理员可以关注项目；关注后可在个人中心查看。").text = "在“现在可以一起做什么”中按时间和角色认领开放任务，或回应工具、材料、技能、知识、空间和志愿时间需求。"
find_paragraph(document, "参与方案 A/B 投票，提交不少于 6 个字的项目建议。").text = "查看下一次碰面时间和协调人，也可关注项目、参与方案 A/B 投票并提交不少于 6 个字的建议。"

# Participation, progress and profile copy.
participation_heading = find_paragraph(document, "6.1 村庄建议")
participation_heading.insert_paragraph_before("参与中心顶部的“发起一个小行动”连接五步向导。建议适合表达意见；微行动适合已经准备组织伙伴、资源和第一次实践的想法。")
progress_heading = find_paragraph(document, "7.2 发展进展")
progress_image = image_paragraph_before(document, "图 7  发展进展与趋势图")
progress_image.insert_paragraph_before("进展页首先展示“社区行动能力”，观察居民是否能够发起行动、匹配本地资源、持续参与并承担长期维护。")
insert_image_before(progress_image, SCREENSHOTS["capacity"], 6.3, "图 9 社区行动能力", "社区行动能力包括居民发起、资源连接、持续参与和长期维护")
progress_image.insert_paragraph_before("图 9  社区行动能力指标", style="Caption")
find_paragraph(document, "打开“发展进展”，先查看更新时间、指标数量和平均完整度。").text = "打开“发展进展”，先查看社区行动能力、更新时间、指标数量和平均完整度。"
find_paragraph(document, "切换“近 3 个月 / 近 6 个月”观察趋势变化。").text = "社区能力包括居民自主发起微行动数、本地资源需求匹配率、重复参与率和行动三个月持续维护率。"
find_paragraph(document, "将鼠标放在图表上查看具体演示数值。").text = "继续切换“近 3 个月 / 近 6 个月”，将鼠标放在图表上查看具体演示数值。"
find_paragraph(document, "游客进入个人中心时需先选择村民或协作者。村民可查看我的上报、建议、活动、问卷、关注和通知；协作者重点查看调研提交；管理员还可进入管理后台。").text = (
    "游客进入个人中心时需先选择村民或协作者。村民和协作者可在“我的行动”查看自己发起或加入的微行动、当前状态、资源需求和下一步；"
    "村民还可查看上报、建议、活动、问卷、关注和通知，协作者可查看调研提交，管理员可进入管理后台。"
)

# Admin, implementation and boundaries.
find_paragraph(document, "数据概览：查看待审核问题、处理中问题、待审核调研和近期操作。").text = "数据概览：查看待审核问题、待核对微行动、处理中问题、待审核调研和近期操作。"
find_paragraph(document, "问题上报：选择问题，修改办理状态，分派责任人并填写处理说明。").insert_paragraph_before(
    "微行动核对：只检查安全、权限、隐私和资源冲突；可匹配行动协调员、要求补充或开放招募。", style="List Bullet"
)
find_paragraph(document, "导出演示数据：下载 JSON 文件，不包含真实村庄或个人数据。").text = "导出演示数据：下载包含问题、微行动、互助资源、项目与指标的 JSON 文件，不包含真实村庄或个人数据。"
find_paragraph(document, "src/data/mockData.ts 提供统一演示数据。").text = "src/data/mockData.ts 提供目标、项目、问题和指标；src/data/communityData.ts 提供微行动与互助资源演示种子。"
find_paragraph(document, "src/services 下的服务层读取议题、项目、指标、活动和内容。").text = "src/services 下的服务层读取议题、项目、指标、活动和公开内容；社区记录由根级 DemoProvider 统一管理。"
find_paragraph(document, "DemoProvider 在当前浏览器会话中保存角色、问题、报名和关注状态。").text = "DemoProvider 在当前浏览器会话中保存角色、问题、微行动、资源、报名、任务认领和关注状态。"
find_paragraph(document, "src/lib/permissions.ts 根据角色判断是否允许上报、报名、评价、审核等操作。").text = "src/lib/permissions.ts 根据角色判断是否允许上报、发起行动、共享资源、报名、评价或审核。"
find_paragraph(document, "真实数据接入方向\n下一阶段会逐个把 Mock Service 替换为 Supabase 查询与写入，并接入真实注册登录、会话、角色和文件上传。界面组件无需整体重写。").text = (
    "真实数据接入方向\n下一阶段优先新增微行动、社区资源、项目任务和资源回应表及 RLS，再逐个把会话状态与 Mock Service 替换为 Supabase 查询和写入。"
    "随后接入真实注册登录、角色和文件上传，界面组件无需整体重写。"
)
privacy_admin = find_paragraph(document, "管理员角色仅为 Demo 切换，不等同于真实登录或安全认证。")
privacy_admin.insert_paragraph_before("互助资源不得公开家庭住址、电话和个人身份；小组内资源在匹配前也不显示坐标。", style="List Bullet")
privacy_admin.insert_paragraph_before("微行动轻量核对不等于施工、资金或公共安全审批；涉及这些事项时仍需另行确认。", style="List Bullet")
map_limit = find_paragraph(document, "页面业务数据全面写入 Supabase。")
map_limit.insert_paragraph_before("微行动、社区资源、项目任务和资源回应的跨会话持久化；当前刷新后恢复演示种子。", style="List Bullet")

# Tables.
role_table = document.tables[0]
role_table.rows[1].cells[1].text = "浏览公开内容、地图、微行动、资源互助、项目和发展进展"
role_table.rows[1].cells[2].text = "首页、地图、项目、进展"
role_table.rows[2].cells[1].text = "上报问题、发起或加入微行动、共享资源、认领任务、报名、建议与评价"
role_table.rows[2].cells[2].text = "微行动、地图、项目、公众参与、个人中心"
role_table.rows[3].cells[1].text = "发起或加入微行动、共享资源、认领任务、提交调研成果与建议"
role_table.rows[3].cells[2].text = "微行动、地图、项目、研究协作、个人中心"
role_table.rows[4].cells[1].text = "轻量核对微行动、处理问题、审核资料、维护项目和查看日志"
role_table.rows[4].cells[2].text = "管理后台"

page_table = document.tables[1]
for row in page_table.rows:
    page = row.cells[0].text.strip()
    if page == "行动地图":
        row.cells[2].text = "问题、微行动、资源供需、空间资料和建筑调研底图"
    elif page == "项目与行动":
        row.cells[2].text = "社区微行动、正式项目、开放任务与资源需求"
    elif page == "发展进展":
        row.cells[2].text = "社区行动能力、趋势、来源与完整度"
    elif page == "公众参与":
        row.cells[2].text = "微行动入口、建议、活动、问卷和共创"
    elif page == "个人中心":
        row.cells[2].text = "我的行动、个人记录与通知"
    elif page == "管理后台":
        row.cells[2].text = "问题处理、微行动核对、审核与维护"
project_row = next(row for row in page_table.rows if row.cells[0].text.strip() == "项目与行动")
add_row_like(page_table, page_table.rows[-1], ["发起微行动", "/actions/new", "五步形成可核对、可招募的小行动"], before_row=project_row)

tech_table = document.tables[2]
add_row_like(tech_table, tech_table.rows[-1], ["会话状态", "DemoProvider + Mock 数据", "保存微行动、资源回应、任务认领等演示交互"])

faq_table = document.tables[3]
for row in faq_table.rows:
    if row.cells[0].text.strip() == "管理员功能只有提示":
        row.cells[2].text = "问题处理、微行动核对和调研审核可完整演示；其他模块待开发"
add_row_like(faq_table, faq_table.rows[-1], ["微行动或资源刷新后消失", "Demo 状态只保存在当前浏览器会话", "通过页面链接连续演示；跨会话保存待接入 Supabase"])

version_table = document.tables[4]
add_row_like(
    version_table,
    version_table.rows[-1],
    [
        "V1.3 Demo",
        "2026-07-18",
        "新增居民微行动五步发起、管理员轻量核对、伙伴招募、社区资源供需地图、项目任务与资源回应，以及社区行动能力指标。",
    ],
)

for table in document.tables:
    keep_table_rows_intact(table)

# The V1.2 manual starts every major chapter with a hard page break. After the
# new screenshots and sections were added, those breaks produced several mostly
# empty pages and one fully blank page. Retain only the cover break.
remove_extra_page_breaks(document)

# Renumber all captions after inserting three new figures.
caption_updates = {
    "图 3  建筑调研底图与建筑档案": "图 4  建筑调研底图与建筑档案",
    "图 4  五步问题上报向导": "图 5  五步问题上报向导",
    "图 6  公众参与中心": "图 8  公众参与中心",
    "图 7  发展进展与趋势图": "图 10  发展进展与趋势图",
    "图 8  管理后台数据概览": "图 11  管理后台数据概览",
    "图 9  数字沙盘概念演示": "图 12  数字沙盘概念演示",
}
for old, new in caption_updates.items():
    find_paragraph(document, old).text = new

# Normalize image names and alt text in document order.
captions = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.style and paragraph.style.name == "Caption" and paragraph.text.strip().startswith("图 ")]
if len(captions) != len(document.inline_shapes):
    raise AssertionError(f"Caption/image mismatch before save: {len(captions)} captions, {len(document.inline_shapes)} images")
for index, (shape, caption) in enumerate(zip(document.inline_shapes, captions), start=1):
    shape._inline.docPr.set("name", f"Figure {index}")
    shape._inline.docPr.set("descr", caption)

document.save(OUTPUT)

# Structural gate before replacing the formal manual.
check = Document(OUTPUT)
assert "V1.3 Demo" in check.paragraphs[5].text
assert len(check.inline_shapes) == 12
assert any(paragraph.text == "3.2 查看社区资源地图与互助板" for paragraph in check.paragraphs)
assert any(paragraph.text == "5.2 五步发起一个小行动" for paragraph in check.paragraphs)
assert any(paragraph.text == "5.3 查找正式项目" for paragraph in check.paragraphs)
assert any(row.cells[0].text.strip() == "发起微行动" for row in check.tables[1].rows)
assert check.tables[4].rows[-1].cells[0].text == "V1.3 Demo"
assert check.tables[3].rows[-1].cells[0].text == "微行动或资源刷新后消失"
assert sum(len(paragraph._p.xpath(".//w:br[@w:type='page']")) for paragraph in check.paragraphs) == 1
assert all(row._tr.xpath("./w:trPr/w:cantSplit") for table in check.tables for row in table.rows)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual: {MANUAL}")
print(f"Backup: {BACKUP}")
print(f"Paragraphs: {len(check.paragraphs)} | tables: {len(check.tables)} | images: {len(check.inline_shapes)}")
