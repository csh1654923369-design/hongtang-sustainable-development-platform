from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
SCREENSHOT = QA_ROOT / "gaussian-home-1440.png"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v125.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.24.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(
        index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p
    )
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def add_version_row(table):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        properties = cell._tc.get_or_add_tcPr()
        for child in list(properties):
            properties.remove(child)
        for child in template_row.cells[index]._tc.tcPr:
            properties.append(deepcopy(child))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        cell.paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


QA_ROOT.mkdir(parents=True, exist_ok=True)
if not SCREENSHOT.exists():
    raise FileNotFoundError(SCREENSHOT)

document = Document(MANUAL)
if "V1.24 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.24 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.25 Demo   |   更新日期：2026年7月28日   |   适用地址：http://localhost:3000"
)

reading_tip = find_paragraph_prefix(document, "阅读提示")
reading_tip.text = (
    "阅读提示\n当前网站是可交互演示版本。首页把红塘村在线3D高斯实景按真实地理坐标嵌入周边三维地形，"
    "远处显示全球卫星影像，村庄附近叠加红塘无人机正射影像。首页展示小花园、茶场、茶厂、村里用水和"
    "光伏设施5个可点击示例点。左键拖动沿村庄水平面平移；右键左右或上下拖动均围绕屏幕画面正中心"
    "对应的三维点运动；滚轮按固定比例缩放，中键拖动已停用。点击图钉只打开事项详情，点击“定位到此处”"
    "才会让镜头飞近。电脑发热或风扇转速较高时，点击右下角齿轮，在“画面质量”中选择“省电”；停止操作后，"
    "静止画面会自动暂停连续绘制。所有三维图钉均为演示位置，待实地核实。"
)

operation = find_paragraph_prefix(document, "3D首页操作：")
operation.text = (
    "3D首页操作：打开首页后，平台会依次加载周边三维地形、远景卫星影像、红塘无人机正射影像和高斯模型。"
    "高斯模型外侧显示连续山地环境；地形继续用于展示周边山势、承载影像和提供高程采样，但不会遮挡高斯实景。"
    "界面右下角纵向排列“回到中心、操作设置、全屏查看”三个图标按钮。点击齿轮打开“操作设置”，可以分别调整"
    "左键平移和滚轮缩放灵敏度，也可以在“省电、均衡、清晰”三档画面质量之间切换；默认使用“均衡”，电脑发热"
    "或风扇转速较高时建议改为“省电”，“清晰”适合需要观察细节或截图时临时使用。三项设置会保存在当前浏览器中。"
    "查看器采用按需绘制，停止拖动、缩放和数据加载后会暂停持续刷新，以降低空闲时的显卡占用。按住左键拖动，"
    "可沿村庄所在地表的水平面平移；右键左右拖动时相机围绕屏幕中心点水平环绕，向上拖动时视线接近水平，"
    "向下拖动时接近垂直俯视。滚轮每次按固定比例放大或缩小，中键拖动不执行任何操作。点击事项图钉只打开详情，"
    "点击“定位到此处”才移动镜头。在线地形或远景影像临时加载失败时，页面会保留高斯模型和本地无人机影像。"
)

technical = find_paragraph_prefix(
    document, "平台参照旧村庄规划互动平台的Cesium实现"
)
technical.text = (
    "平台参照旧村庄规划互动平台的Cesium实现，恢复地球表面并接入Cesium World Terrain和"
    "Cesium World Imagery；红塘0.3米无人机正射影像按真实WGS84范围铺到地形表面。高斯资产根变换的原点为"
    "99.908740607°E、24.636255278°N、1764m，模型、地形和影像使用同一地理坐标空间。查看器将"
    "depthTestAgainstTerrain设为false，使世界地形继续提供周边环境、影像承载和高程采样，但不遮挡"
    "最上层高斯实景。画面质量三档会同时调节高斯瓦片细节、地形细节和内部渲染分辨率：“省电”为24/8/0.72，"
    "“均衡”为16/6/0.88，“清晰”为9/4/1，依次表示模型屏幕空间误差、地形屏幕空间误差和分辨率倍率。"
    "Cesium查看器启用requestRenderMode并将maximumRenderTimeChange设为Infinity；相机、图钉或设置变化时"
    "主动请求新画面，静止且无数据变化时不再持续绘制。三维图钉在initialTilesLoaded事件之后拾取高斯表面，"
    "再以当地地形高度提供最低高度保护。自动测试会校验画面质量切换、按需绘制标记、5个图钉的位置和相机操作。"
)

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载高斯模型、Cesium全球地形和全球影像；按真实经纬度铺设红塘无人机正射影像；"
            "5个事项图钉在模型瓦片加载后贴合高斯表面，并以当地地形高度作为最低高度保护；"
            "保留屏幕中心旋转、水平平移和定比缩放；提供“省电、均衡、清晰”三档画面质量，"
            "通过按需绘制降低画面静止时的显卡占用"
        )
        break
else:
    raise ValueError("3D technology row not found")

heading_105 = find_paragraph_prefix(document, "10.5 无人机正射影像")
lightweight_note = heading_105.insert_paragraph_before(
    "轻量化试验：从现有Cesium 3D Tiles的1,036个末级瓦片恢复出22,519,595点标准高斯PLY，"
    "再使用Spark 2.1.0的build-lod执行--quality、--rad-chunked和--max-sh=0，生成476个流式分块。"
    "RAD索引约54KiB，分块合计531.70MiB，比原735.31MiB瓦片包减少约27.7%；独立预览首屏实测请求"
    "约77.54MiB。该版本保存在外层“3D高斯展示/轻量化模型”中，目前用于技术预览。正式首页仍使用Cesium，"
    "以保留周边地形、无人机正射影像、真实地理坐标和三维图钉的完整联动。"
)
lightweight_note.style = "List Bullet"

heading_13 = find_paragraph_prefix(document, "13. 使用手册维护规则")
faq_heading = heading_13.insert_paragraph_before("12.2 3D首页发热或风扇转速较高")
faq_heading.style = "Heading 2"
faq_step_1 = heading_13.insert_paragraph_before(
    "点击3D画面右下角的齿轮图标，在“画面质量”中选择“省电”。该设置会自动保存在当前浏览器中。"
)
faq_step_1.style = "List Bullet"
faq_step_2 = heading_13.insert_paragraph_before(
    "停止拖动和缩放并等待数据加载结束。静止画面会自动暂停连续绘制；“清晰”档建议只在观察细节或截图时临时启用。"
)
faq_step_2.style = "List Bullet"

caption = find_paragraph_prefix(document, "图 1")
caption.text = "图 1  3D高斯实景首页、五类示例点与画面质量设置"
image_paragraph = image_paragraph_before(document, caption)
clear_paragraph(image_paragraph)
image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
shape = image_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.3))
shape._inline.docPr.set("name", "图 1 3D高斯实景首页与画面质量设置")
shape._inline.docPr.set(
    "descr",
    "红塘村3D高斯实景首页，显示五类示例点和右下角的平移、缩放及画面质量设置",
)

version_row = add_version_row(document.tables[4])
for cell, value in zip(
    version_row.cells,
    [
        "V1.25 Demo",
        "2026-07-28",
        "完成高斯模型轻量化试验；3D首页增加三档画面质量并启用按需绘制，降低静止画面显卡占用。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 27]
assert "V1.25 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "选择“省电”" in all_text
assert "requestRenderMode" in all_text
assert "生成476个流式分块" in all_text
assert "12.2 3D首页发热或风扇转速较高" in all_text
assert any(
    row.cells[0].text.strip() == "V1.25 Demo" for row in check.tables[4].rows
)
first_line_count = len(
    check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]')
)
assert first_line_count >= 18, first_line_count
assert (
    check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
)
assert (
    check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia"))
    == "SimHei"
)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.25: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
