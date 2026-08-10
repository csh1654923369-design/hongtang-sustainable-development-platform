from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
OUTPUT = WORKSPACE_ROOT / "0719讨论_修订版录音稿.docx"

BLACK = RGBColor(0, 0, 0)
BODY_CN = "SimSun"
BODY_WEST = "Times New Roman"
HEADING_CN = "SimHei"


TRANSCRIPT = [
    ("老师", "00:00", "做了什么出来？它是好多平台的一个集合。"),
    ("我", "00:08", "小花园可以单独做一个大的模块进去。"),
    (
        "老师",
        "00:11",
        "是啊，那在这个模块里面做什么呢？别人通过这个模块，要看它跟可持续发展有什么关系。小花园，我不是给你那个东西了吗？就是给你看到底小花园跟它有什么关系：种了什么菜？今年种什么菜？明年种什么菜？有没有记录？我们准备做十年的数据收集。",
    ),
    ("我", "00:41", "每一年、哪一户？"),
    (
        "老师",
        "00:43",
        "不是只看每一年。每一次、每个季度都可能经常换，他们就发布一下，是吧？",
    ),
    ("我", "00:51", "然后他们拍的视频……"),
    (
        "老师",
        "00:52",
        "照片也拍，也可以发。就是把我们给你发的那个小花园平台做出来嘛，对不对？",
    ),
    (
        "老师",
        "01:18",
        "那第二个、第三个是什么呢？我要看茶园。他们都种茶，茶园都是有机茶园。那你是学地理的，有机茶园我要监控什么东西？我怎么看它是有机茶园？他要做茶园，要收茶叶、收青茶、采青茶。我们上面还有一个茶厂，你不是去看过吗？那茶厂不是在收茶叶吗？收多少？谁来收？都要在我们这个体系里体现。",
    ),
    ("我", "01:55", "确实，这个就是针对红塘的，很具体。"),
    ("老师", "02:00", "对，做红塘的。"),
    (
        "我",
        "02:01",
        "然后就是有一个具体的，因为像小花园和茶园，它们两个是红塘很重要的特色。",
    ),
    ("老师", "02:09", "然后还有……"),
    (
        "老师",
        "02:28",
        "还有什么？这个水厂、这个水系统是怎么建的？因为他们的水都是自己集中起来建设。做水系统不是可持续吗？这些呢，是第三个。他们现在还做光伏。第五个，现在他们经常有塌方，塌方点到底在哪里？有几个点？你把我们以前的规划拿来看吧。你们一定要看我们以前做的事。老师的东西从来一直是连续的。有时候你们看我东一榔头西一榔头，其实不是的，其实都是一个整体。老师就一个脑袋，没有那么精神分裂。你把以前我们红塘的规划拿来看吧。",
    ),
    (
        "我",
        "03:30",
        "那我觉得，刚刚这几件事情其实就是一个可以行动的板块。",
    ),
    (
        "老师",
        "03:36",
        "这个板块就是你的行动。那你要把这个板块到底研究什么东西想清楚。假设你十年以后要写一篇报告或论文，那你现在就要开始收数据。然后你把它变成一个众包平台。群众在这里能够经常有获得感，看到自己的变化。领导来了，比如书记一来，我们给他展示这个东西。〔此处原词听辨不清，疑指外国来宾〕、联合国来了，觉得我们这个东西是世界前沿。我们的目标就是这个。",
    ),
    (
        "我",
        "04:21",
        "其实跟刚刚那个〔平台名称听辨不清〕有很多共同之处，就是又要……",
    ),
    (
        "老师",
        "04:26",
        "哎呀，我就是一个脑袋，对，我就是一个脑袋。你别看我一会儿做乡村，一会儿做这个，一会儿做城市，一会儿做社区，它们就是一件事。哪天你可以把这个思维链条拎出来。",
    ),
    (
        "老师",
        "04:52",
        "我觉得，你跟我说，为什么会去做技能？我们当时就知道人没有了。人没有了，如何保证珠三角的产业发展？那就是机器换人。机器换人就必须研究自动化、智能化。我们那个时候，对每一个智能化里面的每一个部分、每一台设备，都要看后面一群人在干嘛，走来走去。听明白。继续想，继续找，你们要跟他们讨论。",
    ),
    ("我", "05:39", "因为我之前就是想做一个通用的平台出来。"),
    ("老师", "05:41", "但是其实不要做通用平台。"),
    ("我", "05:44", "只要给红塘村用就可以了。"),
    ("老师", "05:46", "现在就是这样。"),
    (
        "我",
        "05:48",
        "这样的话，模块就不像现在这样，而是按具体的事情来划分。",
    ),
    (
        "老师",
        "05:53",
        "不是一个通用平台，是一个红塘村的平台。明白？要把这个虚拟的场景做起来，还要把茶马古道的历史做起来。",
    ),
    ("我", "06:07", "嗯。数字沙盘现在还没弄好。"),
    ("老师", "06:09", "准备。可能也只是先做了一个……"),
    ("我", "06:14", "这里面所有东西都是假的，都是演示内容。"),
    ("老师", "06:16", "〔此处听辨不清，可能谈及先放入部分内容。〕"),
    ("我", "06:20", "这个PPT做完了……"),
    ("老师", "06:24", "你看，也挺好的。"),
    ("我", "06:26", "这样我再提要求让它做嘛。"),
    (
        "老师",
        "06:28",
        "你现在做了一个东西出来，我就可以跟你提意见。你别老问我‘要做什么’，你们什么都没有，我怎么提？挺好，我觉得这个东西挺好。这个事情现在就按这个来。你看，我们就可以迭代，这是第一次迭代。开始做。我也理解了你的问题：你想做个通用的，实际上就是做红塘村的事。这个东西只能是这样，好不好？听明白了吧？",
    ),
    (
        "我",
        "07:10",
        "这样就清晰多了。像上面这些不相关的东西都不需要，这些乱七八糟的内容都改成红塘那几件具体事情就行了。",
    ),
    ("老师", "07:19", "对，围绕这些事情就行了。"),
    ("我", "07:21", "可以。"),
    (
        "老师",
        "07:25",
        "〔此处原音识别严重失真。两份转写均无法可靠还原，可能谈及页面定位、工程进展及Kimi生成内容，需回听原音确认。〕",
    ),
]


def set_style_font(style, east_asia, western, size, bold=False):
    style.font.name = western
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")


def set_run_font(run, east_asia=BODY_CN, western=BODY_WEST, size=None, bold=None, italic=None):
    run.font.name = western
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
set_style_font(normal, BODY_CN, BODY_WEST, 11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

title_style = styles["Title"]
set_style_font(title_style, HEADING_CN, BODY_WEST, 22, bold=True)
title_style.paragraph_format.space_before = Pt(0)
title_style.paragraph_format.space_after = Pt(6)

subtitle_style = styles["Subtitle"]
set_style_font(subtitle_style, BODY_CN, BODY_WEST, 10.5)
subtitle_style.paragraph_format.space_before = Pt(0)
subtitle_style.paragraph_format.space_after = Pt(12)

heading1 = styles["Heading 1"]
set_style_font(heading1, HEADING_CN, BODY_WEST, 16, bold=True)
heading1.paragraph_format.space_before = Pt(18)
heading1.paragraph_format.space_after = Pt(10)
heading1.paragraph_format.keep_with_next = True

heading2 = styles["Heading 2"]
set_style_font(heading2, HEADING_CN, BODY_WEST, 13, bold=True)
heading2.paragraph_format.space_before = Pt(14)
heading2.paragraph_format.space_after = Pt(7)
heading2.paragraph_format.keep_with_next = True

speaker_style = styles.add_style("Speaker", WD_STYLE_TYPE.PARAGRAPH)
speaker_style.base_style = normal
set_style_font(speaker_style, HEADING_CN, BODY_WEST, 11, bold=True)
speaker_style.paragraph_format.space_before = Pt(8)
speaker_style.paragraph_format.space_after = Pt(2)
speaker_style.paragraph_format.line_spacing = 1.0
speaker_style.paragraph_format.keep_with_next = True

note_style = styles.add_style("Transcript Note", WD_STYLE_TYPE.PARAGRAPH)
note_style.base_style = normal
set_style_font(note_style, BODY_CN, BODY_WEST, 9.5)
note_style.font.italic = True
note_style.paragraph_format.space_before = Pt(0)
note_style.paragraph_format.space_after = Pt(4)
note_style.paragraph_format.line_spacing = 1.15

header = section.header.paragraphs[0]
header.text = "红塘村可持续发展平台｜讨论录音稿（修订版）"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in header.runs:
    set_run_font(run, size=9)

footer = section.footer.paragraphs[0]
add_page_field(footer)

title = doc.add_paragraph("7月19日红塘村可持续发展平台讨论录音稿", style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
subtitle = doc.add_paragraph("修订版｜根据两份自动转写交叉校订", style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph("修订说明", style="Heading 1")
doc.add_paragraph(
    "本稿依据《0719讨论.docx》和《7月19日 新港西路_原文.docx》交叉校订。已修正高置信度的地名、术语、断句和说话人错误，并适度删除不影响原意的口语重复。"
)
note = doc.add_paragraph(style="Transcript Note")
note.add_run(
    "说明：方括号式标记“〔……〕”表示编辑说明或仍需回听确认的内容；除明显识别错误外，不对发言观点作实质改写。"
)

doc.add_paragraph("录音正文", style="Heading 1")
for speaker, timestamp, text in TRANSCRIPT:
    label = doc.add_paragraph(style="Speaker")
    label.add_run(f"{speaker}  {timestamp}")
    paragraph = doc.add_paragraph(text)
    if text.startswith("〔") and text.endswith("〕"):
        paragraph.style = styles["Transcript Note"]

doc.core_properties.title = "7月19日红塘村可持续发展平台讨论录音稿（修订版）"
doc.core_properties.subject = "红塘村可持续发展平台讨论录音交叉校订稿"
doc.core_properties.comments = "依据两份自动转写交叉校订；听辨不清处保留编辑标记。"

doc.save(OUTPUT)

# Structural and preset audit. Rendering is intentionally skipped by project instruction.
check = Document(OUTPUT)
assert check.paragraphs[0].style.name == "Title"
assert any(p.style.name == "Heading 1" and p.text == "录音正文" for p in check.paragraphs)
assert sum(1 for p in check.paragraphs if p.style.name == "Speaker") == len(TRANSCRIPT)
assert any("众包平台" in p.text for p in check.paragraphs)
assert any("只要给红塘村用" in p.text for p in check.paragraphs)
assert any("数字沙盘现在还没弄好" in p.text for p in check.paragraphs)
assert any("听辨不清" in p.text for p in check.paragraphs)
assert check.sections[0].page_width == Inches(8.5)
assert check.sections[0].left_margin == Inches(1.0)
assert check.styles["Normal"]._element.rPr.rFonts.get(qn("w:eastAsia")) == BODY_CN
assert check.styles["Normal"]._element.rPr.rFonts.get(qn("w:ascii")) == BODY_WEST
assert check.styles["Heading 1"]._element.rPr.rFonts.get(qn("w:eastAsia")) == HEADING_CN

print(f"Created: {OUTPUT}")
print(f"Transcript entries: {len(TRANSCRIPT)}")
print(f"Paragraphs: {len(check.paragraphs)} | sections: {len(check.sections)}")
