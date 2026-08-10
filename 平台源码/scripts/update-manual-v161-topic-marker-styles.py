from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.61 ｜ 更新日期：2026年8月10日",
)

marker_paragraph = next(p for p in document.paragraphs if p.text.startswith("专题地点：当前已接入35处小花园"))
base_text = marker_paragraph.text.split("显示规则：", 1)[0].rstrip()
replace_paragraph_text(
    marker_paragraph,
    base_text
    + " 显示规则：小花园、茶产业、村里用水、塌方与安全、历史与文化分别使用绿色、棕色、蓝色、砖红色和紫色，"
    + "专题点位保留识别图钉；公共服务、生态资源和村景记录等其他资料保持原有颜色，并统一使用低遮挡小圆点。2D和3D使用同一规则。",
)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.61"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.61",
    "2026年8月10日",
    "统一2D与3D地图要素样式：五个专题使用各自固定专题色和识别图钉；其他资料保持原色并统一改为小圆点。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)
    for run in cell.paragraphs[0].runs:
        set_body_font(run)

document.save(MANUAL)

reloaded = Document(MANUAL)
text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.61") for p in reloaded.paragraphs)
assert "其他资料保持原有颜色，并统一使用低遮挡小圆点" in text
assert "2D和3D使用同一规则" in text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.61"
print(MANUAL)
