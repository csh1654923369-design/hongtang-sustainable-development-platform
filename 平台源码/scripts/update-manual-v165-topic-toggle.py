from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        set_body_font(run)


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def remove_row_with_first_cell(table, value: str) -> None:
    for row in list(table.rows):
        if row.cells[0].text.strip() == value:
            table._tbl.remove(row._tr)


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.65 ｜ 更新日期：2026年8月10日",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("首页按照红塘村的具体事情组织为五个专题")),
    "首页按照红塘村的具体事情组织为五个专题：小花园、茶产业、村里用水、塌方与安全、历史与文化。公共服务和村景记录保留为其他资料，不再与五个主题并列。用户通过首页左上角专题卡片进入，2D与3D共用同一套专题状态和地图内容。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("桌面端和手机端都在地图左上角显示同一张白色圆角“专题”卡片")),
    "桌面端和手机端都在地图左上角显示同一张白色圆角“专题”卡片。卡片右侧的“展开”带倒三角图标，点击后打开列表；打开后同一按钮变为“收起”，点击即可关闭。列表直接显示小花园、茶产业、村里用水、塌方与安全、历史与文化五个专题，下面另列公共服务和村景记录等其他资料，不再重复显示“专题、全选、完成”标题行。各项可多选，每个专题右侧的“进入”用于切换到专题阅读状态；数字表示已接入的空间资料数量，没有已核实资料的专题显示“待调查”。切换2D与3D不会重置当前状态。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("当前基础地点数据包含")),
    "当前首页显示204个基础地点和586张现场照片；另载入7个待核实水系统节点、8条待核实线路和3个待核实供水片区，共形成222个地图要素。2D与3D现在优先读取同一个Supabase云端数据库：地点、供水专题和专题记录存放在数据表中，586张WebP照片存放在Storage素材桶中；本地JSON和本地照片继续保留为断网或云端临时不可用时的备用资料。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("专题地点：当前已接入35处小花园")),
    "专题地点：当前已接入35处小花园、9处茶厂和已有水设施资料；茶园、塌方安全点、茶马古道线路及历史地点仍需调查。公共服务和村景记录作为其他资料继续显示。点击已有地点后，2D和3D都会聚焦并在点位旁显示同一气泡详情。 显示规则：小花园、茶产业、村里用水、塌方与安全、历史与文化分别使用绿色、棕色、蓝色、砖红色和紫色，专题点位保留识别图钉；公共服务和村景记录等其他资料保持原有颜色，并统一使用低遮挡小圆点。2D和3D使用同一规则。",
)

replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("进入首页后点击左上角“专题”")),
    "进入首页后，点击左上角专题卡片右侧的“展开”，可以多选五个专题，也可以单独勾选其他资料；查看完列表后点击“收起”。点击某一专题右侧“进入”，地图只保留该专题相关内容，并显示该专题需要回答的问题、已有资料数量和后续调查重点；关闭专题后恢复进入前的勾选状态。",
)

point_table = next(table for table in document.tables if table.cell(0, 0).text == "地点类型")
remove_row_with_first_cell(point_table, "生态资源")
topic_table = document.tables[1]
remove_row_with_first_cell(topic_table, "生态资源")

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.65"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.65",
    "2026年8月10日",
    "专题卡片增加展开与收起状态，列表删除重复标题行；其他资料精简为公共服务和村景记录，并移除原测试点。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
table_text = "\n".join(cell.text for table in reloaded.tables for row in table.rows for cell in row.cells)
assert any(p.text.startswith("版本：V1.65") for p in reloaded.paragraphs)
assert "卡片右侧的“展开”带倒三角图标" in paragraph_text
assert "打开后同一按钮变为“收起”" in paragraph_text
assert "不再重复显示“专题、全选、完成”标题行" in paragraph_text
assert "生态资源" not in paragraph_text
assert "生态资源" not in table_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.65"
print(MANUAL)
