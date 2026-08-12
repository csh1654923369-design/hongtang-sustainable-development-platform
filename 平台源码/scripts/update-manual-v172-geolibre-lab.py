from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_two_character_indent(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), "200")


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_body_font(run)
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def insert_before(anchor, text: str, style: str):
    paragraph = anchor.insert_paragraph_before(style=style)
    run = paragraph.add_run(text)
    set_body_font(run)
    if style == "Normal":
        set_two_character_indent(paragraph)
    return paragraph


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_body_font(paragraph.add_run(text))


document = Document(MANUAL)
version_paragraph = next(p for p in document.paragraphs if p.text.startswith("版本：V1."))
replace_paragraph_text(version_paragraph, "版本：V1.72 ｜ 更新日期：2026年8月11日")

section_anchor = next(p for p in document.paragraphs if p.text == "7. 数据与隐私边界")
if not any(p.text == "6.4 GeoLibre空间数据实验室" for p in document.paragraphs):
    insert_before(section_anchor, "6.4 GeoLibre空间数据实验室", "Heading 2")
    insert_before(
        section_anchor,
        "平台新增独立的GeoLibre空间数据实验室。启动网站后，在浏览器地址栏输入http://localhost:3000/geolibre-lab即可进入；也可点击页内“新窗口打开”，获得更大的专业制图工作区。实验室不是面向村民的首页，而是供规划、调研和数据维护人员检查图层、属性和空间关系的试验工具。",
        "Normal",
    )
    insert_before(
        section_anchor,
        "实验室左侧图层面板当前包含8层：供水分区3个、供水线路8条、水系统节点7个、小花园35处、茶产业9处、村里用水设施2处、公共服务设施9处和村景记录149处。每层都可单独显示或隐藏、调整透明度、缩放至范围；点击对象可检查属性，右侧样式面板可试验颜色、描边和点符号，顶部绘图与编辑工具可用于本次项目中的试验修改。",
        "Normal",
    )
    insert_before(
        section_anchor,
        "GeoLibre与红塘平台读取Supabase中同一套公开空间数据。Supabase Edge Function会在每次打开时把platform_datasets转换为GeoLibre项目和GeoJSON图层，因此云端来源数据更新后，重新打开实验室即可看到新内容。当前桥接采用安全只读模式：GeoLibre内的修改只保存在本次项目或导出文件中，不会自动覆盖正式平台。正式写回功能应在后续增加管理员登录、草稿区、数据校验和审核发布后再启用。",
        "Normal",
    )
    insert_before(
        section_anchor,
        "技术实现包括src/app/geolibre-lab/page.tsx实验室页面、supabase/functions/geolibre-bridge/index.ts数据桥接，以及public/geolibre中的GeoLibre v2.5.0自托管网页构建。运行npm run test:geolibre可检查云端桥接、8个图层、底图请求和实验室实际渲染。",
        "Normal",
    )

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.72"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.72",
    "2026年8月11日",
    "新增GeoLibre空间数据实验室试验：自托管GeoLibre v2.5.0，通过Supabase只读桥接加载8个正式空间图层；支持图层检查、样式调整和试验编辑，当前修改不会自动覆盖正式平台。",
)
for cell, value in zip(new_row.cells, values):
    set_cell_text(cell, value)

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 按项目约定不进行LibreOffice渲染，只执行内容、样式、结构和压缩包完整性检查。
reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.72") for p in reloaded.paragraphs)
assert "6.4 GeoLibre空间数据实验室" in paragraph_text
assert "http://localhost:3000/geolibre-lab" in paragraph_text
assert "当前桥接采用安全只读模式" in paragraph_text
assert "npm run test:geolibre" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.72"
assert reloaded.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
assert reloaded.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
assert len(reloaded.element.body.xpath('.//w:ind[@w:firstLineChars="200"]')) >= 22
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(MANUAL)
print("version=V1.72")
print("zip=ok")
