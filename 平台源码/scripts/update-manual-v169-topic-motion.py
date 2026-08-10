from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.69 ｜ 更新日期：2026年8月10日",
)

topic_intro = next(
    p for p in document.paragraphs if p.text.startswith("桌面端和手机端都在地图左上角")
)
updated_topic_intro = topic_intro.text.replace(
    "点击右侧带倒三角的“展开”后，卡片向下平滑延展，专题内容同时以透明度渐变显现；打开后按钮变为“收起”，点击时执行反向过渡。",
    "点击右侧带倒三角的“展开”后，卡片以简短的缓入缓出动画自然展开，内容同步淡入；打开后按钮变为“收起”，收起时执行反向过渡。动画不使用位移或回弹效果。",
)
replace_paragraph_text(topic_intro, updated_topic_intro)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.69"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.69",
    "2026年8月10日",
    "简化专题卡片展开收起动画：改为短时缓入缓出，不再使用位移或回弹效果。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.69") for p in reloaded.paragraphs)
assert "卡片以简短的缓入缓出动画自然展开" in paragraph_text
assert "动画不使用位移或回弹效果" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.69"
print(MANUAL)
