from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.V1.36.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.35.backup.docx"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "黑体" if is_heading or index == 0 else "宋体")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "宋体")


def copy_cell_format(source_cell, target_cell):
    target_properties = target_cell._tc.get_or_add_tcPr()
    for child in list(target_properties):
        target_properties.remove(child)
    for child in source_cell._tc.tcPr:
        target_properties.append(deepcopy(child))


def add_version_row(table, values):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, value in enumerate(values):
        copy_cell_format(template_row.cells[index], row.cells[index])
        row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[index].text = value
        row.cells[index].paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        row.cells[index].paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def keep_table_rows_intact(table):
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        if not properties.xpath("./w:cantSplit"):
            properties.append(OxmlElement("w:cantSplit"))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    if not header_properties.xpath("./w:tblHeader"):
        header_properties.append(OxmlElement("w:tblHeader"))


QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.36 Demo" in document.paragraphs[5].text:
    print("Manual is already V1.36; no changes applied.")
    raise SystemExit(0)
if "V1.35 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.35 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.36 Demo　｜　更新日期：2026年7月30日　｜　访问地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n"
    "当前网站是可交互演示版本。首页提供“3D实景”和“2D地图”两种模式，打开后默认显示红塘村中心附近的"
    "三维视角，点击画面上方按钮即可切换。3D模式保留CesiumJS的三维相机、地点标记和操作方式，但红塘村"
    "高斯3D Tiles、DSM地形和无人机正射影像已经全部改为读取本机文件，不再调用Cesium ion，也不需要"
    "额度或访问令牌。3D模式载入56个POI和149个村景圆点；二维和三维地图中的POI图钉都使用与地点含义"
    "对应的矢量符号，村景使用小圆点。点击地点后可查看气泡详情和现场照片。2D模式直接复用“村里一张图”"
    "的手绘图、无人机影像、点位筛选和详情功能；切换时只运行当前模式。问题、项目、微行动、互助资源和"
    "指标等互动业务仍为演示数据。"
)

find_paragraph(
    document,
    "以后不必每次输入命令。项目根目录保留“启动网站.bat”和“关闭网站.bat”。网页源码集中放在“平台源码”文件夹中，3D模型、无人机原始影像、真实POI、村景照片清单和地图服务原始资料集中放在“平台素材”文件夹中。正式三维首页需要联网，并要求“平台源码/.env.local”中已经配置可用的 CESIUM_ION_TOKEN。",
).text = (
    "以后不必每次输入命令。项目根目录保留“启动网站.bat”和“关闭网站.bat”。网页源码集中放在"
    "“平台源码”文件夹中，3D模型、无人机原始影像、POI、村景照片清单和地图服务原始资料集中放在"
    "“平台素材”文件夹中。正式三维首页不需要Cesium ion额度或访问令牌；首次打开仍需联网读取"
    "CesiumJS程序库。http://localhost:3000和http://127.0.0.1:3000均可使用。"
)

find_paragraph(
    document,
    "数据与隐私说明\n建筑底图只保留建筑编号、轮廓、新旧类型、高度、估算占地和中心坐标。原始WFS中的姓名、电话、住址、家庭成员、人口等字段不会进入网站。真实地点文件只使用素材中的名称、分类、简介、坐标、照片地址和资料时间。Cesium主令牌只在服务端环境文件中读取，网页只取得模型、地形和影像的资源端点。",
).text = (
    "数据与隐私说明\n建筑底图只保留建筑编号、轮廓、新旧类型、高度、估算占地和中心坐标。原始WFS中的"
    "姓名、电话、住址、家庭成员、人口等字段不会进入网站。地点文件只使用素材中的名称、分类、简介、"
    "坐标、照片地址和资料时间。3D模型由同源的受限接口按需读取指定JSON和GLB分块，不会把素材目录作为"
    "普通公开文件夹暴露。"
)

find_paragraph(
    document,
    "src/app/api/cesium-ion/asset/route.ts 只在服务端读取 CESIUM_ION_TOKEN，分别获取模型、地形和影像端点。主令牌不会写入HTML。正式三维首页需要联网；资源不可用时页面显示失败提示，用户仍可进入“村庄总览”。",
).text = (
    "src/app/api/local-cesium-model/[...segments]/route.ts从“平台素材/3D高斯展示/Cesium本地三维瓦片”"
    "读取tileset.json和GLB分块，支持Range与浏览器缓存；可选环境变量LOCAL_CESIUM_MODEL_DIR用于修改"
    "模型目录。public/cesium-viewer/index.html读取本地DSM高程网格和本地无人机正射影像，不请求"
    "Cesium World Terrain、Bing卫星影像或Cesium ion资产。"
)
find_paragraph(
    document,
    "SparkJS 2.1.0 的RAD/RADC轻量化查看器、本地vendor模块和Range接口继续保留，供离线预览、性能对比或未来混合加载使用，但不再作为正式首页入口。正式首页当前以Cesium承载地形、地理坐标和高斯3D Tiles。",
).text = (
    "SparkJS 2.1.0的RAD/RADC轻量化查看器、本地vendor模块和Range接口继续保留，供离线预览、性能对比"
    "或未来混合加载使用，但不再作为正式首页入口。正式首页当前使用CesiumJS承载本地DSM地形、地理坐标、"
    "本地正射影像和本地高斯3D Tiles；既保留原有功能，也不消耗在线资产额度。"
)
find_paragraph(
    document,
    "“平台素材/Production_1-tif”文件夹保存无人机原始正射影像、DSM及其辅助文件。这些文件体积较大，网站运行时使用已经生成到“平台源码/public”中的网页版本；只有重新生成网页影像时才会读取原始素材。",
).text = (
    "“平台素材/Production_1-tif”文件夹保存无人机原始正射影像、DSM及其辅助文件。这些文件体积较大。"
    "网站运行时使用public/data中的0.3米WebP影像和约1MB的513×513浮点高程网格；只有重新生成网页影像"
    "或地形时才读取原始TIF。"
)
find_paragraph(
    document,
    "scripts/prepare-orthophoto.mjs 从 0.3 米重采样 TIF 生成约 3.7MB 的透明 WebP，并写入坐标系、像素大小、范围和校准状态元数据。",
).text = (
    "scripts/prepare-orthophoto.mjs从0.3米重采样TIF生成约3.7MB的WebP及坐标元数据；"
    "scripts/prepare-local-terrain.py从DSM生成513×513的Float32地形网格，使用邻域低分位采样和平滑"
    "尽量削弱房屋、树木等地表物对基础地形的抬高影响。"
)
find_paragraph(
    document,
    "网页读取 public/data/hongtang-orthophoto-0.3m.webp；配套元数据为 public/data/hongtang-orthophoto.json，坐标系是 EPSG:32647。",
).text = (
    "网页读取public/data/hongtang-orthophoto-0.3m.webp和hongtang-orthophoto.json；三维地形读取"
    "hongtang-terrain-513.f32和hongtang-terrain.json。两组原始资料坐标系均为EPSG:32647。"
)

find_paragraph(
    document,
    "3D高斯实景公开或部署到互联网前，应确认航拍成果、Cesium ion资产和Spark轻量化模型的公开授权。CESIUM_ION_TOKEN必须保存在服务端环境文件中，并按需要设置资源范围、访问控制和用量限制；不要把恢复出的PLY、全部RADC文件或主令牌直接提交到公开Git仓库。",
).text = (
    "3D高斯实景公开或部署到互联网前，应确认航拍成果、本地3D Tiles和Spark轻量化模型的公开授权。"
    "不要把原始PLY、全部GLB或RADC大文件直接提交到公开Git仓库；部署时应使用经过授权的对象存储、"
    "静态资源服务器或受限接口，并设置访问控制。"
)
find_paragraph(
    document,
    "npm run prepare:orthophoto：从外层 0.3 米 TIF 重新生成网页正射影像和元数据。",
).text = (
    "npm run prepare:orthophoto：从外层0.3米TIF重新生成网页正射影像和元数据；"
    "npm run prepare:local-terrain：从外层DSM重新生成三维地形网格和元数据。"
)
find_paragraph(
    document,
    "先点击三维页面右下角齿轮，在“画面质量”中选择“省电”；操作过快或过慢时，可把平移、旋转或缩放灵敏度在0.3—2.0之间单独调整，1.0为默认值。需要减少地点密度时，使用左上角类型筛选。关闭其他占用显卡的页面，并确认网络稳定。Cesium地形、卫星影像和高斯模型都需要联网；若一直加载失败，可重新打开平台或先进入“村庄总览”。",
).text = (
    "先点击三维页面右下角齿轮，在“画面质量”中选择“省电”；操作过快或过慢时，可把平移、旋转或缩放"
    "灵敏度在0.3—2.0之间单独调整，1.0为默认值。需要减少地点密度时，使用左上角类型筛选。关闭其他"
    "占用显卡的页面。本地地形、正射影像和高斯模型不受网络额度影响；若画面一直加载失败，可重新启动"
    "网站并确认“平台素材/3D高斯展示/Cesium本地三维瓦片”完整。首次读取CesiumJS程序库仍需要联网。"
)

implementation_table = document.tables[2]
for row in implementation_table.rows:
    if row.cells[0].text.strip() == "首页地图":
        row.cells[1].text = "HomeExperience + MapExplorer + CesiumJS 1.143（本地资源）"
        row.cells[2].text = (
            "默认3D实景；一键切换完整2D地图；3D读取本地DSM、正射影像和高斯3D Tiles，"
            "不调用Cesium ion；二维与三维共用语义图标，卸载未选地图"
        )

faq_table = document.tables[3]
for row in faq_table.rows:
    if row.cells[0].text.strip() == "3D首页一直加载、图钉不显示或显示失败":
        row.cells[1].text = "本地模型目录缺失、启动服务未重启，或首次读取CesiumJS时网络不可用"
        row.cells[2].text = (
            "确认“平台素材/3D高斯展示/Cesium本地三维瓦片/tileset.json”存在；"
            "双击“关闭网站.bat”后重新启动；首次打开时检查普通网络"
        )

add_version_row(
    document.tables[4],
    [
        "V1.36 Demo",
        "2026-07-30",
        "3D首页改为CesiumJS加载本地DSM地形、本地无人机正射影像和本地高斯3D Tiles；"
        "取消Cesium ion额度与访问令牌依赖，保留205个地点、气泡详情、分类筛选、鼠标操作和2D/3D切换。",
    ],
)

for table in document.tables:
    keep_table_rows_intact(table)
normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 14, 38]
assert "V1.36 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
all_table_text = "\n".join(
    cell.text for table in check.tables for row in table.rows for cell in row.cells
)
assert "不再调用Cesium ion" in all_text
assert "local-cesium-model" in all_text
assert "hongtang-terrain-513.f32" in all_text
assert "CESIUM_ION_TOKEN" not in all_text
assert "V1.36 Demo" in all_table_text
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count
assert check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
assert check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.36: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
