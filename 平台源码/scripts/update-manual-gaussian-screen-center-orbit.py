from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
SCREENSHOT = QA_ROOT / "home-1440.png"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v118.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.17.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


def replace_home_image(document):
    caption = find_paragraph_prefix(document, "图 1")
    image_paragraph = image_paragraph_before(document, caption)
    clear_paragraph(image_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.3))
    shape._inline.docPr.set("name", "图 1 3D高斯实景首页与五类示例点")
    shape._inline.docPr.set(
        "descr",
        "红塘村3D高斯实景首页，右键拖动时围绕当前屏幕画面中心旋转",
    )


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        copy_cell_format(template_row.cells[index], cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if template_row.cells[index].paragraphs and cell.paragraphs:
            cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
            cell.paragraphs[0].paragraph_format.alignment = (
                template_row.cells[index].paragraphs[0].paragraph_format.alignment
            )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_first_line_chars(paragraph, chars=200):
    properties = paragraph._p.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    for attribute in ("firstLine", "hanging", "hangingChars"):
        key = qn(f"w:{attribute}")
        if key in indent.attrib:
            del indent.attrib[key]
    indent.set(qn("w:firstLineChars"), str(chars))


def normalize_document(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")

    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
        if (
            index >= 8
            and paragraph.style
            and paragraph.style.name == "Normal"
            and paragraph.text.strip()
            and "\n" not in paragraph.text
        ):
            set_first_line_chars(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not SCREENSHOT.exists():
    raise FileNotFoundError(SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.17 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.17 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.18 Demo   |   更新日期：2026年7月27日   |   适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页使用红塘村在线3D高斯实景，并展示小花园、茶场、茶厂、"
    "村里用水和光伏设施5个可点击示例点。左键拖动沿村庄水平面平移；右键从画面任何位置按下时，"
    "均以按下瞬间的屏幕画面正中心为固定点旋转，而不是以鼠标所在点旋转；滚轮按固定比例缩放；"
    "中键拖动已停用。点击图钉只打开事项详情；点击“定位到此处”才会让镜头飞近，关闭按钮或按Esc可取消选择。"
    "所有三维图钉均为演示位置，待实地核实，不能作为真实设施、风险或产权判断依据。"
    "原无人机影像首页完整保留为“村庄总览”；3D首页需要保持联网。"
)

operation = find_paragraph_prefix(document, "3D首页操作：")
operation.text = (
    "3D首页操作：按住左键拖动，可沿村庄所在地表的水平面平移；按住右键拖动，可围绕右键按下瞬间的"
    "屏幕画面正中心旋转，无论鼠标从画面哪里按下，中心点都取屏幕正中心；滚轮每次按固定比例放大或缩小；"
    "中键拖动已停用。点击小花园、茶场、茶厂、村里用水或光伏设施图钉，会在桌面右侧或手机底部打开"
    "与“村里一张图”一致的事项详情，但不会自动移动镜头；只有点击“定位到此处”才会平滑靠近。"
    "关闭按钮或Esc会取消高亮。点击“操作设置”后，左键平移和滚轮缩放灵敏度可分别调节、分别保存；"
    "右键旋转使用固定速度。还可使用“回到模型”和“全屏查看”。"
)

technical = find_paragraph_prefix(document, "平台将Cesium相机输入重新映射")
technical.text = (
    "平台将Cesium默认鼠标相机输入替换为左键水平平移、右键屏幕中心轨道旋转和滚轮定比缩放，并移除中键拖动、"
    "默认倾斜和观察输入。右键按下时，先使用Scene.pickPosition从深度缓冲拾取画布正中心的三维坐标；"
    "若正中心暂时没有可拾取模型，则使用Camera.getPickRay生成屏幕中心视线，并与村庄水平面求交。"
    "拖动期间该三维中心保持固定，平台使用四元数更新相机偏移，并让相机方向持续指向该点，因此旋转中心始终"
    "保持在屏幕正中心。左键平移继续使用0.5、0.72、1、1.28和1.6倍五档倍率；滚轮每120毫秒最多执行一次"
    "定比缩放，五档每步比例为6%、10%、15%、22%和32%，默认15%。左键平移与滚轮缩放设置分别写入浏览器本地存储。"
)

for row in document.tables[1].rows:
    if row.cells[0].text.strip() == "首页":
        row.cells[2].text = (
            "全屏红塘村3D高斯实景；5个分类示例图钉可打开事项详情和显式定位；"
            "左键水平平移、右键围绕按下时的屏幕画面正中心旋转、滚轮定比缩放，中键拖动停用"
        )
        break
else:
    raise ValueError("Homepage route row not found")

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载红塘村高斯模型；显示5个可点击分类锚点；二维事项数据复用、显式定位、Esc清除和移动端底部详情；"
            "右键按下时通过深度拾取或中心视线求交确定屏幕中心轨道点；左键水平平移、滚轮定比缩放，中键操作停用"
        )
        break
else:
    raise ValueError("3D technology row not found")

version_table = document.tables[4]
version_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    version_row.cells,
    [
        "V1.18 Demo",
        "2026-07-27",
        "3D首页右键旋转改为以按下瞬间的屏幕画面正中心为固定轨道点；新增深度拾取、中心视线后备求交和中心误差验证。",
    ],
):
    cell.text = value

replace_home_image(document)
normalize_document(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 20]
assert "V1.18 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "屏幕画面正中心为固定点旋转" in all_text
assert "而不是以鼠标所在点旋转" in all_text
assert "Scene.pickPosition" in all_text
assert any(row.cells[0].text.strip() == "V1.18 Demo" for row in check.tables[4].rows)
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count

normal_fonts = check.styles["Normal"].element.rPr.rFonts
heading_fonts = check.styles["Heading 1"].element.rPr.rFonts
assert normal_fonts.get(qn("w:eastAsia")) == "SimSun"
assert normal_fonts.get(qn("w:ascii")) == "Times New Roman"
assert heading_fonts.get(qn("w:eastAsia")) == "SimHei"

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.18: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
