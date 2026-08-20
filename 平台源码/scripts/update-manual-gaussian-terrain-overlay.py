from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
SCREENSHOT = QA_ROOT / "gaussian-terrain-overlay-1440.png"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v121.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.20.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def replace_home_image(document):
    caption = find_paragraph_prefix(document, "图 1")
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    image_paragraph = next(
        paragraph for paragraph in reversed(paragraphs[:index]) if paragraph._p.xpath(".//w:drawing")
    )
    clear_paragraph(image_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.3))
    shape._inline.docPr.set("name", "图 1 地形环境中的3D高斯实景首页")
    shape._inline.docPr.set(
        "descr",
        "红塘村高斯模型按真实坐标叠加在周边三维地形、卫星影像和无人机正射影像中",
    )
    caption.text = "图 1  按真实坐标嵌入周边地形的3D高斯实景首页"


def add_version_row(table):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        properties = cell._tc.get_or_add_tcPr()
        for child in list(properties):
            properties.remove(child)
        for child in template_row.cells[index]._tc.tcPr:
            properties.append(deepcopy(child))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        cell.paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not SCREENSHOT.exists():
    raise FileNotFoundError(SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.20 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.20 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.21 Demo   |   更新日期：2026年7月27日   |   适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页把红塘村在线3D高斯实景按真实地理坐标嵌入周边三维地形，"
    "远处显示全球卫星影像，村庄附近叠加红塘无人机正射影像。高斯模型不再处于黑色空场景中，"
    "可以结合周边山地、道路和村庄环境辨认位置。首页还展示小花园、茶场、茶厂、村里用水和光伏设施"
    "5个可点击示例点。左键拖动沿村庄水平面平移；右键可从画面任意位置按下，左右和上下拖动均围绕"
    "屏幕画面正中心对应的三维点运动；向上接近水平，向下接近垂直俯视。滚轮按固定比例缩放，"
    "中键拖动已停用。点击图钉只打开事项详情，点击“定位到此处”才会让镜头飞近。"
    "所有三维图钉均为演示位置，待实地核实。"
)

operation = find_paragraph_prefix(document, "3D首页操作：")
operation.text = (
    "3D首页操作：打开首页后，平台会依次加载周边三维地形、远景卫星影像、红塘无人机正射影像和高斯模型。"
    "高斯模型外侧显示连续山地环境，不再显示黑色背景。按住左键拖动，可沿村庄所在地表的水平面平移。"
    "右键可从画面任意位置按下；开始拖动后，平台始终以屏幕正中心对应的三维点作为轨道中心。"
    "右键左右拖动时，相机围绕中心点做水平环绕；向上拖动时视线接近水平；向下拖动时接近垂直俯视。"
    "滚轮每次按固定比例放大或缩小，中键拖动已停用。点击事项图钉只打开详情，"
    "点击“定位到此处”才移动镜头。在线地形或远景影像临时加载失败时，页面会保留高斯模型和本地无人机影像。"
)

technical = find_paragraph_prefix(document, "平台将Cesium默认鼠标相机输入替换")
technical.text = (
    "平台参照旧村庄规划互动平台的Cesium实现，恢复地球表面并接入Cesium World Terrain和"
    "Cesium World Imagery；红塘0.3米无人机正射影像按WGS84范围"
    "99.902144°E—99.912024°E、24.631730°N—24.641417°N铺到地形表面。"
    "高斯资产根变换的原点为99.908740607°E、24.636255278°N、1764m，位于无人机影像范围内，"
    "因此模型、地形和影像使用同一地理坐标空间，无需按屏幕位置人工缩放。服务端接口分别取得高斯资产、"
    "全球地形和全球影像的短期访问配置，长期Cesium令牌不会写入公开HTML。在线地形失败时退回椭球地面，"
    "全球影像失败时仍保留本地正射影像。相机继续使用左键水平平移、右键屏幕中心全方向轨道旋转、"
    "滚轮定比缩放，并移除中键拖动；自动测试同时校验三类环境图层、模型地理原点和原有相机操作。"
)

for row in document.tables[1].rows:
    if row.cells[0].text.strip() == "首页":
        row.cells[2].text = (
            "红塘村3D高斯实景按真实坐标叠加周边三维地形、全球卫星影像和无人机正射影像；"
            "5个分类示例图钉；左键水平平移；右键围绕屏幕中心旋转；滚轮定比缩放"
        )
        break
else:
    raise ValueError("Homepage route row not found")

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载高斯模型、Cesium全球地形和全球影像；按真实经纬度铺设红塘无人机正射影像；"
            "高斯资产地理原点为99.908740607°E、24.636255278°N、1764m；"
            "保留5个可点击分类锚点和自定义鼠标操作"
        )
        break
else:
    raise ValueError("3D technology row not found")

version_row = add_version_row(document.tables[4])
for cell, value in zip(
    version_row.cells,
    [
        "V1.21 Demo",
        "2026-07-27",
        "把高斯模型按真实坐标嵌入周边三维地形，接入全球卫星影像和红塘无人机正射影像，消除模型外的黑色空场景。",
    ],
):
    cell.text = value

replace_home_image(document)
normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 23]
assert "V1.21 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "高斯模型不再处于黑色空场景中" in all_text
assert "Cesium World Terrain" in all_text
assert "99.908740607°E、24.636255278°N、1764m" in all_text
assert any(row.cells[0].text.strip() == "V1.21 Demo" for row in check.tables[4].rows)
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count
assert check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
assert check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.21: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
