from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v123.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.22.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def add_version_row(table):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        properties = cell._tc.get_or_add_tcPr()
        for child in list(properties):
            properties.remove(child)
        for child in template_row.cells[index]._tc.tcPr:
            properties.append(deepcopy(child))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        cell.paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.22 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.22 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.23 Demo   |   更新日期：2026年7月28日   |   适用地址：http://localhost:3000"
)

operation = find_paragraph_prefix(document, "3D首页操作：")
operation.text = (
    "3D首页操作：打开首页后，平台会依次加载周边三维地形、远景卫星影像、红塘无人机正射影像和高斯模型。"
    "高斯模型外侧显示连续山地环境，不再显示黑色背景。为避免世界地形在局部略高时把有效的近地高斯点遮住，"
    "平台已关闭“地形遮挡高斯实景”：地形仍用于展示周边山势、承载影像和提供高程采样，但不会再把高斯模型"
    "切出扁平影像块。平台会等待高斯模型表面加载完成后再放置5个事项图钉；图钉优先贴合高斯实景表面，同时"
    "检查当地三维地形高度，若拾取位置低于地形则自动抬升到地表之上，避免模型边缘的图钉钻入地下或跳到错误高度。"
    "按住左键拖动，可沿村庄所在地表的水平面平移。右键可从画面任意位置按下；开始拖动后，平台始终以屏幕正中心"
    "对应的三维点作为轨道中心。右键左右拖动时，相机围绕中心点做水平环绕；向上拖动时视线接近水平；向下拖动时"
    "接近垂直俯视。滚轮每次按固定比例放大或缩小，中键拖动已停用。点击事项图钉只打开详情，点击“定位到此处”"
    "才移动镜头。在线地形或远景影像临时加载失败时，页面会保留高斯模型和本地无人机影像。"
)

technical = find_paragraph_prefix(document, "平台参照旧村庄规划互动平台的Cesium实现")
technical.text = (
    "平台参照旧村庄规划互动平台的Cesium实现，恢复地球表面并接入Cesium World Terrain和"
    "Cesium World Imagery；红塘0.3米无人机正射影像按真实WGS84范围铺到地形表面。高斯资产根变换的原点为"
    "99.908740607°E、24.636255278°N、1764m，模型、地形和影像使用同一地理坐标空间。排查扁平影像块时，"
    "在同一相机视角下分别开启和关闭globe.depthTestAgainstTerrain进行对照，确认世界地形的局部高程或"
    "细节层级会遮挡有效的近地高斯点。查看器现将depthTestAgainstTerrain设为false，使世界地形继续提供"
    "周边环境、影像承载和sampleTerrainMostDetailed高程采样，但不参与遮挡最上层高斯实景。三维图钉在"
    "initialTilesLoaded事件之后沿局部竖直方向调用pickFromRayMostDetailed拾取高斯表面；随后调用"
    "sampleTerrainMostDetailed取得同一经纬度的地形高度。最终锚点高度取高斯表面与地形表面的较高值，"
    "再增加2.5米至6米的图钉抬升量。自动测试会同时校验地形遮挡状态、5个图钉的位置和相机操作。"
)

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载高斯模型、Cesium全球地形和全球影像；按真实经纬度铺设红塘无人机正射影像；"
            "关闭地形对高斯实景的深度遮挡，避免局部露出扁平正射影像，同时保留周边山势、影像承载与高程采样；"
            "5个事项图钉在模型瓦片加载后贴合高斯表面，并以当地地形高度作为最低高度保护；"
            "保留屏幕中心旋转、水平平移和定比缩放"
        )
        break
else:
    raise ValueError("3D technology row not found")

version_row = add_version_row(document.tables[4])
for cell, value in zip(
    version_row.cells,
    [
        "V1.23 Demo",
        "2026-07-28",
        "修复高斯模型与世界地形叠加后局部露出扁平影像块：保留三维地形环境与高程采样，但不再让地形遮挡高斯实景。",
    ],
):
    cell.text = value

normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 25]
assert "V1.23 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "不会再把高斯模型切出扁平影像块" in all_text
assert "depthTestAgainstTerrain设为false" in all_text
assert "pickFromRayMostDetailed" in all_text
assert "sampleTerrainMostDetailed" in all_text
assert any(row.cells[0].text.strip() == "V1.23 Demo" for row in check.tables[4].rows)
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count
assert check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
assert check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.23: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
