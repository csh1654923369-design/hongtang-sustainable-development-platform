from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
SCREENSHOT = QA_ROOT / "home-1440.png"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v120.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.19.backup.docx"


def find_paragraph_prefix(document, prefix):
    compact_prefix = "".join(prefix.split())
    for paragraph in document.paragraphs:
        if "".join(paragraph.text.split()).startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def replace_home_image(document):
    caption = find_paragraph_prefix(document, "图 1")
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    image_paragraph = next(
        paragraph for paragraph in reversed(paragraphs[:index]) if paragraph._p.xpath(".//w:drawing")
    )
    clear_paragraph(image_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.3))
    shape._inline.docPr.set("name", "图 1 3D高斯实景首页与五类示例点")
    shape._inline.docPr.set(
        "descr",
        "红塘村3D高斯实景首页，右键可从任意位置拖动并始终围绕屏幕中心旋转",
    )


def add_version_row(table):
    template_row = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        properties = cell._tc.get_or_add_tcPr()
        for child in list(properties):
            properties.remove(child)
        for child in template_row.cells[index]._tc.tcPr:
            properties.append(deepcopy(child))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
        cell.paragraphs[0].paragraph_format.alignment = (
            template_row.cells[index].paragraphs[0].paragraph_format.alignment
        )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def normalize_fonts(document):
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not SCREENSHOT.exists():
    raise FileNotFoundError(SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.19 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.19 manual, got: {document.paragraphs[5].text}")
if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.20 Demo   |   更新日期：2026年7月27日   |   适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页使用红塘村在线3D高斯实景，并展示小花园、茶场、茶厂、"
    "村里用水和光伏设施5个可点击示例点。左键拖动沿村庄水平面平移。右键可以从画面任意位置按下，"
    "鼠标位置只决定是否开始拖动，不决定旋转中心；左右和上下拖动都围绕屏幕画面正中心对应的三维点运动。"
    "向上拖动时，相机沿中心点轨道降低俯视角，视线越来越接近水平；向下拖动时，相机沿轨道增大俯视角，"
    "视线越来越接近垂直俯视。滚轮按固定比例缩放，中键拖动已停用。点击图钉只打开事项详情；"
    "点击“定位到此处”才会让镜头飞近。所有三维图钉均为演示位置，待实地核实。"
)

operation = find_paragraph_prefix(document, "3D首页操作：")
operation.text = (
    "3D首页操作：按住左键拖动，可沿村庄所在地表的水平面平移。右键可从画面任意位置按下；开始拖动后，"
    "平台始终以屏幕正中心对应的三维点作为轨道中心，而不是以鼠标所在点作为中心。右键左右拖动时，相机围绕"
    "中心点做水平环绕；向上拖动时，相机沿同一中心点轨道降低俯视角，接近水平视线；向下拖动时，相机沿轨道"
    "增大俯视角，接近垂直俯视。上下拖动会改变相机在轨道上的位置，但中心点始终留在屏幕正中。"
    "滚轮每次按固定比例放大或缩小，中键拖动已停用。点击事项图钉只打开详情，点击“定位到此处”才移动镜头。"
)

technical = find_paragraph_prefix(document, "平台将Cesium默认鼠标相机输入替换")
technical.text = (
    "平台将Cesium默认鼠标相机输入替换为左键水平平移、右键屏幕中心全方向轨道旋转、滚轮定比缩放，并移除中键拖动。"
    "右键按下位置不参与中心计算；平台始终使用Scene.pickPosition或屏幕中心视线与村庄水平面的交点取得轨道中心。"
    "水平位移绕该点的局部竖直轴旋转orbitOffset；垂直位移根据鼠标方向调整orbitOffset的仰角，向上减小、向下增大，"
    "范围限制为约3°至88°。每次更新后，相机位置由轨道中心与orbitOffset相加得到，相机方向重新指向轨道中心，"
    "从而保证左右和上下拖动后的中心屏幕误差均不超过1像素。自动测试从画面非中心位置开始拖动，并校验相机位置"
    "随垂直轨道改变、俯视角方向正确且轨道中心仍保持在屏幕正中。"
)

for row in document.tables[1].rows:
    if row.cells[0].text.strip() == "首页":
        row.cells[2].text = (
            "全屏红塘村3D高斯实景；5个分类示例图钉；左键水平平移；右键可从任意位置拖动，"
            "左右和上下均围绕屏幕中心旋转，向上接近水平、向下接近垂直；滚轮定比缩放"
        )
        break
else:
    raise ValueError("Homepage route row not found")

for row in document.tables[2].rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载红塘村高斯模型；显示5个可点击分类锚点；右键按下位置不限制，水平和垂直位移"
            "均围绕屏幕中心轨道点；上移减小俯视角、下移增大俯视角；左键平移、滚轮缩放，中键停用"
        )
        break
else:
    raise ValueError("3D technology row not found")

version_row = add_version_row(document.tables[4])
for cell, value in zip(
    version_row.cells,
    [
        "V1.20 Demo",
        "2026-07-27",
        "明确右键判定区域与旋转中心相互独立：右键可在任意位置按下，左右和上下均围绕屏幕中心轨道点；向上接近水平，向下接近垂直。",
    ],
):
    cell.text = value

replace_home_image(document)
normalize_fonts(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert [len(table.rows) for table in check.tables] == [5, 22, 9, 13, 22]
assert "V1.20 Demo" in check.paragraphs[5].text
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
assert "鼠标位置只决定是否开始拖动，不决定旋转中心" in all_text
assert "左右和上下拖动都围绕屏幕画面正中心" in all_text
assert "上下拖动会改变相机在轨道上的位置" in all_text
assert any(row.cells[0].text.strip() == "V1.20 Demo" for row in check.tables[4].rows)
first_line_count = len(check.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))
assert first_line_count >= 18, first_line_count
assert check.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimSun"
assert check.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "SimHei"

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.20: {MANUAL}")
print(f"Native two-character first-line indents: {first_line_count}")
