from copy import deepcopy
from pathlib import Path
import os
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


SOURCE_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = SOURCE_ROOT.parent
QA_ROOT = SOURCE_ROOT / ".qa"
MANUAL = WORKSPACE_ROOT / "红塘村可持续发展平台使用手册.docx"
OUTPUT = QA_ROOT / "红塘村可持续发展平台使用手册.v113.docx"
BACKUP = QA_ROOT / "红塘村可持续发展平台使用手册.V1.12.backup.docx"
HOME_SCREENSHOT = QA_ROOT / "home-1440.png"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr"):
            paragraph._p.remove(child)


def image_paragraph_before(document, caption):
    paragraphs = document.paragraphs
    index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for paragraph in reversed(paragraphs[:index]):
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
    raise ValueError("Image paragraph not found before caption")


def set_image(paragraph, image_path):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(6.3))
    shape._inline.docPr.set("name", "图 1 3D高斯实景首页")
    shape._inline.docPr.set("descr", "红塘村3D高斯实景铺满首页，滚轮采用缓慢缩放")


def copy_cell_format(source_cell, target_cell):
    properties = target_cell._tc.get_or_add_tcPr()
    for child in list(properties):
        properties.remove(child)
    for child in source_cell._tc.tcPr:
        properties.append(deepcopy(child))


def add_row_like(table, template_row):
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        copy_cell_format(template_row.cells[index], cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if template_row.cells[index].paragraphs and cell.paragraphs:
            cell.paragraphs[0].style = template_row.cells[index].paragraphs[0].style
            cell.paragraphs[0].paragraph_format.alignment = (
                template_row.cells[index].paragraphs[0].paragraph_format.alignment
            )
    return row


def set_run_font(run, east_asia, western="Times New Roman"):
    run.font.name = western
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western="Times New Roman"):
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_first_line_chars(paragraph, chars=200):
    properties = paragraph._p.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    for attribute in ("firstLine", "hanging", "hangingChars"):
        key = qn(f"w:{attribute}")
        if key in indent.attrib:
            del indent.attrib[key]
    indent.set(qn("w:firstLineChars"), str(chars))


def normalize_document(document):
    set_style_font(document.styles["Normal"], "SimSun")
    for level in range(1, 10):
        style_name = f"Heading {level}"
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "SimHei", "SimHei")
    for index, paragraph in enumerate(document.paragraphs):
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            set_run_font(run, "SimHei" if is_heading or index == 0 else "SimSun")
        if (
            index >= 8
            and paragraph.style
            and paragraph.style.name == "Normal"
            and paragraph.text.strip()
            and "\n" not in paragraph.text
        ):
            set_first_line_chars(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "SimSun")


if not HOME_SCREENSHOT.exists():
    raise FileNotFoundError(HOME_SCREENSHOT)

QA_ROOT.mkdir(parents=True, exist_ok=True)
document = Document(MANUAL)
if "V1.12 Demo" not in document.paragraphs[5].text:
    raise ValueError(f"Expected V1.12 manual, got: {document.paragraphs[5].text}")

if not BACKUP.exists():
    shutil.copy2(MANUAL, BACKUP)

document.paragraphs[5].text = (
    "版本：V1.13 Demo   |   更新日期：2026年7月23日   |   适用地址：http://localhost:3000"
)
document.paragraphs[6].text = (
    "阅读提示\n当前网站是可交互演示版本。首页使用红塘村在线3D高斯实景，"
    "可以旋转、缓慢缩放和平移；原无人机影像首页完整保留为“村庄总览”。"
    "3D首页需要联网，业务点位仍未按影像校准真实位置，不能作为真实设施、风险或产权判断依据。"
)

caption = find_paragraph(document, "图 1  3D高斯实景首页")
set_image(image_paragraph_before(document, caption), HOME_SCREENSHOT)

find_paragraph(
    document,
    "桌面端顶栏：平台名称、八个 TAB 和常用操作位于同一行；“首页”打开3D高斯实景，“村庄总览”打开原无人机影像首页，具体事项仍先进入各自的地图首屏。",
).text = (
    "桌面端顶栏：平台名称、八个TAB和常用操作位于同一行；“首页”打开3D高斯实景，"
    "“村庄总览”打开原无人机影像首页，具体事项仍先进入各自的地图首屏。"
)
find_paragraph(
    document,
    "小花园、茶厂、村里用水：进入后首先使用同一张村庄地图，并分别只显示小花园、茶场与茶厂、用水设施点位。",
).insert_paragraph_before(
    "3D首页操作：按住鼠标左键拖动旋转，滚轮缓慢缩放，按住右键拖动平移；可使用“回到模型”和“全屏查看”恢复视角或扩大画面。",
    style="List Bullet",
)

page_table = document.tables[1]
for row in page_table.rows:
    if row.cells[0].text.strip() == "首页":
        row.cells[2].text = "全屏红塘村3D高斯实景，可旋转、缓慢缩放、平移、回到模型和全屏查看"
        break

technology_table = document.tables[2]
for row in technology_table.rows:
    if row.cells[0].text.strip() == "3D实景":
        row.cells[2].text = (
            "在线加载红塘村高斯模型；长期访问令牌只保存在服务端环境变量中；"
            "滚轮缩放速度约为Cesium默认值的27%"
        )
        break

gaussian_heading = find_paragraph(document, "10.4 3D高斯实景首页")
first_gaussian_bullet = find_paragraph(
    document,
    "外层“3D高斯展示”文件夹保存原在线查看页和建筑轮廓参考；实际高斯模型是Cesium ion在线资产，不复制进源码仓库。",
)
first_gaussian_bullet._p.addprevious(gaussian_heading._p)
find_paragraph(document, "10.5 无人机正射影像").insert_paragraph_before(
    "相机控制参数为zoomFactor=1.35、inertiaZoom=0.65、maximumMovementRatio=0.05；"
    "滚轮缩放约为Cesium默认速度的27%，减少一次滚动跨越过大距离的情况。",
    style="List Bullet",
)

version_table = document.tables[4]
version_row = add_row_like(version_table, version_table.rows[-1])
for cell, value in zip(
    version_row.cells,
    [
        "V1.13 Demo",
        "2026-07-23",
        "降低3D首页鼠标滚轮缩放速度、缩放惯性和单帧移动上限，使逐格缩放更细腻；同步更新操作说明与技术参数。",
    ],
):
    cell.text = value

normalize_document(document)
document.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    assert archive.testzip() is None

check = Document(OUTPUT)
assert len(check.inline_shapes) == 13
assert "V1.13 Demo" in check.paragraphs[5].text
assert len(check.tables[4].rows) == 15
heading_index = next(i for i, paragraph in enumerate(check.paragraphs) if paragraph.text.strip() == "10.4 3D高斯实景首页")
bullet_index = next(i for i, paragraph in enumerate(check.paragraphs) if paragraph.text.strip().startswith("外层“3D高斯展示”"))
assert heading_index < bullet_index
assert any("滚轮缓慢缩放" in paragraph.text for paragraph in check.paragraphs)
assert any("zoomFactor=1.35" in paragraph.text for paragraph in check.paragraphs)

os.replace(OUTPUT, MANUAL)
print(f"Updated manual to V1.13: {MANUAL}")
