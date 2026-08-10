from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        set_body_font(run)


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.66 ｜ 更新日期：2026年8月10日",
)

for paragraph in document.paragraphs:
    text = paragraph.text
    updated = text.replace(
        "“高德地图”“卫星遥感”“无人机影像”“手绘图”四种底图按钮",
        "“底图”“卫星遥感”“无人机影像”“手绘图”四种按钮",
    ).replace(
        "可选择“高德地图”“卫星遥感”“无人机影像”或“手绘图”",
        "可选择“底图”“卫星遥感”“无人机影像”或“手绘图”",
    ).replace(
        "在高德地图、卫星遥感、无人机影像和手绘图四种模式下",
        "在底图、卫星遥感、无人机影像和手绘图四种模式下",
    )
    if updated != text:
        replace_paragraph_text(paragraph, updated)

topic_intro = next(p for p in document.paragraphs if p.text.startswith("桌面端和手机端都在地图左上角显示同一张白色圆角“专题”卡片"))
if "平台名称卡片、专题入口和展开后的专题面板保持相同宽度" not in topic_intro.text:
    replace_paragraph_text(
        topic_intro,
        topic_intro.text + " 桌面端的平台名称卡片、专题入口和展开后的专题面板保持相同宽度，左右边缘对齐。",
    )

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.66"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.66",
    "2026年8月10日",
    "二维底图列表中的“高德地图”简化为“底图”并缩窄卡片；左侧平台名称、专题入口与展开面板统一宽度。",
)
for cell, value in zip(new_row.cells, values):
    replace_paragraph_text(cell.paragraphs[0], value)

document.save(MANUAL)

reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.66") for p in reloaded.paragraphs)
assert "“底图”“卫星遥感”“无人机影像”“手绘图”四种按钮" in paragraph_text
assert "可选择“底图”“卫星遥感”“无人机影像”或“手绘图”" in paragraph_text
assert "平台名称卡片、专题入口和展开后的专题面板保持相同宽度" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.66"
print(MANUAL)
