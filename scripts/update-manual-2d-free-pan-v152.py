from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"
OLD_VERSION = "版本：V1.51 ｜ 更新日期：2026年8月9日"
NEW_VERSION = "版本：V1.52 ｜ 更新日期：2026年8月9日"
OLD_2D_GUIDE = "2D模式与3D实景读取同一套数据，并把手绘图或无人机影像作为完整地图底层。点击小花园、茶厂等地点后，二维地图会平滑移动并放大到对应坐标，在点位旁弹出与3D相同结构的气泡详情；点击水源、供排水线路或片区也会聚焦到对象，同时突出与其关联的上下游和服务片区。图钉会保持清晰的屏幕尺寸，不随底图缩放而变得过大。没有资料的字段不会由系统虚构补齐。"
NEW_2D_GUIDE = "2D模式与3D实景读取同一套数据，并把手绘图或无人机影像作为完整地图底层。2D地图在任何缩放级别都可以按住鼠标左键向任意方向拖动；滚轮以鼠标所在位置为中心缩放，双击地图或点击“复位视图”可返回初始位置。平台会自动区分拖动与点击，移动地图时不会误开地点详情。点击小花园、茶厂等地点后，二维地图会平滑移动并放大到对应坐标，在点位旁弹出与3D相同结构的气泡详情；点击水源、供排水线路或片区也会聚焦到对象，同时突出与其关联的上下游和服务片区。图钉会保持清晰的屏幕尺寸，不随底图缩放而变得过大。没有资料的字段不会由系统虚构补齐。"


def find_paragraph(document: Document, exact: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == exact]
    if len(matches) != 1:
        raise ValueError(f"应找到1处段落，实际找到{len(matches)}处：{exact}")
    return matches[0]


def set_two_character_indent(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), "200")


def format_run(run, *, heading: bool = False) -> None:
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "黑体" if heading else "宋体")
    run.font.color.rgb = RGBColor(0, 0, 0)


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    format_run(paragraph.add_run(text), heading=paragraph.style.name.startswith(("Heading", "标题")))
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    format_run(paragraph.add_run(text))


document = Document(MANUAL)
replace_paragraph(find_paragraph(document, OLD_VERSION), NEW_VERSION)
replace_paragraph(find_paragraph(document, OLD_2D_GUIDE), NEW_2D_GUIDE)

faq_table = document.tables[4]
faq_row = faq_table.add_row()
set_cell_text(faq_row.cells[0], "2D地图怎样移动、缩放和复位")
set_cell_text(faq_row.cells[1], "把鼠标放在地图画面上，按住左键即可向任意方向拖动；滚动滚轮进行缩放；双击地图或点击“复位视图”返回初始位置。拖动地图不会打开地点详情。")

version_table = document.tables[5]
template_row = deepcopy(version_table.rows[1]._tr)
version_table.rows[0]._tr.addnext(template_row)
new_row = version_table.rows[1]
set_cell_text(new_row.cells[0], "V1.52")
set_cell_text(new_row.cells[1], "2026年8月9日")
set_cell_text(new_row.cells[2], "2D地图支持在任意缩放级别按住鼠标左键向任意方向拖动；保留鼠标位置中心缩放、双击复位和复位按钮，并增加拖动距离判定，避免移动地图时误触地点详情。")

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 用户已明确不需要LibreOffice渲染，因此只执行结构、内容与压缩包完整性校验。
reopened = Document(MANUAL)
paragraph_texts = [paragraph.text.strip() for paragraph in reopened.paragraphs]
assert NEW_VERSION in paragraph_texts
assert NEW_2D_GUIDE in paragraph_texts
assert reopened.tables[5].rows[1].cells[0].text.strip() == "V1.52"
assert any(row.cells[0].text.strip() == "2D地图怎样移动、缩放和复位" for row in reopened.tables[4].rows)
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(f"updated={MANUAL}")
print(f"version={reopened.tables[5].rows[1].cells[0].text.strip()} faq_rows={len(reopened.tables[4].rows)}")
