from copy import deepcopy
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"


def set_body_font(run) -> None:
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_two_character_indent(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), "200")


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    set_body_font(paragraph.add_run(text))
    if paragraph.style.name == "Normal":
        set_two_character_indent(paragraph)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_body_font(paragraph.add_run(text))


document = Document(MANUAL)
replace_paragraph_text(
    next(p for p in document.paragraphs if p.text.startswith("版本：V1.")),
    "版本：V1.73 ｜ 更新日期：2026年8月11日",
)

mode_paragraph = next(p for p in document.paragraphs if p.text.startswith("首页默认进入红塘村2D地图"))
text = mode_paragraph.text
old_motion = "切换到3D时，下排按钮通过高度与透明度过渡自然收起，返回2D时自然展开；"
new_motion = (
    "切换到3D时，下排按钮通过高度与透明度过渡自然收起，返回2D时自然展开。"
    "为避免按钮动画与地图初始化同时占用性能，平台会先用约260毫秒完成右上角卡片过渡，再挂载高德2D地图、航拍图层和点位；"
)
if old_motion in text:
    text = text.replace(old_motion, new_motion)
old_runtime = "切换时只替换底层地图，左上角“专题”组件始终保留在首页，已经勾选或取消的图层会继续生效；同时只运行当前地图，避免2D与3D同时占用显卡。"
new_runtime = "切换时左上角“专题”组件始终保留，已经勾选或取消的图层继续生效；短暂界面过渡结束后，原地图会被卸载，稳定状态下只运行当前地图，避免2D与3D长期同时占用显卡。"
if old_runtime in text:
    text = text.replace(old_runtime, new_runtime)
replace_paragraph_text(mode_paragraph, text)

version_table = next(table for table in document.tables if table.cell(0, 0).text == "版本")
new_row = next((row for row in version_table.rows if row.cells[0].text == "V1.73"), None)
if new_row is None:
    version_table.rows[0]._tr.addnext(deepcopy(version_table.rows[1]._tr))
    new_row = version_table.rows[1]
values = (
    "V1.73",
    "2026年8月11日",
    "优化2D/3D切换性能：右上角功能卡片先完成约260毫秒的展开或收起，再挂载目标地图，避免高德地图、航拍图层和点位初始化造成动画掉帧。",
)
for cell, value in zip(new_row.cells, values):
    set_cell_text(cell, value)

document.core_properties.modified = datetime.now()
document.save(MANUAL)

# 按项目约定不进行LibreOffice渲染，只执行内容、样式、结构和压缩包完整性检查。
reloaded = Document(MANUAL)
paragraph_text = "\n".join(p.text for p in reloaded.paragraphs)
assert any(p.text.startswith("版本：V1.73") for p in reloaded.paragraphs)
assert "先用约260毫秒完成右上角卡片过渡" in paragraph_text
assert "稳定状态下只运行当前地图" in paragraph_text
assert reloaded.tables[-1].rows[1].cells[0].text == "V1.73"
assert reloaded.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
assert reloaded.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(MANUAL)
print("version=V1.73")
print("zip=ok")
