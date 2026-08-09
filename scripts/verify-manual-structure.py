from pathlib import Path
import zipfile

from docx import Document
from docx.oxml.ns import qn


source_root = Path(__file__).resolve().parents[1]
manual = source_root.parent / "红塘村可持续发展平台使用手册.docx"

with zipfile.ZipFile(manual) as archive:
    assert archive.testzip() is None

document = Document(manual)
paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
table_text = "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)
first_line_count = len(document.element.body.xpath('.//w:ind[@w:firstLineChars="200"]'))

assert any(paragraph.text.startswith("版本：V1.58") for paragraph in document.paragraphs)
assert "首页默认进入红塘村2D地图" in paragraph_text
assert "高德在线地图" in paragraph_text
assert "整数屏幕像素定位" in paragraph_text
assert "npm run test:2d-markers" in paragraph_text
assert "原地图左下角的孤立入口已经删除" in paragraph_text
assert "优先读取同一个Supabase云端数据库" in paragraph_text
assert "hongtang-photos素材桶" in paragraph_text
assert "SUPABASE_SERVICE_KEY不得进入浏览器或Git仓库" in paragraph_text
assert "Supabase：platform_datasets表" in table_text
assert "2D图钉有时模糊，或点击后没有出现详情" in table_text

assert len(document.tables) == 6
assert [cell.text for cell in document.tables[0].rows[0].cells] == ["地点类型", "当前数量", "显示方式"]
assert [cell.text for cell in document.tables[4].rows[0].cells] == ["现象", "处理方法"]
assert [cell.text for cell in document.tables[5].rows[0].cells] == ["版本", "日期", "主要变化"]
assert document.tables[5].rows[1].cells[0].text == "V1.58"
assert first_line_count >= 18
assert document.styles["Normal"].element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
assert document.styles["Normal"].font.name == "Times New Roman"
assert document.styles["Heading 1"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
assert document.styles["Heading 2"].element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"

print(
    {
        "version": next(p.text for p in document.paragraphs if p.text.startswith("版本：")),
        "paragraphs": len(document.paragraphs),
        "table_rows": [len(table.rows) for table in document.tables],
        "native_two_character_indents": first_line_count,
        "zip": "ok",
    }
)

