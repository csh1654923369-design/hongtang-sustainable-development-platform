from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANUAL = ROOT / "红塘村可持续发展平台使用手册.docx"
ORIGINAL = "排水到哪里：只显示2个排水节点和5条排水线路，回答“房前屋后的水在哪里汇集、怎样流动、最终到哪里”。"
UPDATED = "排水到哪里：显示2个待核实排水节点、5条待核实排水线路，并保留2个已有污水设施点位，回答“房前屋后的水在哪里汇集、怎样流动、最终到哪里”。"


document = Document(MANUAL)
matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == ORIGINAL]
if len(matches) != 1:
    raise ValueError(f"应找到1处待更新说明，实际找到{len(matches)}处")

paragraph = matches[0]
paragraph.clear()
run = paragraph.add_run(UPDATED)
run.font.name = "Times New Roman"
run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
run.font.color.rgb = RGBColor(0, 0, 0)
document.save(MANUAL)

reopened = Document(MANUAL)
assert sum(paragraph.text.strip() == UPDATED for paragraph in reopened.paragraphs) == 1
with ZipFile(MANUAL) as archive:
    assert archive.testzip() is None

print(f"updated={MANUAL}")
